import datetime
from datetime import date
from django.utils import timezone
import re
import pymorphy2
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages
from django.db.models import Q
from django.http import HttpResponseRedirect, HttpResponse, FileResponse, JsonResponse
from django.shortcuts import render
from django.utils.encoding import uri_to_iri, escape_uri_path
from django.views.generic import ListView, CreateView, UpdateView
from docx import Document
from docx.shared import Pt
from docxtpl import DocxTemplate
from pandas._libs.tslibs.offsets import BDay
from authen.middleware import GetUserDepAD, GetDepsWithManagerFromAD
from centre.views import group_required, is_ajax, month_from_ru_to_eng, get_students, get_teachers, CourseGroupCode, \
    get_lessons
from config.settings import MEDIA_ROOT
from students.models import Apps, Statuses, EventsForms, Docs, DocsTypes
from .forms import ProgramForm, CoursesForm, EventsForm
from centre.models import Programs, StSchedule, Events, CourseLessons, \
    EventsLessons, Mos, OoTypes, Oos, PositionCategories, EventTypes, PlanningParameters
from centre.tasks import *


class CheckDepMixin(LoginRequiredMixin):
    "Проверка на наличие пользователя в группе Работники центра"
    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Работники центра').exists():
            return HttpResponseRedirect('/access_denied/')
        return super().dispatch(request, *args, **kwargs)


class CheckTeacherMixin(LoginRequiredMixin):
    "Проверка на преподавателя"
    def dispatch(self, request, *args, **kwargs):
        if Profiles.objects.get(user_id=request.user.id).teacher is False:
            return HttpResponseRedirect('/access_denied/')
        return super().dispatch(request, *args, **kwargs)


class ProgramsListView(CheckDepMixin, ListView):
    "Работа с ДПП подразделения пользователя"
    login_url = '/'
    model = Programs
    context_object_name = 'programs'
    template_name = 'dep/study/programs/list.html'

    def get_queryset(self):
        dep = GetUserDepAD(self.request)
        qs = Programs.objects.filter(department=dep).order_by('-id')
        if 'check_deps[]' in self.request.GET:
            chk = self.request.GET.getlist('check_deps[]')
            if len(chk) != 0:
                qs = qs.filter(department__in=chk)
        if 'check_cats[]' in self.request.GET:
            chk = self.request.GET.getlist('check_cats[]')
            if len(chk) != 0:
                qs = qs.filter(categories__in=chk)
        if 'name' in self.request.GET:
            qs = qs.filter(name__contains=self.request.GET.get('name'))
        if 'order_id' in self.request.GET:
            qs = qs.filter(order_id__contains=self.request.GET.get('order_id'))
        return qs[:25]

    def post(self, request, *args, **kwargs):
        if 'order_view' in request.POST:
            file = Programs.objects.get(id=request.POST.get('order_view')).order_file
            if file.path[-3:] == 'pdf':
                return FileResponse(open(file.path, 'rb'), content_type='application/pdf')
            else:
                return FileResponse(open(file.path, 'rb'))
        if 'delete_dpp' in request.POST:
            program = Programs.objects.get(id=request.POST.get('delete_dpp'))
            if program.order_id == '' and program.order_date is None:
                try:
                    program.delete()
                    messages.success(request, 'Программа успешно удалена!')
                except BaseException:
                    messages.error(request, 'Произошла ошибка, попробуйте позже')
                return HttpResponseRedirect('/dep/study/programs')
            else:
                return HttpResponseRedirect('/access_denied/')

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(ProgramsListView, self).get_context_data(**kwargs)
        context['dep'] = GetUserDepAD(self.request)
        return context


class ProgramDetailView(CheckDepMixin, UpdateView):
    "Изменение ДПП подразделения пользователя"
    login_url = '/'
    form_class = ProgramForm
    context_object_name = 'program'
    template_name = 'dep/study/programs/edit.html'
    success_url = '/dep/study/programs'

    def get_object(self, queryset=None):
        program = Programs.objects.get(id=self.kwargs.get('pk'))
        if program.department == GetUserDepAD(self.request) and program.order_date is None and program.order_id == '':
            return program
        else:
            return None

    def form_invalid(self, form):
        messages.error(self.request, 'Произошла ABOBA во время сохранения ДПП. Повторите попытку позже')
        return HttpResponseRedirect('/dep/study/programs')

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Информация успешно обновлена')
        return HttpResponseRedirect('/dep/study/programs')


class ProgramCreateView(CheckDepMixin, CreateView):
    "Внесение новой ДПП подразделения пользователя"
    login_url = '/'
    model = Programs
    form_class = ProgramForm
    context_object_name = 'program'
    template_name = 'dep/study/programs/detail.html'

    def get_form_kwargs(self):
        kwargs = super(ProgramCreateView, self).get_form_kwargs()
        kwargs.update({'dep': GetUserDepAD(self.request)})
        return kwargs

    def form_invalid(self, form):
        messages.error(self.request, 'Произошла ошибка во время сохранения ДПП. Повторите попытку позже')
        return HttpResponseRedirect('/dep/study/programs')

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Новая программа успешно добавлена')
        return HttpResponseRedirect('/dep/study/programs')


def CheckSumHours(id, total, lecture, practice, trainee, individual):
    "Функция контроля количества часов при добавлении темы в разделе"
    module = StSchedule.objects.get(id=id)
    if StSchedule.objects.filter(parent_id=id).exists():
        ch_total = 0
        ch_lecture = 0
        ch_practice = 0
        ch_trainee = 0
        ch_individual = 0
        childrens = StSchedule.objects.filter(parent_id=id)
        for child in childrens:
            ch_total += child.total_hours
            ch_lecture += child.lecture_hours
            ch_practice += child.practice_hours
            ch_trainee += child.trainee_hours
            ch_individual += child.individual_hours
        if module.total_hours < ch_total + total:
            return 'total'
        if module.lecture_hours < ch_lecture + lecture:
            return 'lecture'
        if module.practice_hours < ch_practice + practice:
            return 'practice'
        if module.trainee_hours < ch_trainee + trainee:
            return 'trainee'
        if module.individual_hours < ch_individual + individual:
            return 'individual'
    else:
        if module.total_hours < total:
            return 'total'
        if module.lecture_hours < lecture:
            return 'lecture'
        if module.practice_hours < practice:
            return 'practice'
        if module.trainee_hours < trainee:
            return 'trainee'
        if module.individual_hours < individual:
            return 'individual'
    return None


def CheckSumHoursEdit(id, exec, total, lecture, practice, trainee, individual):
    "Функция контроля количества часов при изменении темы в разделе"
    module = StSchedule.objects.get(id=id)
    if StSchedule.objects.filter(parent_id=id).exists():
        ch_total = 0
        ch_lecture = 0
        ch_practice = 0
        ch_trainee = 0
        ch_individual = 0
        childrens = StSchedule.objects.filter(parent_id=id)
        for child in childrens:
            if child.id != exec:
                ch_total += child.total_hours
                ch_lecture += child.lecture_hours
                ch_practice += child.practice_hours
                ch_trainee += child.trainee_hours
                ch_individual += child.individual_hours
        if module.total_hours < ch_total + total:
            return 'total'
        if module.lecture_hours < ch_lecture + lecture:
            return 'lecture'
        if module.practice_hours < ch_practice + practice:
            return 'practice'
        if module.trainee_hours < ch_trainee + trainee:
            return 'trainee'
        if module.individual_hours < ch_individual + individual:
            return 'individual'
    else:
        if module.total_hours < total:
            return 'total'
        if module.lecture_hours < lecture:
            return 'lecture'
        if module.practice_hours < practice:
            return 'practice'
        if module.trainee_hours < trainee:
            return 'trainee'
        if module.individual_hours < individual:
            return 'individual'
    return None


class KugListView(CheckDepMixin, ListView):
    "Работа с КУГ ДПП подразделения пользователя"
    login_url = '/'
    model = StSchedule
    context_object_name = 'elements'
    template_name = 'dep/study/programs/kug/edit.html'

    def get(self, request, *args, **kwargs):
        if 'close_kug' in request.GET:
            dpp = Programs.objects.get(id=self.kwargs.get('pk'))
            dpp.kug_on_edit = ''
            dpp.save()
            return HttpResponseRedirect('/dep/study/programs')
        if Programs.objects.get(id=self.kwargs.get('pk')).department == GetUserDepAD(self.request):
            dpp = Programs.objects.get(id=self.kwargs.get('pk'))
            prof = Profiles.objects.get(user_id=request.user.id)
            if len(prof.surname) == 0 or len(prof.name) == 0 or len(prof.patronymic) == 0:
                messages.error(request, 'Пожалуйста, заполните ФИО в профиле для доступа к КУГ')
                return HttpResponseRedirect('/dep/study/programs')
            FIO = prof.surname + ' ' + prof.name[:1] + '.' + prof.patronymic[:1] + '.'
            UserKugEdit = dpp.kug_on_edit
            if len(UserKugEdit) != 0 and UserKugEdit != FIO:
                messages.error(request, 'КУГ находится на редактировании у пользователя "' + UserKugEdit + '"')
                return HttpResponseRedirect('/dep/study/programs')
            if dpp.order_id != '' and dpp.order_date is not None and dpp.order_file.path is not None:
                messages.error(request, 'Программа уже утверждена, внесение изменений запрещено')
                return HttpResponseRedirect('/dep/study/programs')
            dpp.kug_on_edit = prof.surname + ' ' + prof.name[:1] + '.' + prof.patronymic[:1] + '.'
            dpp.save()
            return render(request, 'dep/study/programs/kug/edit.html', {
                'elements': StSchedule.objects.filter(program_id=dpp.id),
                'dpp': Programs.objects.get(id=self.kwargs.get('pk'))
            })
        else:
            return None

    def post(self, request, *args, **kwargs):
        if 'new_module' in request.POST:
            total = int(request.POST.get('total_hours'))
            lecture = int(request.POST.get('lecture_hours'))
            practice = int(request.POST.get('practice_hours'))
            trainee = int(request.POST.get('trainee_hours'))
            individual = int(request.POST.get('individual_hours'))
            if lecture + practice + trainee + individual != total:
                messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
                return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('new_module')))
            #try:
            count = StSchedule.objects.filter(program_id=int(request.POST.get('new_module'))).filter(level=0).count()
            new = StSchedule()
            new.program_id = int(request.POST.get('new_module'))
            new.name = 'Раздел ' + str(count + 1) + '. ' + request.POST.get('name')
            new.total_hours = request.POST.get('total_hours')
            new.lecture_hours = request.POST.get('lecture_hours')
            new.practice_hours = request.POST.get('practice_hours')
            new.trainee_hours = request.POST.get('trainee_hours')
            new.individual_hours = request.POST.get('individual_hours')
            new.control_form = request.POST.get('control_form')
            new.save()
            messages.success(request, 'Раздел успешно добавлен')
            #except BaseException:
               # messages.error(request, 'Произошла ошибка при сохранении, повторите попытку позже')
            return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('new_module')))
        elif 'delete_el' in request.POST:
            try:
                el = StSchedule.objects.get(id=int(request.POST.get('delete_el')))
                if StSchedule.objects.filter(parent_id=el.parent_id).count() > 1:
                    children = StSchedule.objects.filter(parent_id=el.parent_id)
                    for child in children:
                        if child.lft > el.rght:
                            name = child.name
                            theme_numb = name[name.find('.'):name.find('.', name.find('.') + 1)]
                            newnumb = int(theme_numb[1:]) - 1
                            newname = name[:name.find('.')] + '.' + str(newnumb) + name[name.find('.', name.find('.') + 1):]
                            child.name = newname
                            child.save()
                if el.parent_id is None:
                    if StSchedule.objects.filter(tree_id__gt=el.tree_id).exists():
                        other = StSchedule.objects.filter(tree_id__gt=el.tree_id).filter(level=0)
                        for ot in other:
                            children = StSchedule.objects.filter(parent_id=ot.id)
                            for child in children:
                                name = child.name
                                theme_numb = name[name.find(' '):name.find('.')]
                                newnumb = int(theme_numb[1:]) - 1
                                newname = name[:name.find(' ') + 1] + str(newnumb) + name[name.find('.'):]
                                child.name = newname
                                child.save()
                            name = ot.name
                            module_numb = name[name.find(' '):name.find('.')]
                            newnumb = int(module_numb) - 1
                            newname = name[:name.find(' ') + 1] + str(newnumb) + name[name.find('.'):]
                            ot.name = newname
                            ot.save()
                    messages.success(request, 'Раздел успешно удален')
                else:
                    messages.success(request, 'Тема успешно удалена')
                StSchedule.objects.get(id=int(request.POST.get('delete_el'))).delete()
            except BaseException:
                messages.error(request, 'Произошла ошибка во время удаления, повторите попытку позже')
            return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('dpp_id')))
        elif 'change_module' in request.POST:
            try:
                total = int(request.POST.get('total_hours'))
                lecture = int(request.POST.get('lecture_hours'))
                practice = int(request.POST.get('practice_hours'))
                trainee = int(request.POST.get('trainee_hours'))
                individual = int(request.POST.get('individual_hours'))
                if lecture + practice + individual + trainee != total:
                    messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
                    return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('dpp')))
                module = StSchedule.objects.get(id=request.POST.get('change_module'))
                name = module.name
                module.name = name[:name.find('.') + 2] + request.POST.get('name')
                module.total_hours = request.POST.get('total_hours')
                module.lecture_hours = request.POST.get('lecture_hours')
                module.practice_hours = request.POST.get('practice_hours')
                module.trainee_hours = request.POST.get('trainee_hours')
                module.individual_hours = request.POST.get('individual_hours')
                module.control_form = request.POST.get('control_form')
                module.save()
                messages.success(request, 'Информация о разделе успешно обновлена')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
            return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('dpp')))
        elif 'new_theme' in request.POST:
            total = int(request.POST.get('total_hours'))
            lecture = int(request.POST.get('lecture_hours'))
            practice = int(request.POST.get('practice_hours'))
            trainee = int(request.POST.get('trainee_hours'))
            individual = int(request.POST.get('individual_hours'))
            if lecture + practice + trainee + individual != total:
                messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
                return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('new_theme')))
            check = CheckSumHours(int(request.POST.get('module_id')), total, lecture, practice, trainee, individual)
            if check is not None:
                if check == 'total':
                    messages.error(request, 'Общая сумма часов всех тем превышает указанное в разделе значение')
                elif check == 'lecture':
                    messages.error(request, 'Сумма лекционных часов всех тем превышает указанное в разделе значение')
                elif check == 'practice':
                    messages.error(request, 'Сумма часов практики всех тем превышает указанное в разделе значение')
                elif check == 'trainee':
                    messages.error(request, 'Сумма часов стажировок всех тем превышает указанное в разделе значение')
                else:
                    messages.error(request,
                                   'Сумма часов индивидуальных занятий всех тем превышает указанное в разделе значение')
                return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('new_theme')))
            try:
                count = StSchedule.objects.filter(program_id=int(request.POST.get('new_theme'))) \
                    .filter(parent_id=int(request.POST.get('module_id'))).count()
                module_number = StSchedule.objects.get(id=int(request.POST.get('module_id'))).name
                new = StSchedule()
                new.program_id = int(request.POST.get('new_theme'))
                new.name = 'Тема' + module_number[module_number.find(' '):module_number.find('.')] + '.' + str(
                    count + 1) + '. ' + request.POST.get('name')
                new.total_hours = request.POST.get('total_hours')
                new.lecture_hours = request.POST.get('lecture_hours')
                new.practice_hours = request.POST.get('practice_hours')
                new.trainee_hours = request.POST.get('trainee_hours')
                new.individual_hours = request.POST.get('individual_hours')
                new.control_form = request.POST.get('control_form')
                new.parent_id = int(request.POST.get('module_id'))
                new.save()
                messages.success(request, 'Тема успешно добавлена')
            except BaseException:
                messages.error(request, 'Произошла ошибка при сохранении, повторите попытку позже')
            return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('new_theme')))
        elif 'change_theme' in request.POST:
            total = int(request.POST.get('total_hours'))
            lecture = int(request.POST.get('lecture_hours'))
            practice = int(request.POST.get('practice_hours'))
            trainee = int(request.POST.get('trainee_hours'))
            individual = int(request.POST.get('individual_hours'))
            if lecture + practice + trainee + individual != total:
                messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
                return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('dpp_id')))
            id_parent = StSchedule.objects.get(id=int(request.POST.get('change_theme'))).parent_id
            check = CheckSumHoursEdit(id_parent, int(request.POST.get('change_theme')), total, lecture, practice, trainee,  individual)
            if check is not None:
                if check == 'total':
                    messages.error(request, 'Общая сумма часов всех тем превышает указанное в разделе значение')
                elif check == 'lecture':
                    messages.error(request, 'Сумма лекционных часов всех тем превышает указанное в разделе значение')
                elif check == 'practice':
                    messages.error(request, 'Сумма часов практики всех тем превышает указанное в разделе значение')
                elif check == 'trainee':
                    messages.error(request, 'Сумма часов стажировок всех тем превышает указанное в разделе значение')
                else:
                    messages.error(request,
                                   'Сумма часов индивидуальных занятий всех тем превышает указанное в разделе значение')
                return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('dpp_id')))
            try:
                theme = StSchedule.objects.get(id=request.POST.get('change_theme'))
                name = theme.name
                theme.name = name[:name.find('.', name.find('.') + 1) + 2] + request.POST.get('name')
                theme.total_hours = request.POST.get('total_hours')
                theme.lecture_hours = request.POST.get('lecture_hours')
                theme.practice_hours = request.POST.get('practice_hours')
                theme.trainee_hours = request.POST.get('trainee_hours')
                theme.individual_hours = request.POST.get('individual_hours')
                theme.control_form = request.POST.get('control_form')
                theme.save()
                messages.success(request, 'Информация о теме успешно обновлена')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
            return HttpResponseRedirect('/dep/study/kug_' + str(request.POST.get('dpp_id')))
        elif 'edit_finish' in request.POST:
            try:
                dpp = Programs.objects.get(id=int(request.POST.get('edit_finish')))
                dpp.kug_on_edit = ''
                dpp.save()
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        else:
            messages.error(request, 'Действие не найдено')
        return HttpResponseRedirect('/dep/study/programs')


class CoursesList(CheckDepMixin, ListView):
    "Работа с курсами на основе ДПП подразлеления пользователя"
    login_url = '/'
    model = Courses
    context_object_name = 'courses'
    template_name = 'dep/study/planning/courses/list.html'

    def get_queryset(self):
        qs = Courses.objects.filter(program__in=Programs.objects.filter(department=GetUserDepAD(self.request))).\
            order_by('-id')
        if 'name' in self.request.GET:
            if len(self.request.GET.get('name')) > 0:
                qs = qs.filter(program__in=Programs.objects.filter(name__contains=self.request.GET.get('name')))
        if 'duration' in self.request.GET:
            if len(self.request.GET.get('duration')) > 0:
                qs = qs.filter(program__in=Programs.objects.filter(duration=self.request.GET.get('duration')))
        if 'study_date' in self.request.GET:
            try:
                post = self.request.GET.get('study_date')
                dt = date(int(post[:4]), int(post[-5:7]), int(post[8:]))
                qs = qs.filter(date_start__lte=date, date_finish__gte=dt)
            except BaseException:
                pass
        return qs[:25]

    def post(self, request):
        if 'change_course' in request.POST:
            try:
                course = Courses.objects.get(id=request.POST.get('change_course'))
                course.place = request.POST.get('place')
                course.date_start = request.POST.get('date_start')
                course.date_finish = request.POST.get('date_finish')
                course.save()
                messages.success(request, 'Информация о курсе успешно обновлена')
            except BaseException:
                messages.success(request, 'Произошла ошибка, повторите попытку позже')
        elif 'delete_course' in request.POST:
            try:
                Courses.objects.get(id=request.POST.get('delete_course')).delete()
                messages.success(request, 'Курс успешно удален')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        else:
            messages.error(request, 'Действие не найдено')
        return HttpResponseRedirect('/dep/study/planning/courses')

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(CoursesList, self).get_context_data(**kwargs)
        context['dep_name'] = GetUserDepAD(self.request)
        return context


class CoursesCreate(CheckDepMixin, CreateView):
    "Создание нового курса на основе ДПП подразделения пользователя"
    login_url = '/'
    model = Courses
    form_class = CoursesForm
    context_object_name = 'course'
    template_name = 'dep/study/planning/courses/detail.html'

    def get_form_kwargs(self):
        kwargs = super(CoursesCreate, self).get_form_kwargs()
        kwargs.update({'dep': GetUserDepAD(self.request)})
        return kwargs

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'al' in request.GET:
                return JsonResponse({
                    'val': PlanningParameters.objects.get(alias=request.GET.get('al')).value
                })
            duration = Programs.objects.get(id=request.GET.get('id')).duration
            type_dpp = Programs.objects.get(id=request.GET.get('id')).type_dpp
            date_order = Programs.objects.get(id=request.GET.get('id')).order_date.strftime('%d.%m.%Y')
            return JsonResponse({
                'duration': str(duration),
                'type_dpp': type_dpp,
                'date_order': date_order
            })
        else:
            return super(CoursesCreate, self).get(request)

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Курс успешно сохранен')
        return HttpResponseRedirect('/dep/study/planning/courses')

    def form_invalid(self, form):
        messages.error(self.request, 'Произошла ошибка во время сохранения курса, повторите попытку позже')
        return HttpResponseRedirect('/dep/study/planning/courses')


class EventsList(CheckDepMixin, ListView):
    "Работа с мероприятиями подразделения пользователя"
    login_url = '/'
    model = Events
    context_object_name = 'events'
    template_name = 'dep/study/planning/events/list.html'

    def get_queryset(self):
        qs = Events.objects.filter(department=GetUserDepAD(self.request)).order_by('-id')
        if 'name' in self.request.GET:
            if len(self.request.GET.get('name')) > 0:
                qs = qs.filter(name__contains=self.request.GET.get('name'))
        if 'duration' in self.request.GET:
            if len(self.request.GET.get('duration')) > 0:
                qs = qs.filter(duration=int(self.request.GET.get('duration')))
        if 'study_date' in self.request.GET:
            try:
                post = self.request.GET.get('study_date')
                dt = date(int(post[:4]), int(post[-5:7]), int(post[8:]))
                qs = qs.filter(date_start__lte=date, date_finish__gte=dt)
            except BaseException:
                pass
        if 'check_deps[]' in self.request.GET:
            chk = self.request.GET.getlist('check_deps[]')
            if len(chk) != 0:
                qs = qs.filter(department__in=chk)
        if 'check_cats[]' in self.request.GET:
            chk = self.request.GET.getlist('check_cats[]')
            if len(chk) != 0:
                qs = qs.filter(categories__in=chk)
        return qs[:25]

    def post(self, request, *args, **kwargs):
        if 'delete_id' in request.POST:
            try:
                Events.objects.get(id=request.POST.get('delete_id')).delete()
                messages.success(request, 'Мероприятие успешно удалено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        else:
            messages.error(request, 'Действие не найдено')
        return HttpResponseRedirect('/dep/study/planning/events')

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(EventsList, self).get_context_data(**kwargs)
        context['dep_name'] = GetUserDepAD(self.request)
        return context


class EventDetailView(CheckDepMixin, UpdateView):
    "Изменение мероприятия подразделения пользователя"
    login_url = '/'
    form_class = EventsForm
    context_object_name = 'event'
    template_name = 'dep/study/planning/events/detail.html'
    success_url = '/dep/study/planning/events'

    def get_object(self, queryset=None):
        event = Events.objects.get(id=self.kwargs.get('pk'))
        if event.department == GetUserDepAD(self.request):
            return event
        else:
            return None

    def form_invalid(self, form):
        messages.error(self.request, 'Произошла ошибка во время сохранения мероприятия. Повторите попытку позже')
        return HttpResponseRedirect('/dep/study/planning/events')

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Информация о мероприятии успешно обновлена')
        return HttpResponseRedirect('/dep/study/planning/events')


class EventCreateView(CheckDepMixin, CreateView):
    "Внесение нового мероприятия подразделения пользователя"
    login_url = '/'
    model = Events
    form_class = EventsForm
    context_object_name = 'event'
    template_name = 'dep/study/planning/events/detail.html'

    def get_form_kwargs(self):
        kwargs = super(EventCreateView, self).get_form_kwargs()
        kwargs.update({'dep': GetUserDepAD(self.request)})
        return kwargs

    def form_invalid(self, form):
        messages.error(self.request, 'Произошла ошибка во время сохранения мероприятия. Повторите попытку позже')
        return HttpResponseRedirect('/dep/study/planning/events')

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Новое мероприятие успешно добавлено')
        return HttpResponseRedirect('/dep/study/planning/events')

    def get_context_data(self, **kwargs):
        context = super(EventCreateView, self).get_context_data()
        context['event'] = 'create'
        return context


class StudentsGroupList(CheckDepMixin, ListView):
    "Работа с учебными группами подразделения пользователя"
    login_url = '/'
    model = StudentGroups
    context_object_name = 'groups'
    template_name = 'dep/study/studentgroups/list.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            teacher = Profiles.objects.get(id=request.GET.get('id_teach'))
            return JsonResponse({
                'fio': teacher.surname+' '+teacher.name+' '+teacher.patronymic,
                'email': User.objects.get(id=teacher.user_id).email,
                'phone': teacher.phone
            })
        if 'cert_list' in request.GET:
            group = StudentGroups.objects.select_related('course').get(id=request.GET.get('cert_list'))
            dep = group.course.program.department
            date_start = group.course.date_start
            date_finish = group.course.date_finish
            if group.course.program.type_dpp == 'Повышение квалификации':
                type_dpp = 'повышения квалификации'
            else:
                type_dpp = 'профессиональной переподготовки'
            deps = GetDepsWithManagerFromAD()
            for key, value in deps.items():
                if key == dep:
                    manager = value
            if group.course.program.type_dpp == 'Повышение квалификации':
                type_dpp = 'повышении квалификации'
            else:
                type_dpp = 'профессиональной переподготовке'
            path = STATIC_ROOT + '\\doc_templates\\xlsx\\cert_list.xlsx'
            writer = BookWriter(path)
            info = {
                'dep': dep,
                'type_dpp': type_dpp,
                'name_dpp': group.course.program.name,
                'duration': group.course.program.duration,
                'code': group.code,
                'manager': manager,
                'day_start': date_start.strftime('%d'),
                'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                'year_start': date_start.strftime('%Y'),
                'day_finish': date_finish.strftime('%d'),
                'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                'year_finish': date_finish.strftime('%Y'),
                'date_get': group.course.date_finish.strftime('%d.%m.%Y'),
                'date_enroll': group.date_enroll.strftime('%d.%m.%Y'),
                'date_expl': group.date_exp.strftime('%d.%m.%Y'),
                'number_enroll': group.enroll_number,
                'number_expl': group.exp_number,
            }
            info['students'] = get_students(group.id)
            writer.render_sheet(info, 'Ведомость выдачи удостоверений', 0)
            newpath = MEDIA_ROOT + '\\Ведомости\\' + group.code
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            writer.save(newpath + "\\ВыдачаУдостоверений_.xlsx")
            filename = newpath + "\\ВыдачаУдостоверений_.xlsx"
            strin = "ВыдачаУдостоверений_" + group.code + ".xlsx"
            data = open(filename, "rb").read()
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Length'] = os.path.getsize(filename)
            response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
            return response
        qs = StudentGroups.objects.\
            filter(Q(event__in=Events.objects.filter(department=GetUserDepAD(self.request)))|
                   Q(course__in=Courses.objects.filter(program__in=Programs.objects.filter(department=GetUserDepAD(self.request))))).\
            order_by('-id')
        if 'code' in request.GET:
            if len(request.GET.get('code')) > 0:
                qs = qs.filter(code__contains=request.GET.get('code'))
        if 'name' in request.GET:
            if len(request.GET.get('name')) > 0:
                events = Events.objects.filter(name__contains=request.GET.get('name'))
                courses = Courses.objects.filter(program__in=Programs.objects.filter(name__contains=request.GET.get('name')))
                qs = qs.filter(Q(event__in=events) | Q(course__in=courses))
        if 'type' in request.GET:
            typ = uri_to_iri(request.GET.get('type'))
            if typ in ['Повышение квалификации', 'Профессиональная переподготовка']:
                qs = qs.filter(event_id=None).filter(course__in=Courses.objects.filter(program__in=Programs.objects.filter(type_dpp=request.GET.get('type'))))
            elif request.GET.get('type') == 'all':
                pass
            else:
                qs = qs.filter(event__in=Events.objects.filter(type_id=request.GET.get('type')))
        if 'grstatuses' in request.GET:
            qs = qs.filter(status_id__in=request.GET.getlist('grstatuses'))
            return render(request, 'dep/study/studentgroups/list.html', context={
                'groups': qs[:25],
                'filter_status': 'yes',
                'dep_name': GetUserDepAD(self.request)
            })
        return render(request, 'dep/study/studentgroups/list.html', context={
            'groups': qs[:25],
            'dep_name': GetUserDepAD(self.request)
        })

    def post(self, request, *args, **kwargs):
        #try:
        if is_ajax(request=request):
            group = StudentGroups.objects.get(id=request.POST.get('id_group'))
            if 'url_study' in request.POST:
                group.event_url = request.POST.get('url_study')
                group.status = StGroupStatuses.objects.get(name='Идет обучение')
                group.save()
                messages.success(request, 'Ссылка успешно сохранена')
                return JsonResponse({})
            if 'url_survey' in request.POST:
                group.survey_url = request.POST.get('url_survey')
                group.save()
                messages.success(request, 'Ссылка на опрос успешно изменена')
                return JsonResponse({})
            else:
                group.offer = request.FILES.get('file')
                group.status = StGroupStatuses.objects.get(name='Ожидает ссылку на обучение')
                group.save()
                students = group.students.all()
                emails = []
                for student in students:
                    app = Apps.objects.filter(group_id=request.POST.get('id_group')).get(profile_id=student.id)
                    app.status_id = Statuses.objects.get(name='Ждем оплату').id
                    app.save()
                    emails.append(student.user.email)
                if group.event is None:
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        typ = 'курсе повышения квалификации'
                    else:
                        typ = 'курсе профессиональной переподготовки'
                    name = group.course.program.name
                    start = group.course.date_start
                else:
                    type_classic = group.event.type.name
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(type_classic)[0]
                    word = type.inflect({'loct'})
                    typ = word.word
                    name = group.event.name
                    start = group.event.date_start
                pay_date = start - BDay(5)
                pay_deadline = pay_date.strftime('%d')+' '+month_from_ru_to_eng(pay_date.strftime('%B'))+' '+pay_date.strftime('%Y')
                EmailOfferPay.delay(emails, typ, name, pay_deadline)
                messages.success(request, 'Скан договора оферты успешно загружен!')
                return JsonResponse({})
        else:
            if 'approve_list' in request.POST:
                group = StudentGroups.objects.get(id=request.POST.get('approve_list'))
                group.status = StGroupStatuses.objects.get(name='Ожидает подгрузку скана договора оферты')
                group.save()
                messages.success(request, 'Состав группы успешно утвержден')
            elif 'change_list' in request.POST:
                group = StudentGroups.objects.get(id=request.POST.get('change_list'))
                if group.event is None:
                    group.status = StGroupStatuses.objects.get(name='Идет регистрация')
                    group.save()
                    messages.success(request, 'Регистрация в учебную группу успешно открыта')
                else:
                    count = group.students.count()
                    if count >= group.students_number:
                        messages.error(request, 'В учебной группе достигнуто плановое количество мест')
                    else:
                        group.status = StGroupStatuses.objects.get(name='Идет регистрация')
                        group.save()
                        messages.success(request, 'Регистрация в учебную группу успешно открыта')
            elif 'order_enroll' in request.POST:
                group = StudentGroups.objects.get(id=request.POST.get('order_enroll'))
                start = group.course.date_start
                finish = group.course.date_finish
                doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/order_enroll_ou.docx")
                context = {
                    'prog_name': group.course.program.name,
                    'duration': group.course.program.duration,
                    'day_start': start.strftime('%d'),
                    'month_start': month_from_ru_to_eng(start.strftime('%B')),
                    'year_start': start.strftime('%Y'),
                    'day_finish': finish.strftime('%d'),
                    'month_finish': month_from_ru_to_eng(finish.strftime('%B')),
                    'year_finish': finish.strftime('%Y'),
                    'students_count': group.students.count(),
                    'code': group.code,
                }
                doc.render(context)
                newpath = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\"
                if not os.path.exists(newpath):
                    os.makedirs(newpath)
                doc.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_зачисление.docx")
                document = Document(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_зачисление.docx")
                tbl = document.tables[0]
                for index, student in enumerate(group.students.all()):
                    row_cells = tbl.add_row().cells
                    row_cells[0].text = str(index+1)
                    row_cells[1].text = student.surname+' '+student.name+' '+student.patronymic
                for row in tbl.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(12)
                document.save(MEDIA_ROOT + "\\Приказы\\ОУ\\Приказ_зачисление_" + group.code + ".docx")
                filename = MEDIA_ROOT + "\\Приказы\\ОУ\\Приказ_зачисление_" + group.code + ".docx"
                strin = 'Приказ_зачисление_' + group.code + '.docx'
                data = open(filename, "rb").read()
                response = HttpResponse(data,
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Length'] = os.path.getsize(filename)
                response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                return response
            elif 'order_expulsion' in request.POST:
                group = StudentGroups.objects.get(id=request.POST.get('order_expulsion'))
                finish = group.course.date_finish
                if finish.strftime('%w') == '6':
                    finish = finish + timedelta(days=2)
                if finish.strftime('%w') == '0':
                    finish = finish + timedelta(days=1)
                doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/order_expulsion_ou.docx")
                context = {
                    'prog_name': group.course.program.name,
                    'duration': group.course.program.duration,
                    'day_exp': finish.strftime('%d'),
                    'month_exp': month_from_ru_to_eng(finish.strftime('%B')),
                    'year_exp': finish.strftime('%Y'),
                }
                doc.render(context)
                newpath = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\"
                if not os.path.exists(newpath):
                    os.makedirs(newpath)
                doc.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx")
                document = Document(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx")
                tbl = document.tables[0]
                for index, student in enumerate(group.students.all()):
                    row_cells = tbl.add_row().cells
                    row_cells[0].text = str(index+1)+'.'
                    row_cells[1].text = student.surname+' '+student.name+' '+student.patronymic
                for row in tbl.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(12)
                document.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx")
                filename = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx"
                strin = 'Приказ_отчисление_' + group.code + '.docx'
                data = open(filename, "rb").read()
                response = HttpResponse(data,
                                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Length'] = os.path.getsize(filename)
                response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                return response
            elif 'delete_gr' in request.POST:
                StudentGroups.objects.get(id=request.POST.get('delete_gr')).delete()
                messages.success(request, 'Группа успешно удалена')
            elif 'close_study' in request.POST:
                group = StudentGroups.objects.get(id=request.POST.get('close_study'))
                group.survey_show = True
                group.status_id = StGroupStatuses.objects.get(name='Обучение завершено').id
                group.save()
                messages.success(request, 'Обучение в группе успешно завершено')
            elif 'journal_course' in request.POST:
                if not CourseLessons.objects.filter(group_id=request.POST.get('journal_course')).exists():
                    messages.error(request, 'Не найдены занятия в расписании учебной группы')
                    return HttpResponseRedirect('/centre/study/studentgroups')
                if not CourseLessons.objects.filter(group_id=request.POST.get('journal_course')). \
                        exclude(stschedule__in=StSchedule.objects.filter(control_form='')).exists():
                    messages.error(request, 'Не найдены занятия с формой контроля в расписании учебной группы')
                    return HttpResponseRedirect('/centre/study/studentgroups')
                group = StudentGroups.objects.get(id=request.POST.get('journal_course'))
                dep = group.course.program.department
                list = re.split(' |-', dep)
                short_dep = ''
                for el in list:
                    short_dep += el[:1].upper()
                date_start = group.course.date_start
                date_finish = group.course.date_finish
                if group.course.program.type_dpp == 'Повышение квалификации':
                    type_dpp = 'повышения квалификации'
                else:
                    type_dpp = 'профессиональной переподготовки'
                cats = ''
                for cat in group.course.program.categories.all():
                    cats += cat.name + '; '
                deps = GetDepsWithManagerFromAD()
                for key, value in deps.items():
                    if key == dep:
                        manager = value
                        fio = value.split(' ')
                        io_family_manager = fio[1][:1] + '.' + fio[2][:1] + '. ' + fio[0]
                        break
                total_lecture = total_practice = total_individual = total_trainee = total = 0
                lessons = CourseLessons.objects.filter(group_id=group.id)
                for les in lessons:
                    total_lecture += les.lecture_hours
                    total_practice += les.practice_hours
                    total_trainee += les.trainee_hours
                    total_individual += les.individual_hours
                    total += les.lecture_hours + les.practice_hours + les.individual_hours + les.trainee_hours
                path = STATIC_ROOT + '\\doc_templates\\xlsx\\journal_course.xlsx'
                writer = BookWriter(path)
                info = {
                    'type_dpp': type_dpp,
                    'name_dpp': group.course.program.name,
                    'duration': group.course.program.duration,
                    'place': group.course.place,
                    'code': group.code,
                    'dep': dep,
                    'cats': cats,
                    'manager': manager,
                    'day_start': date_start.strftime('%d'),
                    'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                    'year_start': date_start.strftime('%Y'),
                    'day_finish': date_finish.strftime('%d'),
                    'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                    'year_finish': date_finish.strftime('%Y'),
                    'io_family_manager': io_family_manager
                }
                info2 = {}
                info2['students'] = get_students(group.id)
                if group.study_form != 'Без использования ДОТ':
                    title = 'учебной аудиторной нагрузки,\nв т.ч. с использованием электронного обучения  и дистанционных образовательных технологий'
                else:
                    title = 'учебной аудиторной нагрузки'
                info3 = {
                    'total_lecture': total_lecture,
                    'total_practice': total_practice,
                    'total_trainee': total_trainee,
                    'total_individual': total_individual,
                    'total': total,
                    'short_dep': short_dep,
                    'manager': io_family_manager,
                    'title': title
                }
                info3['teachers'] = get_teachers(group.id)
                writer.render_sheet(info, 'тит', 0)
                writer.render_sheet(info2, 'список', 1)
                writer.render_sheet(info3, 'часы ауд', 2)
                count = 3
                exams = lessons.exclude(stschedule__in=StSchedule.objects.filter(control_form='')).order_by(
                    'lesson_time_start')
                newpath = MEDIA_ROOT + '\\Журналы\\' + group.code
                if not os.path.exists(newpath):
                    os.makedirs(newpath)
                for index, ex in enumerate(exams):
                    date_lesson = ex.lesson_time_start.date()
                    indexes = [m.start() for m in re.finditer(' ', ex.stschedule.name)]
                    info_pa = {
                        'theme': ex.stschedule.name[indexes[1]:],
                        'cats': cats,
                        'lector': ex.teacher.surname + ' ' + ex.teacher.name + ' ' + ex.teacher.patronymic,
                        'day_lesson': date_lesson.strftime('%d'),
                        'month_lesson': month_from_ru_to_eng(date_lesson.strftime('%B')),
                        'year_lesson': date_lesson.strftime('%Y'),
                        'control_form': ex.stschedule.control_form,
                        'manager': manager
                    }
                    info_pa['students'] = get_students(group.id)
                    if index < len(exams) - 1:
                        writer.render_sheet(info_pa, 'зачетная ведомость ПА ' + str(count - 2), 3)
                    else:
                        writer.render_sheet(info_pa, 'зачетная ведомость ИА ', 4)
                    count += 1
                writer.save(MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx")
                wb = openpyxl.load_workbook(MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx")
                ws = wb.worksheets[0]
                img = openpyxl.drawing.image.Image(STATIC_ROOT + '\\doc_templates\\xlsx\\logo.png')
                img.anchor = 'C1'
                ws.add_image(img)
                wb.save(MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx")
                filename = MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx"
                strin = "Журнал_" + group.code + ".xlsx"
                data = open(filename, "rb").read()
                response = HttpResponse(data,
                                        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Length'] = os.path.getsize(filename)
                response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                return response
            elif 'journal_event' in request.POST:
                group = StudentGroups.objects.get(id=request.POST.get('journal_event'))
                dep = group.event.department
                list = re.split(' |-', dep)
                short_dep = ''
                for el in list:
                    short_dep += el[:1].upper()
                date_start = group.event.date_start
                date_finish = group.event.date_finish
                morph = pymorphy2.MorphAnalyzer()
                type = morph.parse(group.event.type.name)[0]
                word = type.inflect({'loct'})
                cats = ''
                for cat in group.event.categories.all():
                    cats += cat.name + '; '
                deps = GetDepsWithManagerFromAD()
                for key, value in deps.items():
                    if key == dep:
                        manager = value
                        fio = value.split(' ')
                        io_family_manager = fio[1][:1] + '.' + fio[2][:1] + '. ' + fio[0]
                        break
                total_lecture = total_practice = total = 0
                lessons = EventsLessons.objects.filter(group_id=group.id)
                for les in lessons:
                    total_lecture += les.lecture_hours
                    total_practice += les.practice_hours
                    total += les.lecture_hours + les.practice_hours
                path = STATIC_ROOT + '\\doc_templates\\xlsx\\journal_event.xlsx'
                writer = BookWriter(path)
                info = {
                    'type_event': word.word,
                    'event_name': group.event.name,
                    'code': group.code,
                    'place': group.event.place,
                    'dep': dep,
                    'cats': cats,
                    'manager': manager,
                    'day_start': date_start.strftime('%d'),
                    'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                    'year_start': date_start.strftime('%Y'),
                    'day_finish': date_finish.strftime('%d'),
                    'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                    'year_finish': date_finish.strftime('%Y'),
                }
                writer.render_sheet(info, 'тит', 0)
                info2 = {}
                info2['students'] = get_students(group.id)
                writer.render_sheet(info2, 'Список', 1)
                info3 = {}
                total_l = total_p = total = 0
                for lesson in EventsLessons.objects.filter(group_id=group.id):
                    total_l += lesson.lecture_hours
                    total_p += lesson.practice_hours
                    total += lesson.lecture_hours + lesson.practice_hours
                info3['teachers'] = get_teachers(group.id)
                info3['total_lecture'] = total_l
                info3['total_practice'] = total_p
                info3['total'] = total
                writer.render_sheet(info3, 'часы ауд', 2)
                info4 = {}
                forms = EventsForms.objects.filter(group_id=group.id)
                total_mo = 0
                for mo in Mos.objects.all():
                    if forms.filter(mo_id=mo.id).exists():
                        info4[mo.tpl_name] = forms.filter(mo_id=mo.id).count()
                        total_mo += forms.filter(mo_id=mo.id).count()
                    else:
                        info4[mo.tpl_name] = '0'
                info4['total_mo'] = str(total_mo)
                total_oo = 0
                for type in OoTypes.objects.all():
                    if forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).exists():
                        info4[type.tpl_name] = forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).count()
                        total_oo += forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).count()
                    else:
                        info4[type.tpl_name] = '0'
                info4['total_oo'] = str(total_oo)
                total_pos = 0
                for type_pos in PositionCategories.objects.all():
                    if type_pos.tpl_name not in ['spec_1', 'spec_2']:
                        if forms.filter(position_cat_id=type_pos.id).exists():
                            info4[type_pos.tpl_name] = forms.filter(position_cat_id=type_pos.id).count()
                            total_pos += forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).count()
                        else:
                            info4[type_pos.tpl_name] = '0'
                spec = 0
                if forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_1').id).exists():
                    spec += forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_1').id).count()
                    total_pos += forms.filter(
                        position_cat_id=PositionCategories.objects.get(tpl_name='spec_1').id).count()
                if forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_2').id).exists():
                    spec += forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_2').id).count()
                    total_pos += forms.filter(
                        position_cat_id=PositionCategories.objects.get(tpl_name='spec_2').id).count()
                info4['spec'] = spec
                info4['total_pos'] = total_pos
                writer.render_sheet(info4, 'cтат', 3)
                newpath = MEDIA_ROOT + '\\Журналы\\' + group.code
                if not os.path.exists(newpath):
                    os.makedirs(newpath)
                writer.save(newpath + "\\Журнал.xlsx")
                filename = newpath + "\\Журнал.xlsx"
                strin = "Журнал_" + group.code + ".xlsx"
                data = open(filename, "rb").read()
                response = HttpResponse(data,
                                        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Length'] = os.path.getsize(filename)
                response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                return response
            else:
                str_teacher = request.POST.get('teacher')
                group = StudentGroups.objects.get(id=request.POST.get('curator'))
                group.curator_id = int(str_teacher[str_teacher.find(':')+1:str_teacher.find(')')])
                group.save()
                messages.success(request, 'Куратор успешно назначен')
        #except BaseException:
        #    messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/dep/study/studentgroups')


class StudentGroupCreate(CheckDepMixin, CreateView):
    "Создание новой учебной группы подразделения пользователя"
    login_url = '/'
    model = StudentGroups
    context_object_name = 'group'
    template_name = 'dep/study/studentgroups/create.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'id_type' in request.GET:
                typ = uri_to_iri(request.GET.get('id_type'))
                flag = True
                if typ in ['Повышение квалификации', 'Профессиональная переподготовка']:
                    courses = Courses.objects.filter(program__in=Programs.objects.filter(department=GetUserDepAD(self.request), type_dpp=typ)).order_by('-id')
                    if courses.count() != 0:
                        voc = {}
                        for course in courses:
                            voc[course.id] = course.program.name
                    else:
                        voc = None
                else:
                    flag = False
                    events = Events.objects.filter(department=GetUserDepAD(self.request)).\
                        filter(type_id=int(request.GET.get('id_type'))).order_by('-id')
                    if events.count() != 0:
                        voc = {}
                        for event in events:
                            voc[event.id] = event.name
                    else:
                        voc = None
                return JsonResponse({
                    'voc': voc,
                    'type': flag
                })
        return render(request, 'dep/study/studentgroups/create.html', {
            'courses': Courses.objects.all().order_by('-id')[:50],
            'events': Events.objects.all().order_by('-id')[:50],
        })

    def post(self, request, *args, **kwargs):
        try:
            if 'type' in request.POST:
                type = request.POST.get('type')
                if type in ['Повышение квалификации', 'Профессиональная переподготовка']:
                    if type == 'Повышение квалификации':
                        typ = 'Курс повышения квалификации (онлайн)'
                    else:
                        typ = 'Курс профессиональной переподготовки (онлайн)'
                    crs = request.POST.get('course')
                    dep_name = Programs.objects.get(id=Courses.objects.get(id=crs).program_id).department
                    if Courses.objects.filter(program__in=Programs.objects.filter(department=dep_name)).exists():
                        count = Courses.objects.filter(program__in=Programs.objects.filter(department=dep_name)).count()
                    else:
                        count = 0
                    list = re.split(' |-', dep_name)
                    dep = ''
                    for el in list:
                        dep += el[:1].upper()
                    date_start = Courses.objects.get(id=crs).date_start
                    date_service = date_start.strftime('%B')+' '+str(date_start.year)+' года'
                    start = date_start.strftime('%d.%m.%Y')
                    finish = Courses.objects.get(id=crs).date_finish.strftime('%d.%m.%Y')
                    name = Courses.objects.get(id=crs).program.name
                    duration = Courses.objects.get(id=crs).program.duration
                    code = CourseGroupCode(dep_name, str(count), date_start, 'course')
                    new_group = StudentGroups()
                    new_group.code = code
                    new_group.course_id = crs
                    new_group.curator = None
                    new_group.save()
                    messages.success(request, 'Группа успешно создана')
                else:
                    typ = EventTypes.objects.get(id=type).name+' (онлайн)'
                    vnt = request.POST.get('event')
                    dep_name = Events.objects.get(id=vnt).department
                    if Events.objects.filter(department=dep_name).exists():
                        count = Events.objects.filter(department=dep_name).count()
                    else:
                        count = 0
                    list = re.split(' |-', dep_name)
                    dep = ''
                    for el in list:
                        dep += el[:1].upper()
                    date_start = Events.objects.get(id=vnt).date_start
                    date_service = date_start.strftime('%B') + ' ' + str(date_start.year) + ' года'
                    start = date_start.strftime('%d.%m.%Y')
                    finish = Events.objects.get(id=vnt).date_finish.strftime('%d.%m.%Y')
                    name = Events.objects.get(id=vnt).name
                    duration = Events.objects.get(id=vnt).duration
                    code = CourseGroupCode(dep_name, str(count), date_start, 'event')
                    new_group = StudentGroups()
                    new_group.code = code
                    new_group.event_id = vnt
                    new_group.curator = None
                    new_group.students_number = int(request.POST.get('studentsnumber'))
                    new_group.save()
                    messages.success(request, 'Группа успешно создана')
                url = 'https://edu-dev.coko38.ru/student/detail/' + str(new_group.id)
                doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/create_stgroup/email.docx")
                context = {
                    'date_service': date_service,
                    'dep': dep,
                    'type': typ,
                    'name': name,
                    'url': url,
                    'duration': duration,
                    'start': start,
                    'finish': finish,
                }
                doc.render(context)
                doc.save(MEDIA_ROOT + "\\Технические файлы\\Email\\" + code + ".docx")
        except BaseException:
            messages.error(request, 'Произошла ошибка, повторите попытку позже...')
        return HttpResponseRedirect('/dep/study/studentgroups')


class StGrStudentsList(CheckDepMixin, ListView):
    "Работа с обучающимися учебной группы подразделения пользователя"
    login_url = '/'
    model = StudentGroups
    context_object_name = 'students'
    template_name = 'dep/study/studentgroups/students_list.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'list_sts' in request.GET:
                apps = Apps.objects.filter(group_id=request.GET.get('group')).\
                    filter(check_diploma_info=False)
                list_st = []
                for app in apps:
                    list_st.append(app.profile_id)
                return JsonResponse({'list_st': list_st})
            elif 'list_pay' in request.GET:
                apps = Apps.objects.filter(group_id=request.GET.get('group')).filter(check_pay=False).\
                    exclude(pay_doc_id=None)
                list_pay = []
                for app in apps:
                    list_pay.append(app.profile_id)
                return JsonResponse({'list_pay': list_pay})
            elif 'pay_student' in request.GET:
                try:
                    pay_doc = Apps.objects.filter(group_id=request.GET.get('group')).\
                        get(profile_id=request.GET.get('pay_student')).pay_doc_id
                    status = Apps.objects.filter(group_id=request.GET.get('group')).\
                        get(profile_id=request.GET.get('pay_student')).status.name
                    return JsonResponse({'pay_doc': pay_doc, 'status': status})
                except BaseException:
                    return JsonResponse({'pay_doc': None})
            elif 'pay_accept' in request.GET:
                app = Apps.objects.filter(group_id=request.GET.get('group')).\
                    get(profile_id=request.GET.get('pay_accept'))
                app.status_id = Statuses.objects.get(name='Оплачено').id
                app.save()
                group = StudentGroups.objects.get(id=request.GET.get('group'))
                if group.event is None:
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        typ = 'курсе повышения квалификации'
                    else:
                        typ = 'курс профессиональной переподготовки'
                    name = group.course.program.name
                else:
                    type_classic = group.event.type.name
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(type_classic)[0]
                    word = type.inflect({'loct'})
                    typ = word.word
                    name = group.event.name
                mail = Profiles.objects.get(id=request.GET.get('pay_accept')).user.email
                EmailAcceptPay.delay(mail, typ, name)
                messages.success(request, 'Оплата успешно подтверждена')
                return JsonResponse({})
            elif 'pay_denied' in request.GET:
                app = Apps.objects.get(id=request.GET.get('pay_denied'))
                id = app.pay_doc_id
                app.pay_doc_id = None
                app.save()
                Docs.objects.get(id=id).delete()
                app.status_id = Statuses.objects.get(name='Ждем оплату').id
                app.check_pay = False
                app.save()
                group = StudentGroups.objects.get(id=request.GET.get('group'))
                if group.event is None:
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        typ = 'курсе повышения квалификации'
                    else:
                        typ = 'курс профессиональной переподготовки'
                    name = group.course.program.name
                else:
                    type_classic = group.event.type.name
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(type_classic)[0]
                    word = type.inflect({'loct'})
                    typ = word.word
                    name = group.event.name
                mail = Profiles.objects.get(id=app.profile_id).user.email
                EmailDeniedPay.delay(mail, typ, name, request.GET.get('message'))
                messages.success(request, 'Оплата успешно отклонена')
                return JsonResponse({})
            elif 'pay_doc_id' in request.GET:
                app = Apps.objects.filter(group_id=request.GET.get('group')).get(profile_id=request.GET.get('pay_doc_id'))
                pd_id = app.pay_doc_id
                student = Profiles.objects.get(id=request.GET.get('pay_doc_id'))
                return JsonResponse({
                    'app': app.id,
                    'pd_id': pd_id,
                    'fio': student.surname+' '+student.name+' '+student.patronymic
                })
            else:
                form = CoursesForms.objects.filter(group_id=request.GET.get('group')).\
                    get(profile_id=request.GET.get('student'))
                if form.change_surname is not None:
                    change_sur = form.change_surname.id
                else:
                    change_sur = None
                student = Profiles.objects.get(id=request.GET.get('student'))
                if form.edu_cat is not None:
                    name_educat = form.edu_cat.name
                else:
                    name_educat = None
                return JsonResponse({
                    'fio': student.surname+' '+student.name+' '+student.patronymic,
                    'id_form': form.id,
                    'id_doc': form.edu_doc.id,
                    'surname': form.check_surname,
                    'id_changesur': change_sur,
                    'serial': form.edu_serial,
                    'number': form.edu_number,
                    'date': form.edu_date,
                    'edu_cat': name_educat
                })
        else:
            group = StudentGroups.objects.get(id=request.GET.get('id_group'))
            if group.event is None:
                if Programs.objects.get(id=Courses.objects.get(id=group.course_id).program_id).department != GetUserDepAD(self.request):
                    return HttpResponseRedirect('/access_denied/')
            else:
                if Events.objects.get(id=group.event_id).department != GetUserDepAD(self.request):
                    return HttpResponseRedirect('/access_denied/')
            qs = group.students.all().order_by('surname')
            if 'surname' in request.GET:
                if len(request.GET.get('surname')) > 0:
                    qs = qs.filter(surname__contains=request.GET.get('surname'))
            if 'name' in request.GET:
                if len(request.GET.get('name')) > 0:
                    qs = qs.filter(name__contains=request.GET.get('name'))
            if 'patronymic' in request.GET:
                if len(request.GET.get('patronymic')) > 0:
                    qs = qs.filter(patronymic__contains=request.GET.get('patronymic'))
            if 'email' in request.GET:
                if len(request.GET.get('email')) > 0:
                    qs = qs.filter(user__in=User.objects.filter(email__contains=request.GET.get('email')))
            NoCheckDip = True
            if Apps.objects.filter(group_id=request.GET.get('id_group')).filter(check_diploma_info=False).exists():
                NoCheckDip = False
            NoCheckPay = True
            if Apps.objects.filter(group_id=request.GET.get('id_group')).filter(check_pay=False).\
                    exclude(pay_doc_id=None).exists():
                NoCheckPay = False
            return render(request, 'dep/study/studentgroups/students_list.html', context={
                'students': qs[:25],
                'group': group,
                'NoCheckDip': NoCheckDip,
                'NoCheckPay': NoCheckPay
            })

    def post(self, request, *args, **kwargs):
        if 'surname' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.check_surname = request.POST.get('surname')
            form.save()
            return JsonResponse({'check': 'ok'})
        if 'serial' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.edu_serial = request.POST.get('serial')
            form.save()
            return JsonResponse({'check': 'ok'})
        if 'number' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.edu_number = request.POST.get('number')
            form.save()
            return JsonResponse({'check': 'ok'})
        if 'date' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.edu_date = request.POST.get('date')
            form.save()
            return JsonResponse({'check': 'ok'})
        if 'end' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            app = Apps.objects.filter(group_id=form.group_id).get(profile_id=form.profile_id)
            app.check_diploma_info = True
            app.save()
            return JsonResponse({'check': 'ok'})
        if 'pay_end' in request.POST:
            app = Apps.objects.get(id=request.POST.get('pay_end'))
            app.check_pay = True
            app.save()
            return JsonResponse({'check': 'ok'})
        if 'del_student' in request.POST:
            app = Apps.objects.get(id=request.POST.get('del_student'))
            group = StudentGroups.objects.get(id=app.group_id)
            if group.course is None:
                EventsForms.objects.filter(group_id=group.id).get(profile_id=app.profile_id).delete()
            else:
                CoursesForms.objects.filter(group_id=group.id).get(profile_id=app.profile_id).delete()
            StInGr = group.students.get(id=app.profile_id)
            group.students.remove(StInGr)
            app.delete()
            messages.success(request, 'Обучающийся успешно удален из группы')
            return HttpResponseRedirect('/dep/study/students?id_group='+request.POST.get('id_group'))
        if 'scan' in request.FILES:
            app = Apps.objects.filter(group_id=request.POST.get('group')).get(profile_id=request.POST.get('student'))
            new_doc = Docs()
            new_doc.profile_id = request.POST.get('student')
            new_doc.doc_type_id = DocsTypes.objects.get(name='Скан удостоверения').id
            new_doc.file = request.FILES.get('scan')
            new_doc.save()
            id_doc = new_doc.id
            app.certificate_id = id_doc
            app.status_id = Statuses.objects.get(name='Архив').id
            app.save()
            messages.success(request, 'Скан удостоверения успешно загружен')
            return JsonResponse({})


class ShedulesList(CheckDepMixin, ListView):
    "Расписания учебных групп подразделения пользователя"
    login_url = '/'
    model = StudentGroups
    context_object_name = 'groups'
    template_name = 'dep/study/schedule/list.html'

    def get_queryset(self):
        request = self.request
        qs = StudentGroups.objects. \
            filter(Q(event__in=Events.objects.filter(department=GetUserDepAD(self.request))) |
                   Q(course__in=Courses.objects.filter(
                       program__in=Programs.objects.filter(department=GetUserDepAD(self.request))))). \
            order_by('-id')
        if 'code' in self.request.GET:
            if len(request.GET.get('code')) > 0:
                qs = qs.filter(code__contains=request.GET.get('code'))
        if 'name' in self.request.GET:
            if len(request.GET.get('name')) > 0:
                courses = Courses.objects.filter(program__in=Programs.objects.filter(name__contains=request.GET.get('name')))
                events = Events.objects.filter(name__contains=request.GET.get('name'))
                qs = qs.filter(Q(event__in=events) | Q(course__in=courses))
        if 'study_date' in self.request.GET:
            try:
                post = self.request.GET.get('study_date')
                dt = date(int(post[:4]), int(post[-5:7]), int(post[8:]))
                courses = Courses.objects.filter(date_start__lte=dt, date_finish__gte=dt)
                events = Events.objects.filter(date_start__lte=dt, date_finish__gte=dt)
                qs = qs.filter(Q(event__in=events) | Q(course__in=courses))
            except BaseException:
                pass
        return qs

    def post(self, request, *args, **kwargs):
        group = StudentGroups.objects.get(id=request.POST.get('group_id'))
        if group.course is None:
            dep = group.event.department
            date_start = group.event.date_start
            date_finish = group.event.date_finish
            event_cats = ''
            for cat in group.event.categories.all():
                event_cats += cat.name+'; '
            lessons = EventsLessons.objects.filter(group_id=group.id)
            total_lecture = total_practice = total_hours = 0
            for el in lessons:
                total_lecture += el.lecture_hours
                total_practice += el.practice_hours
            total_hours = total_lecture + total_practice
            short_dep = ''
            list = re.split(' |-', dep)
            for el in list:
                short_dep += el[:1].upper()
            deps = GetDepsWithManagerFromAD()
            for key,value in deps.items():
                if key == dep:
                    fio = value.split(' ')
                    io_family_manager = fio[1][:1]+'.'+fio[2][:1]+'. '+fio[0]
            path = STATIC_ROOT + '\\doc_templates\\xlsx\\schedule_event.xlsx'
            writer = BookWriter(path)
            info = {
                'dep': dep,
                'code': group.code,
                'event_name': group.event.name,
                'events_cats': event_cats,
                'day_start': date_start.strftime('%d'),
                'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                'year_start': date_start.strftime('%Y'),
                'day_finish': date_finish.strftime('%d'),
                'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                'year_finish': date_finish.strftime('%Y'),
                'total_lecture': total_lecture,
                'total_practice': total_practice,
                'total_hours': total_hours,
                'short_dep': short_dep,
                'io_family_manager': io_family_manager
            }
        else:
            dep = group.course.program.department
            date_start = group.course.date_start
            date_finish = group.course.date_finish
            course_cats = ''
            for cat in group.course.program.categories.all():
                course_cats += cat.name + '; '
            lessons = CourseLessons.objects.filter(group_id=group.id)
            total_lecture = total_practice = total_individual = total_hours = 0
            for el in lessons:
                total_lecture += el.lecture_hours
                total_practice += el.practice_hours
                total_individual += el.individual_hours
            total_hours = total_lecture + total_practice + total_individual
            short_dep = ''
            list = re.split(' |-', dep)
            for el in list:
                short_dep += el[:1].upper()
            deps = GetDepsWithManagerFromAD()
            for key, value in deps.items():
                if key == dep:
                    fio = value.split(' ')
                    io_family_manager = fio[1][:1] + '.' + fio[2][:1] + '. ' + fio[0]
            path = STATIC_ROOT + '\\doc_templates\\xlsx\\schedule_course.xlsx'
            writer = BookWriter(path)
            info = {
                'dep': dep,
                'code': group.code,
                'course_name': group.course.program.name,
                'course_cats': course_cats,
                'day_start': date_start.strftime('%d'),
                'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                'year_start': date_start.strftime('%Y'),
                'day_finish': date_finish.strftime('%d'),
                'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                'year_finish': date_finish.strftime('%Y'),
                'total_lecture': total_lecture,
                'total_practice': total_practice,
                'total_individual': total_individual,
                'total_hours': total_hours,
                'short_dep': short_dep,
                'io_family_manager': io_family_manager
            }
        info['items'] = get_lessons(group.id)
        payloads = [info]
        writer.render_book(payloads=payloads)
        newpath = MEDIA_ROOT + '\\Расписания\\'+group.code
        if not os.path.exists(newpath):
            os.makedirs(newpath)
        writer.save(MEDIA_ROOT + '\\Расписания\\'+group.code + "\\Расписание.xlsx")
        filename = MEDIA_ROOT + '\\Расписания\\'+group.code + "\\Расписание.xlsx"
        strin = "Расписание_"+group.code+".xlsx"
        data = open(filename, "rb").read()
        response = HttpResponse(data,
                                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Length'] = os.path.getsize(filename)
        response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
        return response

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super(ShedulesList, self).get_context_data(**kwargs)
        context['dep_name'] = GetUserDepAD(self.request)
        return context


class CourseLessonsList(CheckDepMixin, ListView):
    "Расписание занятий курса подразделения пользователя"
    login_url = '/'
    model = CourseLessons
    context_object_name = 'lessons'
    template_name = 'dep/study/schedule/lessons_course.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'timestart' in request.GET:
                check_timestart = True
                check_break = True
                check_lunch = True
                if CourseLessons.objects.filter(group_id=self.kwargs.get('group')).exists():
                    str_start = request.GET.get('timestart')
                    timestart = datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                                  int(str_start[11:13]) - 8, int(str_start[14:16]))
                    tf = timestart + timedelta(hours=8, minutes=45)
                    timezone.make_aware(timestart, timezone.get_default_timezone())
                    lessons = CourseLessons.objects.exclude(id=request.GET.get('lesson')).filter(
                        group_id=self.kwargs.get('group')).order_by('lesson_time_start')
                    for lesson in lessons:
                        t_s = lesson.lesson_time_start.replace(tzinfo=None)
                        t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                        if t_s <= timestart < t_f:
                            check_timestart = False
                            break
                    for previous, current in zip(lessons, lessons[1:]):
                        t_s = current.lesson_time_start.replace(tzinfo=None)
                        t_f = previous.lesson_time_finish.replace(tzinfo=None)
                        if t_f <= timestart <= t_s:
                            main_tf = t_f
                            check_twohours = False
                            for les in lessons:
                                if les.lesson_time_finish == previous.lesson_time_start:
                                    check_twohours = True
                            if check_twohours is True:
                                break_between = timestart - previous.lesson_time_finish.replace(tzinfo=None)
                                if break_between.seconds // 60 < 10:
                                    check_break = False
                                    break
                    if check_timestart is True and check_break is True:
                        count_hours = 0
                        for lesson in lessons:
                            t_s = lesson.lesson_time_start.replace(tzinfo=None)
                            t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                            if t_f.date() == timestart.date() and t_s < timestart:
                                count_hours += 1
                            if count_hours >= 4:
                                was_lunch = False
                                for previous, current in zip(lessons, lessons[1:]):
                                    t_s = current.lesson_time_start.replace(tzinfo=None)
                                    t_f = previous.lesson_time_finish.replace(tzinfo=None)
                                    if t_s.date() == t_f.date() and t_s < timestart:
                                        br = t_s - t_f
                                        if br.seconds // 60 >= 30:
                                            was_lunch = True
                                            break
                                if was_lunch is False:
                                    break_lunch = timestart - main_tf
                                    if break_lunch.seconds // 60 < 30:
                                        check_lunch = False
                    else:
                        pass
                return JsonResponse({
                    'check_timestart': check_timestart,
                    'check_break': check_break,
                    'check_lunch': check_lunch,
                    'tf': tf.strftime('%H:%M')
                })
            elif 'timefinish' in request.GET:
                check_timefinish = True
                if CourseLessons.objects.exclude(id=request.GET.get('lesson')).filter(
                        group_id=self.kwargs.get('group')).exists():
                    str_finish = request.GET.get('timefinish')
                    timefinish = datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                                   int(str_finish[11:13]), int(str_finish[14:16]))
                    if CourseLessons.objects.exclude(id=request.GET.get('lesson')).filter(
                            group_id=self.kwargs.get('group')) \
                            .filter(lesson_time_start__lt=timefinish). \
                            filter(lesson_time_finish__gt=timefinish).exists():
                        check_timefinish = False
                    else:
                        pass
                else:
                    pass
                return JsonResponse({'check_timefinish': check_timefinish})
            elif 'get_themes' in request.GET:
                if StSchedule.objects.filter(
                        program_id=StudentGroups.objects.get(id=self.kwargs.get('group')).course.program.id).exists():
                    themes = StSchedule.objects.filter(
                        program_id=StudentGroups.objects.get(id=self.kwargs.get('group')).course.program.id)
                    dict_themes = {}
                    count = 1
                    for theme in themes:
                        list_th = []
                        list_th.append(theme.id)
                        list_th.append(theme.name)
                        dict_themes[count] = list_th
                        count += 1
                    return JsonResponse({
                        'themes': dict_themes
                    })
                else:
                    return JsonResponse({})
            elif 'theme_id' in request.GET:
                theme = StSchedule.objects.get(id=request.GET.get('theme_id'))
                l_hours = theme.lecture_hours
                p_hours = theme.practice_hours
                t_hours = theme.trainee_hours
                i_hours = theme.individual_hours
                if CourseLessons.objects.filter(group_id=self.kwargs.get('group')). \
                        filter(stschedule_id=request.GET.get('theme_id')).exists():
                    lessons = CourseLessons.objects.filter(group_id=self.kwargs.get('group')). \
                        filter(stschedule_id=request.GET.get('theme_id'))
                    for lesson in lessons:
                        l_hours -= lesson.lecture_hours
                        p_hours -= lesson.practice_hours
                        t_hours -= lesson.trainee_hours
                        i_hours -= lesson.individual_hours
                teachers = Profiles.objects.filter(teacher=True)
                list_t = []
                for teacher in teachers:
                    list_t.append(teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic + ' (ID:' + str(
                        teacher.id) + ')')
                return JsonResponse({
                    'lecture': l_hours,
                    'practice': p_hours,
                    'individual': i_hours,
                    'trainee': t_hours,
                    'list_t': list_t
                })
            elif 'id_teach' in request.GET:
                str_start = request.GET.get('time_start')
                str_finish = request.GET.get('time_finish')
                teacher = Profiles.objects.get(id=int(request.GET.get('id_teach')))
                check_free = True
                timestart = datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                               int(str_finish[11:13]), int(str_finish[14:16]))
                if CourseLessons.objects.filter(teacher_id=request.GET.get('id_teach')). \
                        filter(lesson_time_start__lte=timestart). \
                        filter(lesson_time_finish__gt=timestart).exists():
                    check_free = False
                if CourseLessons.objects.filter(teacher_id=request.GET.get('id_teach')). \
                        filter(lesson_time_start__lte=timefinish). \
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('id_teach')). \
                        filter(lesson_time_start__lte=timestart). \
                        filter(lesson_time_finish__gt=timestart).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('id_teach')). \
                        filter(lesson_time_start__lte=timefinish). \
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check_free = False
                return JsonResponse({
                    'fio': teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic,
                    'email': User.objects.get(id=teacher.user_id).email,
                    'phone': teacher.phone,
                    'freetime': check_free,
                })
            elif 'form_study' in request.GET:
                group = StudentGroups.objects.get(id=request.GET.get('group'))
                if request.GET.get('form_study') == 'without_dot':
                    group.study_form = 'Без использования ДОТ'
                elif request.GET.get('form_study') == 'only_dot':
                    group.study_form = 'Исключительно ДОТ'
                else:
                    group.study_form = 'С использованием ДОТ'
                group.save()
                return JsonResponse({})
            elif 'change_lesson' in request.GET:
                less = CourseLessons.objects.select_related('stschedule').get(id=request.GET.get('change_lesson'))
                group = StudentGroups.objects.select_related('course').get(id=less.group_id)
                date_s = group.course.date_start
                date_f = group.course.date_finish
                teachs = Profiles.objects.filter(teacher=True).order_by('surname')
                if less.lecture_hours == less.practice_hours == less.trainee_hours == 0:
                    if less.control == '':
                        type = 'Самостоятельная работа'
                    else:
                        type = 'Контрольное занятие'
                elif less.individual_hours == less.practice_hours == less.trainee_hours == 0:
                    type = 'Лекция'
                elif less.individual_hours == less.lecture_hours == less.trainee_hours == 0:
                    if less.control == '':
                        type = 'Практика'
                    else:
                        type = 'Контрольное занятие'
                elif less.individual_hours == less.lecture_hours == less.practice_hours == 0:
                    type = 'Стажировка'
                else:
                    type = 'Смешанное занятие'
                dict_teachs = {}
                current_teacher = ''
                for teach in teachs:
                    if less.teacher_id == teach.id:
                        current_teacher = teach.surname + ' ' + teach.name + ' ' + teach.patronymic + ' (ID:' + str(
                            teach.id) + ')'
                    dict_teachs[teach.id] = teach.surname + ' ' + teach.name + ' ' + teach.patronymic
                start = less.lesson_time_start.replace(tzinfo=None)
                finish = less.lesson_time_finish.replace(tzinfo=None)
                start = start + timedelta(hours=8)
                finish = finish + timedelta(hours=8)
                return JsonResponse({
                    'date_s': date_s.strftime('%d.%m.%Y'),
                    'date_f': date_f.strftime('%d.%m.%Y'),
                    'theme_name': less.stschedule.name,
                    'date_start': less.lesson_time_start.strftime('%Y-%m-%d'),
                    'time_start': start.strftime('%H:%M'),
                    'time_finish': finish.strftime('%H:%M'),
                    'distance': less.distance,
                    'control': less.control,
                    'teachers': dict_teachs,
                    'current_teacher': current_teacher,
                    'type': type
                })
            elif 'check_date' in request.GET:
                course = StudentGroups.objects.get(id=self.kwargs.get('group')).course
                d = date(int(request.GET.get('check_date')[:4]), int(request.GET.get('check_date')[5:7]),
                                  int(request.GET.get('check_date')[8:10]))
                if course.date_start <= d <= course.date_finish:
                    check = True
                else:
                    check = False
                return JsonResponse({
                    'check': check
                })
            elif 'generate' in request.GET:
                group = StudentGroups.objects.select_related('course').get(id=self.kwargs.get('group'))
                str_form = request.GET.get('generate')
                max_day_hours = int(request.GET.get('count'))
                if not StSchedule.objects.filter(program_id=group.course.program.id).exists():
                    messages.error(request, 'Не найдены темы и разделы в КУГ')
                    return HttpResponseRedirect(
                        '/centre/study/schedule/course_lessons_' + str(self.kwargs.get('group')) + '#close')
                start_datetime = datetime(int(str_form[:4]),
                                                   int(str_form[5:7]),
                                                   int(str_form[8:10]),
                                                   int(str_form[11:13]),
                                                   int(str_form[14:16]))
                ls = CourseLessons.objects.filter(group_id=self.kwargs.get('group')). \
                    filter(lesson_time_start__year=start_datetime.year,
                           lesson_time_start__month=start_datetime.month,
                           lesson_time_start__day=start_datetime.day)
                for l in ls:
                    l.delete()
                day_hours = 0
                while day_hours < max_day_hours:
                    new_les = CourseLessons()
                    new_les.group_id = group.id
                    if day_hours % 2 == 0 and day_hours != 0:
                        if day_hours % 4 == 0:
                            new_les.lesson_time_start = start_datetime + timedelta(minutes=30)
                            new_les.lesson_time_finish = start_datetime + timedelta(minutes=75)
                            start_datetime = start_datetime + timedelta(minutes=75)
                        else:
                            new_les.lesson_time_start = start_datetime + timedelta(minutes=10)
                            new_les.lesson_time_finish = start_datetime + timedelta(minutes=55)
                            start_datetime = start_datetime + timedelta(minutes=55)
                    else:
                        new_les.lesson_time_start = start_datetime
                        new_les.lesson_time_finish = start_datetime + timedelta(minutes=45)
                        start_datetime = start_datetime + timedelta(minutes=45)
                    new_les.lecture_hours = 0
                    new_les.practice_hours = 0
                    new_les.trainee_hours = 0
                    new_les.individual_hours = 0
                    new_les.teacher = None
                    new_les.distance = False
                    new_les.control = ''
                    new_les.save()
                    day_hours += 1
                ls = CourseLessons.objects.filter(group_id=self.kwargs.get('group')). \
                    filter(lesson_time_start__year=start_datetime.year,
                           lesson_time_start__month=start_datetime.month,
                           lesson_time_start__day=start_datetime.day)
                dict_l = {}
                for l in ls:
                    list_l = []
                    list_l.append(l.id)
                    lts = l.lesson_time_start.replace(tzinfo=None) + timedelta(hours=8)
                    ltf = l.lesson_time_finish.replace(tzinfo=None) + timedelta(hours=8)
                    list_l.append(lts.strftime('%H:%M'))
                    list_l.append(ltf.strftime('%H:%M'))
                    dict_l[l.id] = list_l
                themes = StSchedule.objects.filter(program_id=Programs.objects.get(id=group.course.program.id).id)
                dict_th = {}
                for theme in themes:
                    if theme.parent_id is not None or theme.control_form != '':
                        list_th = []
                        list_th.append(theme.id)
                        if 'Раздел ' in theme.name:
                            list_th.append(theme.name[7:])
                        else:
                            list_th.append(theme.name[5:])
                        free_hours = theme.total_hours
                        if CourseLessons.objects.filter(group_id=self.kwargs.get('group')).filter(
                                stschedule_id=theme.id).exists():
                            for l in CourseLessons.objects.filter(group_id=self.kwargs.get('group')).filter(
                                    stschedule_id=theme.id):
                                if l.lecture_hours != 0 and l.lecture_hours is not None:
                                    free_hours -= l.lecture_hours
                                if l.practice_hours != 0 and l.practice_hours is not None:
                                    free_hours -= l.practice_hours
                                if l.trainee_hours != 0 and l.trainee_hours is not None:
                                    free_hours -= l.trainee_hours
                                if l.individual_hours != 0 and l.individual_hours is not None:
                                    free_hours -= l.individual_hours
                        list_th.append(free_hours)
                        dict_th[theme.id] = list_th
                teachers = Profiles.objects.filter(teacher=True)
                dict_teach = {}
                for teacher in teachers:
                    dict_teach[
                        teacher.id] = teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic + ' (ID:' + str(
                        teacher.id) + ')'
                return JsonResponse({
                    'lessons': dict_l,
                    'themes': dict_th,
                    'teachers': dict_teach
                })
            elif 'getthemechoose' in request.GET:
                group = StudentGroups.objects.get(id=self.kwargs.get('group'))
                lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                theme = StSchedule.objects.filter(name__contains=request.GET.get('getthemechoose')).get(
                    program_id=group.course.program.id)
                lect_h = theme.lecture_hours
                prac_h = theme.practice_hours
                trai_h = theme.trainee_hours
                indi_h = theme.individual_hours
                for less in CourseLessons.objects.filter(group_id=self.kwargs.get('group')):
                    if less.stschedule_id == theme.id:
                        if less.lecture_hours != 0 and less.lecture_hours is not None:
                            lect_h -= less.lecture_hours
                        if less.practice_hours != 0 and less.practice_hours is not None:
                            prac_h -= less.practice_hours
                        if less.trainee_hours != 0 and less.trainee_hours is not None:
                            trai_h -= less.trainee_hours
                        if less.individual_hours != 0 and less.individual_hours is not None:
                            indi_h -= less.individual_hours
                lesson.stschedule_id = theme.id
                lesson.lecture_hours = 0
                lesson.practice_hours = 0
                lesson.trainee_hours = 0
                lesson.individual_hours = 0
                lesson.control = theme.control_form
                lesson.save()
                return JsonResponse({
                    'lect_h': lect_h,
                    'prac_h': prac_h,
                    'trai_h': trai_h,
                    'indi_h': indi_h,
                    'control': theme.control_form
                })
            elif 'lestypechoose' in request.GET:
                group = StudentGroups.objects.get(id=self.kwargs.get('group'))
                lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                type = request.GET.get('lestypechoose')
                if type == 'lecture':
                    lesson.lecture_hours = 1
                    lesson.practice_hours = 0
                    lesson.trainee_hours = 0
                    lesson.individual_hours = 0
                elif type == 'practice':
                    lesson.lecture_hours = 0
                    lesson.practice_hours = 1
                    lesson.trainee_hours = 0
                    lesson.individual_hours = 0
                elif type == 'trainee':
                    lesson.lecture_hours = 0
                    lesson.practice_hours = 0
                    lesson.trainee_hours = 1
                    lesson.individual_hours = 0
                else:
                    lesson.lecture_hours = 0
                    lesson.practice_hours = 0
                    lesson.trainee_hours = 0
                    lesson.individual_hours = 1
                lesson.save()
                return JsonResponse({})
            elif 'lesformatchoose' in request.GET:
                lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                if request.GET.get('lesformatchoose') == 'dot':
                    lesson.distance = True
                else:
                    lesson.distance = False
                lesson.save()
                return JsonResponse({})
            elif 'teacher_gen' in request.GET:
                check = True
                str_start = request.GET.get('time_start')
                str_finish = request.GET.get('time_finish')
                timestart = datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                               int(str_finish[11:13]), int(str_finish[14:16]))
                if CourseLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')). \
                        filter(lesson_time_start__lte=timestart). \
                        filter(lesson_time_finish__gt=timestart).exists():
                    check = False
                if CourseLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')). \
                        filter(lesson_time_start__lte=timefinish). \
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')). \
                        filter(lesson_time_start__lte=timestart). \
                        filter(lesson_time_finish__gt=timestart).exists():
                    check = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')). \
                        filter(lesson_time_start__lte=timefinish). \
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check = False
                if check is True:
                    lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                    lesson.teacher_id = request.GET.get('teacher_gen')
                    lesson.save()
                return JsonResponse({
                    'check': check
                })
            else:
                return HttpResponseRedirect('/access_denied/')
        else:
            try:
                prog = StudentGroups.objects.get(id=self.kwargs.get('group')).course.program_id
                if Programs.objects.get(id=prog).department == GetUserDepAD(self.request):
                    if 'save_less' in request.GET:
                        timestart = datetime(int(request.GET.get('date_less')[:4]),
                                                      int(request.GET.get('date_less')[5:7]),
                                                      int(request.GET.get('date_less')[8:10]),
                                                      int(request.GET.get('time_start')[:2]),
                                                      int(request.GET.get('time_start')[3:]))
                        timefinish = datetime(int(request.GET.get('date_less')[:4]),
                                                       int(request.GET.get('date_less')[5:7]),
                                                       int(request.GET.get('date_less')[8:10]),
                                                       int(request.GET.get('time_finish')[:2]),
                                                       int(request.GET.get('time_finish')[3:]))
                        lesson = CourseLessons.objects.get(id=request.GET.get('save_less'))
                        LessonDay = date(int(request.GET.get('date_less')[:4]),
                                                  int(request.GET.get('date_less')[5:7]),
                                                  int(request.GET.get('date_less')[8:10]), )
                        if CourseLessons.objects.filter(group_id=lesson.group_id).exclude(
                                id=request.GET.get('save_less')). \
                                filter(lesson_time_start__lte=timestart).\
                                filter(lesson_time_finish__gt=timestart).exists():
                            check = True
                            CountPrevLessons = None
                            TargetLesson = None
                            NextLessons = None
                            if CourseLessons.objects.exclude(id=lesson.id). \
                                    filter(group_id=lesson.group_id). \
                                    filter(lesson_time_finish__day=LessonDay.day). \
                                    filter(lesson_time_finish__month=LessonDay.month). \
                                    filter(lesson_time_finish__year=LessonDay.year). \
                                    filter(lesson_time_start__lt=timestart).exists():
                                CountPrevLessons = CourseLessons.objects.exclude(id=lesson.id). \
                                    filter(group_id=lesson.group_id). \
                                    filter(lesson_time_finish__day=LessonDay.day). \
                                    filter(lesson_time_finish__month=LessonDay.month). \
                                    filter(lesson_time_finish__year=LessonDay.year). \
                                    filter(lesson_time_start__lt=timestart).count()
                                TargetLesson = CourseLessons.objects.filter(group_id=lesson.group_id). \
                                    filter(lesson_time_finish__day=LessonDay.day). \
                                    filter(lesson_time_finish__month=LessonDay.month). \
                                    filter(lesson_time_finish__year=LessonDay.year). \
                                    filter(lesson_time_start__lt=timestart).latest('lesson_time_start')
                            if CourseLessons.objects.filter(group_id=lesson.group_id). \
                                    filter(lesson_time_finish__day=LessonDay.day). \
                                    filter(lesson_time_finish__month=LessonDay.month). \
                                    filter(lesson_time_finish__year=LessonDay.year). \
                                    filter(lesson_time_start__gte=timestart).exists():
                                NextLessons = CourseLessons.objects.filter(group_id=lesson.group_id). \
                                    filter(lesson_time_finish__day=LessonDay.day). \
                                    filter(lesson_time_finish__month=LessonDay.month). \
                                    filter(lesson_time_finish__year=LessonDay.year). \
                                    filter(lesson_time_start__gte=timestart).order_by('lesson_time_start')
                            if NextLessons is not None:
                                if TargetLesson is not None:
                                    new_finish = TargetLesson.lesson_time_finish.replace(tzinfo=None)
                                    if check is True:
                                        if CountPrevLessons % 2 == 0:
                                            if CountPrevLessons % 4 == 0:
                                                lesson.lesson_time_start = new_finish + timedelta(hours=8, minutes=30)
                                                lesson.lesson_time_finish = new_finish + timedelta(hours=8, minutes=75)
                                            else:
                                                lesson.lesson_time_start = new_finish + timedelta(hours=8, minutes=10)
                                                lesson.lesson_time_finish = new_finish + timedelta(hours=8, minutes=55)
                                        else:
                                            lesson.lesson_time_start = new_finish + timedelta(hours=8)
                                            lesson.lesson_time_finish = new_finish + timedelta(hours=8, minutes=45)
                                    for curr, next in zip(NextLessons, NextLessons[1:]):
                                        curr.lesson_time_start = next.lesson_time_start
                                        curr.lesson_time_finish = next.lesson_time_finish
                                        curr.save()
                                    LastNext = NextLessons.latest('lesson_time_start')
                                    if (CountPrevLessons + NextLessons.count()) % 2 == 0:
                                        if (CountPrevLessons + NextLessons.count()) % 4 == 0:
                                            LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(
                                                minutes=75)
                                        else:
                                            LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(
                                                minutes=55)
                                    else:
                                        LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                                    LastNext.save()
                                    LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                                    LastNext.save()
                                else:
                                    lesson.lesson_time_start = timestart
                                    lesson.lesson_time_finish = timefinish
                                    for curr, next in zip(NextLessons, NextLessons[1:]):
                                        curr.lesson_time_start = next.lesson_time_start
                                        curr.lesson_time_finish = next.lesson_time_finish
                                        curr.save()
                                    FirstNext = NextLessons.latest('-lesson_time_start')
                                    FirstNext.lesson_time_start = timefinish
                                    FirstNext.save()
                                    FirstNext.lesson_time_finish = FirstNext.lesson_time_start + timedelta(minutes=45)
                                    FirstNext.save()
                                    LastNext = NextLessons.latest('lesson_time_start')
                                    if NextLessons.count() % 2 == 0:
                                        if NextLessons.count() % 4 == 0:
                                            LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(
                                                minutes=75)
                                        else:
                                            LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(
                                                minutes=55)
                                    else:
                                        LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                                    LastNext.save()
                                    LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                                    LastNext.save()
                            else:
                                lesson.lesson_time_start = timestart
                                lesson.lesson_time_finish = timefinish
                        else:
                            lesson.lesson_time_start = timestart
                            lesson.lesson_time_finish = timefinish
                        str_teacher = request.GET.get('teacher')
                        lesson.teacher_id = int(str_teacher[str_teacher.find(':') + 1:str_teacher.find(')')])
                        lesson.control = request.GET.get('control')
                        lesson.distance = request.GET.get('dist')
                        lesson.save()
                        messages.success(request, 'Информация о занятии успешно изменена')
                        return HttpResponseRedirect(
                            '/dep/study/schedule/course_lessons_' + str(self.kwargs.get('group')) + '#close')
                    return render(request, 'dep/study/schedule/lessons_course.html', context={
                        'lessons': CourseLessons.objects.filter(group_id=self.kwargs.get('group')).order_by('lesson_time_start'),
                        'group': StudentGroups.objects.get(id=self.kwargs.get('group')),
                    })
                else:
                    return HttpResponseRedirect('/access_denied/')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
                return HttpResponseRedirect('/dep/study/schedule')

    def post(self, request, *args, **kwargs):
        if 'delete_id' in request.POST:
            try:
                CourseLessons.objects.get(id=request.POST.get('delete_id')).delete()
                messages.success(request, 'Занятие успешно удалено')
            except BaseException:
                messages.error(request, 'Произошшла ошибка, повторите попытку позже')
        else:
            try:
                str_start = request.POST.get('lesson_time_start')
                str_finish = request.POST.get('lesson_time_finish')
                timestart = datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                               int(str_finish[11:13]), int(str_finish[14:16]))
                str_teacher = request.POST.get('teacher')
                new = CourseLessons()
                new.group_id = int(self.kwargs.get('group'))
                new.stschedule_id = int(request.POST.get('stschedule'))
                new.lecture_hours = int(request.POST.get('lecture_hours'))
                new.practice_hours = int(request.POST.get('practice_hours'))
                new.trainee_hours = int(request.POST.get('trainee_hours'))
                new.individual_hours = int(request.POST.get('individual_hours'))
                new.lesson_time_start = timestart
                new.lesson_time_finish = timefinish
                new.teacher_id = int(str_teacher[str_teacher.find(':')+1:str_teacher.find(')')])
                new.save()
                messages.success(request, 'Занятие успешно добавлено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/dep/study/schedule/course_lessons_'+str(self.kwargs.get('group')))

    def get_context_data(self, **kwargs):
        context = super(CourseLessonsList, self).get_context_data(**kwargs)
        context.update({
            'course': Courses.objects.get(id=StudentGroups.objects.get(id=self.kwargs.get('group')).course_id)
        })
        return context


class EventLessonsList(CheckDepMixin, ListView):
    "Расписание занятий мероприятия подразделения пользователя"
    login_url = '/'
    model = EventsLessons
    context_object_name = 'lessons'
    template_name = 'dep/study/schedule/lessons_event.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'timestart' in request.GET:
                check_timestart = True
                check_break = True
                check_lunch = True
                str_start = request.GET.get('timestart')
                timestart = datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]) - 8, int(str_start[14:16]))
                if EventsLessons.objects.filter(group_id=self.kwargs.get('group')).exists():
                    timezone.make_aware(timestart, timezone.get_default_timezone())
                    lessons = EventsLessons.objects.exclude(id=request.GET.get('lesson')).filter(
                        group_id=self.kwargs.get('group')).order_by('lesson_time_start')
                    for lesson in lessons:
                        t_s = lesson.lesson_time_start.replace(tzinfo=None)
                        t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                        if t_s <= timestart < t_f:
                            check_timestart = False
                            break
                    for previous, current in zip(lessons, lessons[1:]):
                        t_s = current.lesson_time_start.replace(tzinfo=None)
                        t_f = previous.lesson_time_finish.replace(tzinfo=None)
                        if t_f <= timestart <= t_s:
                            main_tf = t_f
                            check_twohours = False
                            for les in lessons:
                                if les.lesson_time_finish == previous.lesson_time_start:
                                    check_twohours = True
                            if check_twohours is True:
                                break_between = timestart - previous.lesson_time_finish.replace(tzinfo=None)
                                if break_between.seconds // 60 < 10:
                                    check_break = False
                                    break
                    if check_timestart is True and check_break is True:
                        count_hours = 0
                        for lesson in lessons:
                            t_s = lesson.lesson_time_start.replace(tzinfo=None)
                            t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                            if t_f.date() == timestart.date() and t_s < timestart:
                                count_hours += 1
                            if count_hours >= 4:
                                was_lunch = False
                                for previous, current in zip(lessons, lessons[1:]):
                                    t_s = current.lesson_time_start.replace(tzinfo=None)
                                    t_f = previous.lesson_time_finish.replace(tzinfo=None)
                                    if t_s.date() == t_f.date() and t_s < timestart:
                                        br = t_s - t_f
                                        if br.seconds // 60 >= 30:
                                            was_lunch = True
                                            break
                                if was_lunch is False:
                                    break_lunch = timestart - main_tf
                                    if break_lunch.seconds // 60 < 30:
                                        check_lunch = False
                    else:
                        pass
                return JsonResponse({
                    'check_timestart': check_timestart,
                    'check_break': check_break,
                    'check_lunch': check_lunch,
                })
            elif 'timefinish' in request.GET:
                check_timefinish = True
                if EventsLessons.objects.filter(group_id=self.kwargs.get('group')).exists():
                    str_finish = request.GET.get('timefinish')
                    timefinish = datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                                   int(str_finish[11:13]), int(str_finish[14:16]))
                    if EventsLessons.objects.filter(group_id=self.kwargs.get('group')) \
                            .filter(lesson_time_start__lte=timefinish). \
                            filter(lesson_time_finish__gte=timefinish).exists():
                        check_timefinish = False
                teachers = Profiles.objects.filter(teacher=True)
                list_t = []
                for teacher in teachers:
                    list_t.append(teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic + ' (ID:' + str(
                        teacher.id) + ')')
                return JsonResponse({
                    'check_timefinish': check_timefinish,
                    'list_t': list_t
                })
            elif 'id_teach' in request.GET:
                str_start = request.GET.get('time_start')
                str_finish = request.GET.get('time_finish')
                teacher = Profiles.objects.get(id=int(request.GET.get('id_teach')))
                check_free = True
                timestart = datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                               int(str_finish[11:13]), int(str_finish[14:16]))
                if CourseLessons.objects.filter(teacher_id=teacher.id). \
                        filter(lesson_time_start__lte=timestart). \
                        filter(lesson_time_finish__gte=timestart).exists():
                    check_free = False
                if CourseLessons.objects.filter(teacher_id=teacher.id). \
                        filter(lesson_time_start__lte=timefinish). \
                        filter(lesson_time_finish__gte=timefinish).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=teacher.id). \
                        filter(lesson_time_start__lte=timestart). \
                        filter(lesson_time_finish__gte=timestart).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=teacher.id). \
                        filter(lesson_time_start__lte=timefinish). \
                        filter(lesson_time_finish__gte=timefinish).exists():
                    check_free = False
                return JsonResponse({
                    'fio': teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic,
                    'email': User.objects.get(id=teacher.user_id).email,
                    'phone': teacher.phone,
                    'freetime': check_free,
                })
            else:
                pass
        else:
            gr = StudentGroups.objects.get(id=self.kwargs.get('group'))
            if Events.objects.get(id=gr.event_id).department != GetUserDepAD(self.request):
                return HttpResponseRedirect('/access_denied/')
            return render(request, 'dep/study/schedule/lessons_event.html', context={
                'lessons': EventsLessons.objects.filter(group_id=self.kwargs.get('group')).order_by('lesson_time_start'),
                'group': StudentGroups.objects.get(id=self.kwargs.get('group'))
            })

    def post(self, request, *args, **kwargs):
        if 'delete_id' in request.POST:
            try:
                EventsLessons.objects.get(id=request.POST.get('delete_id')).delete()
                messages.success(request, 'Занятие успешно удалено')
            except BaseException:
                messages.error(request, 'Произошшла ошибка, повторите попытку позже')
        else:
            try:
                gr = request.POST.get('group_id')
                lesson_day = request.POST.get('lesson_day')
                str_start = request.POST.get('lesson_time_start')
                str_finish = request.POST.get('lesson_time_finish')
                timestart = datetime(int(lesson_day[:4]), int(lesson_day[5:7]), int(lesson_day[8:10]),
                                              int(str_start[:2]), int(str_start[3:5]))
                timefinish = datetime(int(lesson_day[:4]), int(lesson_day[5:7]), int(lesson_day[8:10]),
                                               int(str_finish[:2]), int(str_finish[3:5]))
                str_teacher = request.POST.get('teacher')
                new = EventsLessons()
                new.group_id = int(self.kwargs.get('group'))
                new.theme = request.POST.get('theme')
                new.lecture_hours = int(request.POST.get('lecture_hours'))
                new.practice_hours = int(request.POST.get('practice_hours'))
                if EventsLessons.objects.filter(group_id=gr). \
                        filter(lesson_time_start__lte=timestart).filter(lesson_time_finish__gt=timestart).exists():
                    check = True
                    CountPrevLessons = None
                    TargetLesson = None
                    NextLessons = None
                    if EventsLessons.objects. \
                            filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__lt=timestart).exists():
                        CountPrevLessons = EventsLessons.objects. \
                            filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__lt=timestart).count()
                        TargetLesson = EventsLessons.objects.filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__lt=timestart).latest('lesson_time_start')
                    if EventsLessons.objects.filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__gte=timestart).exists():
                        NextLessons = EventsLessons.objects.filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__gte=timestart).order_by('lesson_time_start')
                    if NextLessons is not None:
                        if TargetLesson is not None:
                            new_finish = TargetLesson.lesson_time_finish.replace(tzinfo=None)
                            if check is True:
                                if CountPrevLessons % 2 == 0:
                                    if CountPrevLessons % 4 == 0:
                                        new.lesson_time_start = new_finish + timedelta(hours=8, minutes=30)
                                        new.lesson_time_finish = new_finish + timedelta(hours=8, minutes=75)
                                    else:
                                        new.lesson_time_start = new_finish + timedelta(hours=8, minutes=10)
                                        new.lesson_time_finish = new_finish + timedelta(hours=8, minutes=55)
                                else:
                                    new.lesson_time_start = new_finish + timedelta(hours=8)
                                    new.lesson_time_finish = new_finish + timedelta(hours=8, minutes=45)
                            for curr, next in zip(NextLessons, NextLessons[1:]):
                                curr.lesson_time_start = next.lesson_time_start
                                curr.lesson_time_finish = next.lesson_time_finish
                                curr.save()
                            LastNext = NextLessons.latest('lesson_time_start')
                            if (CountPrevLessons + NextLessons.count()) % 2 == 0:
                                if (CountPrevLessons + NextLessons.count()) % 4 == 0:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=75)
                                else:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=55)
                            else:
                                LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                            LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                        else:
                            new.lesson_time_start = timestart
                            new.lesson_time_finish = timefinish
                            for curr, next in zip(NextLessons, NextLessons[1:]):
                                curr.lesson_time_start = next.lesson_time_start
                                curr.lesson_time_finish = next.lesson_time_finish
                                curr.save()
                            FirstNext = NextLessons.latest('-lesson_time_start')
                            FirstNext.lesson_time_start = timefinish
                            FirstNext.save()
                            FirstNext.lesson_time_finish = FirstNext.lesson_time_start + timedelta(minutes=45)
                            FirstNext.save()
                            if NextLessons.count() > 1:
                                LastNext = NextLessons.latest('lesson_time_start')
                                if NextLessons.count() % 2 == 0:
                                    if NextLessons.count() % 4 == 0:
                                        LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=75)
                                    else:
                                        LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=55)
                                else:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                                LastNext.save()
                                LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                                LastNext.save()
                    else:
                        new.lesson_time_start = timestart
                        new.lesson_time_finish = timefinish
                else:
                    new.lesson_time_start = timestart
                    new.lesson_time_finish = timefinish
                new.teacher_id = int(str_teacher[str_teacher.find(':') + 1:str_teacher.find(')')])
                new.save()
                messages.success(request, 'Занятие успешно добавлено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/dep/study/schedule/event_lessons_'+str(self.kwargs.get('group')))

    def get_context_data(self, **kwargs):
        context = super(EventLessonsList, self).get_context_data(**kwargs)
        context.update({
            'event': Events.objects.get(id=StudentGroups.objects.get(id=self.kwargs.get('group')).event_id)
        })
        return context


class PersonalSchedule(CheckTeacherMixin, ListView):
    "Расписание занятий преподавателя"
    login_url = '/'
    model = CourseLessons
    context_object_name = 'courselessons'
    template_name = 'profile/personal_schedule.html'

    def get_queryset(self):
        return CourseLessons.objects.filter(teacher_id=Profiles.objects.get(user_id=self.request.user.id).id)\
            .order_by('lesson_time_start')

    def get_context_data(self, **kwargs):
        context = super(PersonalSchedule, self).get_context_data(**kwargs)
        context.update({
            'eventlessons': EventsLessons.objects.filter(teacher_id=Profiles.objects.get(user_id=self.request.user.id).id).
                order_by('lesson_time_start'),
            'teacher': Profiles.objects.get(id=Profiles.objects.get(user_id=self.request.user.id).id)
        })
        return context


@login_required(login_url='/')
@group_required('Работники центра')
def PlanningChoose(request):
    return render(request, 'dep/study/planning/choose.html')

# Create your views here.

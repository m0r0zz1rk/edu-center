import os
import re
from datetime import timedelta
import openpyxl
from docxcompose.composer import Composer
import docx
import pymorphy2
from django.core.mail import EmailMessage
from django.utils import timezone
from django.utils.encoding import escape_uri_path, uri_to_iri
from django.views import View
from docx import Document
from docx.shared import Pt
from num2words import num2words
from openpyxl import Workbook
from pandas._libs.tslibs.offsets import BDay
from petrovich.main import Petrovich
from petrovich.enums import Case, Gender
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import User
from django.db.models import Q
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse, FileResponse
from django.shortcuts import render
from django.views.generic.edit import CreateView, UpdateView
from django.views.generic import DetailView, ListView
from docxtpl import DocxTemplate
from xlsxtpl.writerx import BookWriter
from authen.middleware import GetDepsWithManagerFromAD
from config.settings import STATIC_ROOT, MEDIA_ROOT
from students.models import CoursesForms, Apps, EventsForms, Statuses, Docs, DocsTypes
from .forms import ProgramForm, CoursesForm, EventsForm
from .models import *
import locale
from .tasks import EmailOfferPay, EmailAcceptPay, EmailDeniedPay, generate_pk1, frdo_report

locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')


class CheckAdminMixin(LoginRequiredMixin):
    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Администраторы').exists():
            return HttpResponseRedirect('/access_denied/')
        return super().dispatch(request, *args, **kwargs)


class CheckCentreMixin(LoginRequiredMixin):
    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name__in=['Работники центра', 'Администраторы']).exists():
            return HttpResponseRedirect('/access_denied/')
        return super().dispatch(request, *args, **kwargs)


class EventTypesList(CheckAdminMixin, ListView):
    login_url = '/'
    model = EventTypes
    context_object_name = 'types'
    template_name = 'centre/guides/eventtypes.html'

    def get_queryset(self):
        return EventTypes.objects.all().order_by('-id')

    def post(self, request, *args, **kwargs):
        try:
            if 'create_type' in request.POST:
                new = EventTypes()
                new.name = request.POST.get('create_type')
                new.save()
                messages.success(request, 'Тип мероприятия успешно добавлен')
            elif 'delete_type' in request.POST:
                EventTypes.objects.get(id=request.POST.get('delete_type')).delete()
                messages.success(request, 'Тип мероприятия успешно удален')
            else:
                pass
        except BaseException:
            messages.success(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/guides/eventtypes')


class OosListView(CheckAdminMixin, ListView):
    login_url = '/'
    model = Oos
    context_object_name = 'oos'
    template_name = 'centre/guides/oos.html'

    def get_queryset(self):
        qs = Oos.objects.all().order_by('-id')
        if self.request.GET:
            if 'check_mos[]' in self.request.GET:
                chk = self.request.GET.getlist('check_mos[]')
                if len(chk) != 0:
                    qs = qs.filter(mo_id__in=chk)
            if 'check_types[]' in self.request.GET:
                chk = self.request.GET.getlist('check_types[]')
                if len(chk) != 0:
                    qs = qs.filter(type_oo_id__in=chk)
            if 'find_short' in self.request.GET:
                qs = qs.filter(short_name__contains=self.request.GET.get('find_short'))
            if 'find_full' in self.request.GET:
                qs = qs.filter(full_name__contains=self.request.GET.get('find_full'))
        return qs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.GET:
            context['find'] = 'yes'
        return context


class PositionsListView(CheckAdminMixin, ListView):
    login_url = '/'
    model = Positions
    context_object_name = 'positions'
    template_name = 'centre/guides/positions.html'

    def get_queryset(self):
        return Positions.objects.all().order_by('name')


class PosCategoriesListView(CheckAdminMixin, ListView):
    login_url = '/'
    model = PositionCategories
    context_object_name = 'categories'
    template_name = 'centre/guides/poscats.html'

    def get_queryset(self):
        return PositionCategories.objects.all().order_by('name')


class AudCategoriesListView(CheckAdminMixin, ListView):
    login_url = '/'
    model = AudienceCategories
    context_object_name = 'categories'
    template_name = 'centre/guides/audcats.html'

    def get_queryset(self):
        return AudienceCategories.objects.all().order_by('name')


class UsersListView(CheckAdminMixin, ListView):
    login_url = '/'
    model = Profiles
    context_object_name = 'users'
    template_name = 'centre/guides/users.html'

    def get_queryset(self):
        qs = Profiles.objects.all().order_by('-id')
        if self.request.GET:
            if 'check_states[]' in self.request.GET:
                chk = self.request.GET.getlist('check_states[]')
                if len(chk) != 0:
                    qs = qs.filter(state_id__in=chk)
            if 'select' in self.request.GET:
                if self.request.GET.get('select') == 'Преподаватели':
                    qs = qs.filter(teacher=True)
            if 'find_surname' in self.request.GET:
                qs = qs.filter(surname__contains=self.request.GET.get('find_surname'))
            if 'find_name' in self.request.GET:
                qs = qs.filter(name__contains=self.request.GET.get('find_name'))
            if 'find_patronymic' in self.request.GET:
                qs = qs.filter(name__contains=self.request.GET.get('find_patronymic'))
            if 'find_email' in self.request.GET:
                if User.objects.filter(email=self.request.GET.get('find_email')).exists():
                    qs = qs.filter(user_id=User.objects.get(email=self.request.GET.get('find_email')))
            if 'find_phone' in self.request.GET:
                if qs.filter(phone=self.request.GET.get('find_phone')).exists():
                    qs = qs.get(phone=self.request.GET.get('find_phone'))
            if 'find_snils' in self.request.GET:
                if qs.filter(snils=self.request.GET.get('find_snils')).exists():
                    qs = qs.get(snils=self.request.GET.get('find_snils'))
        return qs[:25]


class UserDetailView(CheckAdminMixin, DetailView):
    login_url = '/'
    model = Profiles
    context_object_name = 'user'
    template_name = 'centre/guides/user_detail.html'


class ProgramsListView(CheckAdminMixin, ListView):
    login_url = '/'
    model = Programs
    context_object_name = 'programs'
    template_name = 'centre/study/programs/list.html'

    def get_queryset(self):
        qs = Programs.objects.all().order_by('-id')
        if 'check_deps[]' in self.request.GET:
            chk = self.request.GET.getlist('check_deps[]')
            if len(chk) != 0:
                qs = qs.filter(department__in=chk)
        if 'check_cats[]' in self.request.GET:
            chk = self.request.GET.getlist('check_cats[]')
            if len(chk) != 0:
                qs = qs.filter(categories__in=chk)
        if 'name' in self.request.GET:
            qs = qs.filter(name__contains=self.request.GET.get('name'))
        if 'order_id' in self.request.GET:
            qs = qs.filter(order_id__contains=self.request.GET.get('order_id'))
        return qs[:25]


class ProgramCreateView(CheckAdminMixin, CreateView):
    login_url = '/'
    model = Programs
    context_object_name = 'program'
    template_name = 'centre/study/programs/detail.html'

    def get(self, request, *args, **kwargs):
        return render(request, 'centre/study/programs/detail.html', {'form': ProgramForm})

    def post(self, request, *args, **kwargs):
        form = ProgramForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, 'Новая программа успешно добавлена')
        else:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
        return HttpResponseRedirect('/centre/study/programs')


class ProgramDetailView(CheckAdminMixin, DetailView):
    login_url = '/'
    form_class = ProgramForm
    context_object_name = 'program'
    template_name = 'centre/study/programs/edit.html'

    def get_object(self, queryset=None):
        return Programs.objects.get(id=self.kwargs.get('pk'))

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Информация успешно обновлена')
        return HttpResponseRedirect('/centre/study/programs')

    def get_context_data(self, **kwargs):

        context = super(ProgramDetailView, self).get_context_data(**kwargs)
        prog = Programs.objects.get(id=self.kwargs.get('pk'))
        data = {
            'department': prog.department,
            'name': prog.name,
            'type_dpp': prog.type_dpp,
            'duration': prog.duration,
            'categories': prog.categories.all(),
            'annotation': prog.annotation,
            'order_id': prog.order_id,
            'order_date': prog.order_date.isoformat(),
            'order_file': prog.order_file,
            'price': prog.price
        }
        form = ProgramForm(initial=data)
        context['form'] = form
        context['id_prog'] = self.kwargs.get('pk')
        return context


class CoursesList(CheckAdminMixin, ListView):
    login_url = '/'
    model = Courses
    context_object_name = 'courses'
    template_name = 'centre/study/planning/courses/list.html'

    def get_queryset(self):
        qs = Courses.objects.all().order_by('-id')
        if 'name' in self.request.GET:
            if len(self.request.GET.get('name')) > 0:
                qs = qs.filter(program__in=Programs.objects.filter(name__contains=self.request.GET.get('name')))
        if 'duration' in self.request.GET:
            if len(self.request.GET.get('duration')) > 0:
                qs = qs.filter(program__in=Programs.objects.filter(duration=self.request.GET.get('duration')))
        if 'study_date' in self.request.GET:
            try:
                post = self.request.GET.get('study_date')
                date = datetime.date(int(post[:4]), int(post[-5:7]), int(post[8:]))
                qs = qs.filter(date_start__lte=date, date_finish__gte=date)
            except BaseException:
                pass
        return qs[:25]

    def post(self, request):
        try:
            course = Courses.objects.get(id=request.POST.get('id'))
            course.place = request.POST.get('place')
            course.date_start = request.POST.get('date_start')
            course.date_finish = request.POST.get('date_finish')
            course.save()
            messages.success(request, 'Курс успешно изменен')
        except BaseException:
            messages.success(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/planning/courses')


class CoursesCreate(CheckAdminMixin, CreateView):
    login_url = '/'
    model = Courses
    context_object_name = 'course'
    template_name = 'centre/study/planning/courses/detail.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'al' in request.GET:
                return JsonResponse({
                    'val': PlanningParameters.objects.get(alias=request.GET.get('al')).value
                })
            else:
                duration = Programs.objects.get(id=request.GET.get('id')).duration
                type_dpp = Programs.objects.get(id=request.GET.get('id')).type_dpp
                date_order = Programs.objects.get(id=request.GET.get('id')).order_date.strftime('%d.%m.%Y')
                return JsonResponse({
                    'duration': str(duration),
                    'type_dpp': type_dpp,
                    'date_order': date_order
                })
        else:
            return render(request, 'centre/study/planning/courses/detail.html', {
                'form': CoursesForm(),
            })

    def post(self, request, *args, **kwargs):
        form = CoursesForm(request.POST)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, 'Курс успешно сохранен')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        else:
            messages.error(request, 'Некорректно заполнены данные нового курса')
        return HttpResponseRedirect('/centre/study/planning/courses')


class EventsList(CheckAdminMixin, ListView):
    login_url = '/'
    model = Events
    context_object_name = 'events'
    template_name = 'centre/study/planning/events/list.html'

    def get_queryset(self):
        qs = Events.objects.all().order_by('-id')
        if 'name' in self.request.GET:
            if len(self.request.GET.get('name')) > 0:
                qs = qs.filter(name__contains=self.request.GET.get('name'))
        if 'duration' in self.request.GET:
            if len(self.request.GET.get('duration')) > 0:
                qs = qs.filter(duration=int(self.request.GET.get('duration')))
        if 'study_date' in self.request.GET:
            try:
                post = self.request.GET.get('study_date')
                date = datetime.date(int(post[:4]), int(post[-5:7]), int(post[8:]))
                qs = qs.filter(date_start__lte=date, date_finish__gte=date)
            except BaseException:
                pass
        if 'check_deps[]' in self.request.GET:
            chk = self.request.GET.getlist('check_deps[]')
            if len(chk) != 0:
                qs = qs.filter(department__in=chk)
        if 'check_cats[]' in self.request.GET:
            chk = self.request.GET.getlist('check_cats[]')
            if len(chk) != 0:
                qs = qs.filter(categories__in=chk)
        return qs[:25]

    def post(self, request, *args, **kwargs):
        if 'id' in request.POST:
            try:
                Events.objects.get(id=request.POST.get('id')).delete()
                messages.success(request, 'Мероприятие успешно удалено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        else:
            pass
        return HttpResponseRedirect('/centre/study/planning/events')


class EventsCreate(CheckAdminMixin, CreateView):
    login_url = '/'
    model = Events
    context_object_name = 'event'
    template_name = 'centre/study/planning/events/detail.html'

    def get(self, request, *args, **kwargs):
        return render(request, 'centre/study/planning/events/detail.html', {
            'form': EventsForm(),
        })

    def post(self, request, *args, **kwargs):
        form = EventsForm(request.POST)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, 'Мероприятие успешно добавлено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        else:
            messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/planning/events')


class EventDetailView(CheckAdminMixin, UpdateView):
    login_url = '/'
    form_class = EventsForm
    context_object_name = 'event'
    template_name = 'centre/study/planning/events/detail.html'

    def get_object(self, queryset=None):
        return Events.objects.get(id=self.kwargs.get('pk'))

    def form_valid(self, form):
        form.save()
        messages.success(self.request, 'Мероприятие успешно изменено')
        return HttpResponseRedirect('/centre/study/planning/events')

    def get_context_data(self, **kwargs):
        context = super(EventDetailView, self).get_context_data(**kwargs)
        event = Events.objects.get(id=self.kwargs.get('pk'))
        start = str(event.date_start)
        finish = str(event.date_finish)
        data = {
            'department': event.department,
            'name': event.name,
            'type': event.type,
            'duration': event.duration,
            'categories': event.categories.all(),
            'place': event.place,
            'date_start': start[:4]+'-'+start[5:7]+'-'+start[8:10],
            'date_finish': finish[:4]+'-'+finish[5:7]+'-'+finish[8:10],
            'price': event.price
        }
        form = EventsForm(initial=data)
        context['form'] = form
        return context


class Item_event():
    def __init__(self, weekday, date, time, theme, lections, practice, teacher):
        self.weekday = weekday
        self.date = date
        self.time = time
        self.theme = theme
        self.lections =lections
        self.practice = practice
        self.teacher = teacher


class Item_course():
    def __init__(self, weekday, date, time, theme, lections, practice, trainee, individual, teacher):
        self.weekday = weekday
        self.date = date
        self.time = time
        self.theme = theme
        self.lections =lections
        self.practice = practice
        self.trainee = trainee
        self.individual = individual
        self.teacher = teacher


def get_lessons(group_id):
    less = []
    if StudentGroups.objects.get(id=group_id).event is not None:
        lessons = EventsLessons.objects.filter(group_id=group_id).order_by('lesson_time_start')
        for lesson in lessons:
            time_st = lesson.lesson_time_start + timedelta(hours=8)
            time_fin = lesson.lesson_time_finish + timedelta(hours=8)
            l = Item_event(
                time_st.strftime('%A'),
                time_fin.strftime('%d.%m.%Y'),
                time_st.strftime('%H.%M') + '-' + time_fin.strftime('%H.%M'),
                lesson.theme,
                lesson.lecture_hours,
                lesson.practice_hours,
                lesson.teacher.surname + ' ' + lesson.teacher.name[:1] + '.' + lesson.teacher.patronymic[:1] + '.'
            )
            less.append(l)
    else:
        lessons = CourseLessons.objects.filter(group_id=group_id).order_by('lesson_time_start')
        for lesson in lessons:
            time_st = lesson.lesson_time_start + timedelta(hours=8)
            time_fin = lesson.lesson_time_finish + timedelta(hours=8)
            indexes = [m.start() for m in re.finditer(' ', lesson.stschedule.name)]
            l = Item_course(
                time_st.strftime('%A'),
                time_fin.strftime('%d.%m.%Y'),
                time_st.strftime('%H.%M') + '-' + time_fin.strftime('%H.%M'),
                lesson.stschedule.name[indexes[1]:],
                lesson.lecture_hours,
                lesson.practice_hours,
                lesson.trainee_hours,
                lesson.individual_hours,
                lesson.teacher.surname + ' ' + lesson.teacher.name[:1] + '.' + lesson.teacher.patronymic[:1] + '.'
            )
            less.append(l)
    return less


class ShedulesList(CheckAdminMixin, ListView):
    login_url = '/'
    model = StudentGroups
    context_object_name = 'groups'
    template_name = 'centre/study/schedule/list.html'

    def get_queryset(self):
        request = self.request
        qs = StudentGroups.objects.all().order_by('-id')
        if 'code' in self.request.GET:
            if len(request.GET.get('code')) > 0:
                qs = qs.filter(code__contains=request.GET.get('code'))
        if 'name' in self.request.GET:
            if len(request.GET.get('name')) > 0:
                courses = Courses.objects.filter(program__in=Programs.objects.filter(name__contains=request.GET.get('name')))
                events = Events.objects.filter(name__contains=request.GET.get('name'))
                qs = qs.filter(Q(event__in=events) | Q(course__in=courses))
        if 'study_date' in self.request.GET:
            try:
                post = self.request.GET.get('study_date')
                date = datetime.date(int(post[:4]), int(post[-5:7]), int(post[8:]))
                courses = Courses.objects.filter(date_start__lte=date, date_finish__gte=date)
                events = Events.objects.filter(date_start__lte=date, date_finish__gte=date)
                qs = qs.filter(Q(event__in=events) | Q(course__in=courses))
            except BaseException:
                pass
        return qs[:25]

    def post(self, request, *args, **kwargs):
        group = StudentGroups.objects.get(id=request.POST.get('group_id'))
        if group.course is None:
            dep = group.event.department
            date_start = group.event.date_start
            date_finish = group.event.date_finish
            event_cats = ''
            for cat in group.event.categories.all():
                event_cats += cat.name+'; '
            lessons = EventsLessons.objects.filter(group_id=group.id)
            total_lecture = total_practice = total_hours = 0
            for el in lessons:
                total_lecture += el.lecture_hours
                total_practice += el.practice_hours
            total_hours = total_lecture + total_practice
            short_dep = ''
            list = re.split(' |-', dep)
            for el in list:
                short_dep += el[:1].upper()
            deps = GetDepsWithManagerFromAD()
            for key,value in deps.items():
                if key == dep:
                    fio = value.split(' ')
                    io_family_manager = fio[1][:1]+'.'+fio[2][:1]+'. '+fio[0]
            path = STATIC_ROOT + '\\doc_templates\\xlsx\\schedule_event.xlsx'
            writer = BookWriter(path)
            info = {
                'dep': dep,
                'code': group.code,
                'event_name': group.event.name,
                'events_cats': event_cats,
                'day_start': date_start.strftime('%d'),
                'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                'year_start': date_start.strftime('%Y'),
                'day_finish': date_finish.strftime('%d'),
                'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                'year_finish': date_finish.strftime('%Y'),
                'total_lecture': total_lecture,
                'total_practice': total_practice,
                'total_hours': total_hours,
                'short_dep': short_dep,
                'io_family_manager': io_family_manager
            }
        else:
            dep = group.course.program.department
            date_start = group.course.date_start
            date_finish = group.course.date_finish
            course_cats = ''
            for cat in group.course.program.categories.all():
                course_cats += cat.name + '; '
            lessons = CourseLessons.objects.filter(group_id=group.id)
            total_lecture = total_practice = total_individual = total_trainee = 0
            for el in lessons:
                total_lecture += el.lecture_hours
                total_practice += el.practice_hours
                total_trainee += el.trainee_hours
                total_individual += el.individual_hours
            total_hours = total_lecture + total_practice + total_trainee + total_individual
            short_dep = ''
            list = re.split(' |-', dep)
            for el in list:
                short_dep += el[:1].upper()
            deps = GetDepsWithManagerFromAD()
            for key, value in deps.items():
                if key == dep:
                    fio = value.split(' ')
                    io_family_manager = fio[1][:1] + '.' + fio[2][:1] + '. ' + fio[0]
            path = STATIC_ROOT + '\\doc_templates\\xlsx\\schedule_course.xlsx'
            writer = BookWriter(path)
            info = {
                'dep': dep,
                'code': group.code,
                'course_name': group.course.program.name,
                'course_cats': course_cats,
                'day_start': date_start.strftime('%d'),
                'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                'year_start': date_start.strftime('%Y'),
                'day_finish': date_finish.strftime('%d'),
                'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                'year_finish': date_finish.strftime('%Y'),
                'total_lecture': total_lecture,
                'total_practice': total_practice,
                'total_individual': total_individual,
                'total_hours': total_hours,
                'short_dep': short_dep,
                'io_family_manager': io_family_manager
            }
        info['items'] = get_lessons(group.id)
        payloads = [info]
        writer.render_book(payloads=payloads)
        newpath = MEDIA_ROOT + '\\Расписания\\'+group.code
        if not os.path.exists(newpath):
            os.makedirs(newpath)
        writer.save(MEDIA_ROOT + '\\Расписания\\'+group.code + "\\Расписание.xlsx")
        filename = MEDIA_ROOT + '\\Расписания\\'+group.code + "\\Расписание.xlsx"
        strin = "Расписание_"+group.code+".xlsx"
        data = open(filename, "rb").read()
        response = HttpResponse(data,
                                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Length'] = os.path.getsize(filename)
        response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
        return response


class CourseLessonsList(CheckAdminMixin, ListView):
    login_url = '/'
    model = CourseLessons
    context_object_name = 'lessons'
    template_name = 'centre/study/schedule/lessons_course.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'timestart' in request.GET:
                    check_timestart = True
                    check_break = True
                    check_lunch = True
                    if CourseLessons.objects.filter(group_id=self.kwargs.get('group')).exists():
                        str_start = request.GET.get('timestart')
                        timestart = datetime.datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                                      int(str_start[11:13])-8, int(str_start[14:16]))
                        tf = timestart + timedelta(hours=8, minutes=45)
                        timezone.make_aware(timestart, timezone.get_default_timezone())
                        lessons = CourseLessons.objects.exclude(id=request.GET.get('lesson')).filter(group_id=self.kwargs.get('group')).order_by('lesson_time_start')
                        for lesson in lessons:
                            t_s = lesson.lesson_time_start.replace(tzinfo=None)
                            t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                            if t_s <= timestart < t_f:
                                check_timestart = False
                                break
                        for previous, current in zip(lessons, lessons[1:]):
                            t_s = current.lesson_time_start.replace(tzinfo=None)
                            t_f = previous.lesson_time_finish.replace(tzinfo=None)
                            if t_f <= timestart <= t_s:
                                main_tf = t_f
                                check_twohours = False
                                for les in lessons:
                                    if les.lesson_time_finish == previous.lesson_time_start:
                                        check_twohours = True
                                if check_twohours is True:
                                    break_between = timestart - previous.lesson_time_finish.replace(tzinfo=None)
                                    if break_between.seconds//60 < 10:
                                        check_break = False
                                        break
                        if check_timestart is True and check_break is True:
                            count_hours = 0
                            for lesson in lessons:
                                t_s = lesson.lesson_time_start.replace(tzinfo=None)
                                t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                                if t_f.date() == timestart.date() and t_s < timestart:
                                    count_hours += 1
                                if count_hours >= 4:
                                    was_lunch = False
                                    for previous, current in zip(lessons, lessons[1:]):
                                        t_s = current.lesson_time_start.replace(tzinfo=None)
                                        t_f = previous.lesson_time_finish.replace(tzinfo=None)
                                        if t_s.date() == t_f.date() and t_s < timestart:
                                            br = t_s - t_f
                                            if br.seconds//60 >= 30:
                                                was_lunch = True
                                                break
                                    if was_lunch is False:
                                        break_lunch = timestart - main_tf
                                        if break_lunch.seconds//60 < 30:
                                            check_lunch = False
                        else:
                            pass
                    return JsonResponse({
                        'check_timestart': check_timestart,
                        'check_break': check_break,
                        'check_lunch': check_lunch,
                        'tf': tf.strftime('%H:%M')
                    })
            elif 'timefinish' in request.GET:
                check_timefinish = True
                if CourseLessons.objects.exclude(id=request.GET.get('lesson')).filter(group_id=self.kwargs.get('group')).exists():
                    str_finish = request.GET.get('timefinish')
                    timefinish = datetime.datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                                  int(str_finish[11:13]), int(str_finish[14:16]))
                    if CourseLessons.objects.exclude(id=request.GET.get('lesson')).filter(group_id=self.kwargs.get('group'))\
                            .filter(lesson_time_start__lt=timefinish).\
                            filter(lesson_time_finish__gt=timefinish).exists():
                        check_timefinish = False
                    else:
                        pass
                else:
                    pass
                return JsonResponse({'check_timefinish': check_timefinish})
            elif 'get_themes' in request.GET:
                if StSchedule.objects.filter(program_id=StudentGroups.objects.get(id=self.kwargs.get('group')).course.program.id).exists():
                    themes = StSchedule.objects.filter(program_id=StudentGroups.objects.get(id=self.kwargs.get('group')).course.program.id)
                    dict_themes = {}
                    count = 1
                    for theme in themes:
                        list_th = []
                        list_th.append(theme.id)
                        list_th.append(theme.name)
                        dict_themes[count] = list_th
                        count += 1
                    return JsonResponse({
                        'themes': dict_themes
                    })
                else:
                    return JsonResponse({})
            elif 'theme_id' in request.GET:
                theme = StSchedule.objects.get(id=request.GET.get('theme_id'))
                l_hours = theme.lecture_hours
                p_hours = theme.practice_hours
                t_hours = theme.trainee_hours
                i_hours = theme.individual_hours
                if CourseLessons.objects.filter(group_id=self.kwargs.get('group')).\
                        filter(stschedule_id=request.GET.get('theme_id')).exists():
                    lessons = CourseLessons.objects.filter(group_id=self.kwargs.get('group')).\
                        filter(stschedule_id=request.GET.get('theme_id'))
                    for lesson in lessons:
                        l_hours -= lesson.lecture_hours
                        p_hours -= lesson.practice_hours
                        t_hours -= lesson.trainee_hours
                        i_hours -= lesson.individual_hours
                teachers = Profiles.objects.filter(teacher=True)
                list_t = []
                for teacher in teachers:
                    list_t.append(teacher.surname+' '+teacher.name+' '+teacher.patronymic+' (ID:'+str(teacher.id)+')')
                return JsonResponse({
                    'lecture': l_hours,
                    'practice': p_hours,
                    'individual': i_hours,
                    'trainee': t_hours,
                    'list_t': list_t
                })
            elif 'id_teach' in request.GET:
                str_start = request.GET.get('time_start')
                str_finish = request.GET.get('time_finish')
                teacher = Profiles.objects.get(id=int(request.GET.get('id_teach')))
                check_free = True
                timestart = datetime.datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]), int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime.datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]), int(str_finish[11:13]), int(str_finish[14:16]))
                if CourseLessons.objects.filter(teacher_id=request.GET.get('id_teach')).\
                        filter(lesson_time_start__lte=timestart).\
                        filter(lesson_time_finish__gt=timestart).exists():
                    check_free = False
                if CourseLessons.objects.filter(teacher_id=request.GET.get('id_teach')).\
                        filter(lesson_time_start__lte=timefinish).\
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('id_teach')).\
                        filter(lesson_time_start__lte=timestart).\
                        filter(lesson_time_finish__gt=timestart).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('id_teach')).\
                        filter(lesson_time_start__lte=timefinish).\
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check_free = False
                return JsonResponse({
                    'fio': teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic,
                    'email': User.objects.get(id=teacher.user_id).email,
                    'phone': teacher.phone,
                    'freetime': check_free,
                })
            elif 'form_study' in request.GET:
                group = StudentGroups.objects.get(id=request.GET.get('group'))
                if request.GET.get('form_study') == 'without_dot':
                    group.study_form = 'Без использования ДОТ'
                elif request.GET.get('form_study') == 'only_dot':
                    group.study_form = 'Исключительно ДОТ'
                else:
                    group.study_form = 'С использованием ДОТ'
                group.save()
                return JsonResponse({})
            elif 'change_lesson' in request.GET:
                less = CourseLessons.objects.select_related('stschedule').get(id=request.GET.get('change_lesson'))
                group = StudentGroups.objects.select_related('course').get(id=less.group_id)
                date_s = group.course.date_start
                date_f = group.course.date_finish
                sts = StSchedule.objects.get(id=less.stschedule_id)
                teachs = Profiles.objects.filter(teacher=True).order_by('surname')
                if less.lecture_hours == less.practice_hours == less.trainee_hours == 0:
                    if less.control == '':
                        type = 'Самостоятельная работа'
                    else:
                        type = 'Контрольное занятие'
                elif less.individual_hours == less.practice_hours == less.trainee_hours == 0:
                    type = 'Лекция'
                elif less.individual_hours == less.lecture_hours == less.trainee_hours == 0:
                    if less.control == '':
                        type = 'Практика'
                    else:
                        type = 'Контрольное занятие'
                elif less.individual_hours == less.lecture_hours == less.practice_hours == 0:
                    type = 'Стажировка'
                else:
                    type = 'Смешанное занятие'
                dict_teachs = {}
                current_teacher = ''
                for teach in teachs:
                    if less.teacher_id == teach.id:
                        current_teacher = teach.surname + ' ' + teach.name + ' ' + teach.patronymic+' (ID:'+str(teach.id)+')'
                    dict_teachs[teach.id] = teach.surname + ' ' + teach.name + ' ' + teach.patronymic

                start = less.lesson_time_start.replace(tzinfo=None)
                finish = less.lesson_time_finish.replace(tzinfo=None)
                start = start + timedelta(hours=8)
                finish = finish + timedelta(hours=8)
                return JsonResponse({
                    'date_s': date_s.strftime('%d.%m.%Y'),
                    'date_f': date_f.strftime('%d.%m.%Y'),
                    'theme_name': less.stschedule.name,
                    'date_start': less.lesson_time_start.strftime('%Y-%m-%d'),
                    'time_start': start.strftime('%H:%M'),
                    'time_finish': finish.strftime('%H:%M'),
                    'distance': less.distance,
                    'control': less.control,
                    'teachers': dict_teachs,
                    'current_teacher': current_teacher,
                    'type': type
                })
                #else:
                  #  return HttpResponseRedirect('/access_denied/')
            elif 'check_date' in request.GET:
                course = StudentGroups.objects.get(id=self.kwargs.get('group')).course
                d = datetime.date(int(request.GET.get('check_date')[:4]), int(request.GET.get('check_date')[5:7]), int(request.GET.get('check_date')[8:10]))
                if course.date_start <= d <= course.date_finish:
                    check = True
                else:
                    check = False
                return JsonResponse({
                    'check': check
                })
            elif 'generate' in request.GET:
                group = StudentGroups.objects.select_related('course').get(id=self.kwargs.get('group'))
                str_form = request.GET.get('generate')
                max_day_hours = int(request.GET.get('count'))
                if not StSchedule.objects.filter(program_id=group.course.program.id).exists():
                    messages.error(request, 'Не найдены темы и разделы в КУГ')
                    return HttpResponseRedirect(
                        '/centre/study/schedule/course_lessons_' + str(self.kwargs.get('group')) + '#close')
                start_datetime = datetime.datetime(int(str_form[:4]),
                                                   int(str_form[5:7]),
                                                   int(str_form[8:10]),
                                                   int(str_form[11:13]),
                                                   int(str_form[14:16]))
                ls = CourseLessons.objects.filter(group_id=self.kwargs.get('group')). \
                    filter(lesson_time_start__year=start_datetime.year,
                           lesson_time_start__month=start_datetime.month,
                           lesson_time_start__day=start_datetime.day)
                for l in ls:
                    l.delete()
                day_hours = 0
                while day_hours < max_day_hours:
                    new_les = CourseLessons()
                    new_les.group_id = group.id
                    if day_hours % 2 == 0 and day_hours != 0:
                        if day_hours % 4 == 0:
                            new_les.lesson_time_start = start_datetime + timedelta(minutes=30)
                            new_les.lesson_time_finish = start_datetime + timedelta(minutes=75)
                            start_datetime = start_datetime + timedelta(minutes=75)
                        else:
                            new_les.lesson_time_start = start_datetime + timedelta(minutes=10)
                            new_les.lesson_time_finish = start_datetime + timedelta(minutes=55)
                            start_datetime = start_datetime + timedelta(minutes=55)
                    else:
                        new_les.lesson_time_start = start_datetime
                        new_les.lesson_time_finish = start_datetime + timedelta(minutes=45)
                        start_datetime = start_datetime + timedelta(minutes=45)
                    new_les.lecture_hours = 0
                    new_les.practice_hours = 0
                    new_les.trainee_hours = 0
                    new_les.individual_hours = 0
                    new_les.teacher = None
                    new_les.distance = False
                    new_les.control = ''
                    new_les.save()
                    day_hours += 1
                ls = CourseLessons.objects.filter(group_id=self.kwargs.get('group')).\
                    filter(lesson_time_start__year=start_datetime.year,
                           lesson_time_start__month=start_datetime.month,
                           lesson_time_start__day=start_datetime.day)
                dict_l = {}
                for l in ls:
                    list_l = []
                    list_l.append(l.id)
                    lts = l.lesson_time_start.replace(tzinfo=None) + timedelta(hours=8)
                    ltf = l.lesson_time_finish.replace(tzinfo=None) + timedelta(hours=8)
                    list_l.append(lts.strftime('%H:%M'))
                    list_l.append(ltf.strftime('%H:%M'))
                    dict_l[l.id] = list_l
                themes = StSchedule.objects.filter(program_id=Programs.objects.get(id=group.course.program.id).id)
                dict_th = {}
                for theme in themes:
                    if theme.parent_id is not None or theme.control_form != '':
                        list_th = []
                        list_th.append(theme.id)
                        if 'Раздел ' in theme.name:
                            list_th.append(theme.name[7:])
                        else:
                            list_th.append(theme.name[5:])
                        free_hours = theme.total_hours
                        if CourseLessons.objects.filter(group_id=self.kwargs.get('group')).filter(stschedule_id=theme.id).exists():
                            for l in CourseLessons.objects.filter(group_id=self.kwargs.get('group')).filter(stschedule_id=theme.id):
                                if l.lecture_hours != 0 and l.lecture_hours is not None:
                                    free_hours -= l.lecture_hours
                                if l.practice_hours != 0 and l.practice_hours is not None:
                                    free_hours -= l.practice_hours
                                if l.trainee_hours != 0 and l.trainee_hours is not None:
                                    free_hours -= l.trainee_hours
                                if l.individual_hours != 0 and l.individual_hours is not None:
                                    free_hours -= l.individual_hours
                        list_th.append(free_hours)
                        dict_th[theme.id] = list_th
                teachers = Profiles.objects.filter(teacher=True)
                dict_teach = {}
                for teacher in teachers:
                    dict_teach[teacher.id] = teacher.surname+' '+teacher.name+' '+teacher.patronymic+' (ID:'+str(teacher.id)+')'
                return JsonResponse({
                    'lessons': dict_l,
                    'themes': dict_th,
                    'teachers': dict_teach
                })
            elif 'getthemechoose' in request.GET:
                group = StudentGroups.objects.get(id=self.kwargs.get('group'))
                lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                theme = StSchedule.objects.filter(name__contains=request.GET.get('getthemechoose')).get(program_id=group.course.program.id)
                lect_h = theme.lecture_hours
                prac_h = theme.practice_hours
                trai_h = theme.trainee_hours
                indi_h = theme.individual_hours
                for less in CourseLessons.objects.filter(group_id=self.kwargs.get('group')):
                    if less.stschedule_id == theme.id:
                        if less.lecture_hours != 0 and less.lecture_hours is not None:
                            lect_h -= less.lecture_hours
                        if less.practice_hours != 0 and less.practice_hours is not None:
                            prac_h -= less.practice_hours
                        if less.trainee_hours != 0 and less.trainee_hours is not None:
                            trai_h -= less.trainee_hours
                        if less.individual_hours != 0 and less.individual_hours is not None:
                            indi_h -= less.individual_hours
                lesson.stschedule_id = theme.id
                lesson.lecture_hours = 0
                lesson.practice_hours = 0
                lesson.trainee_hours = 0
                lesson.individual_hours = 0
                lesson.control = theme.control_form
                lesson.save()
                return JsonResponse({
                    'lect_h': lect_h,
                    'prac_h': prac_h,
                    'trai_h': trai_h,
                    'indi_h': indi_h,
                    'control': theme.control_form
                })
            elif 'lestypechoose' in request.GET:
                group = StudentGroups.objects.get(id=self.kwargs.get('group'))
                lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                type = request.GET.get('lestypechoose')
                if type == 'lecture':
                    lesson.lecture_hours = 1
                    lesson.practice_hours = 0
                    lesson.trainee_hours = 0
                    lesson.individual_hours = 0
                elif type == 'practice':
                    lesson.lecture_hours = 0
                    lesson.practice_hours = 1
                    lesson.trainee_hours = 0
                    lesson.individual_hours = 0
                elif type == 'trainee':
                    lesson.lecture_hours = 0
                    lesson.practice_hours = 0
                    lesson.trainee_hours = 1
                    lesson.individual_hours = 0
                else:
                    lesson.lecture_hours = 0
                    lesson.practice_hours = 0
                    lesson.trainee_hours = 0
                    lesson.individual_hours = 1
                lesson.save()
                return JsonResponse({})
            elif 'lesformatchoose' in request.GET:
                lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                if request.GET.get('lesformatchoose') == 'dot':
                    lesson.distance = True
                else:
                    lesson.distance = False
                lesson.save()
                return JsonResponse({})
            elif 'teacher_gen' in request.GET:
                check = True
                str_start = request.GET.get('time_start')
                str_finish = request.GET.get('time_finish')
                timestart = datetime.datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime.datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                               int(str_finish[11:13]), int(str_finish[14:16]))
                if CourseLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')).\
                        filter(lesson_time_start__lte=timestart).\
                        filter(lesson_time_finish__gt=timestart).exists():
                    check = False
                if CourseLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')).\
                        filter(lesson_time_start__lte=timefinish).\
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')).\
                        filter(lesson_time_start__lte=timestart).\
                        filter(lesson_time_finish__gt=timestart).exists():
                    check = False
                if EventsLessons.objects.filter(teacher_id=request.GET.get('teacher_gen')).\
                        filter(lesson_time_start__lte=timefinish).\
                        filter(lesson_time_finish__gt=timefinish).exists():
                    check = False
                if check is True:
                    lesson = CourseLessons.objects.get(id=request.GET.get('lesson'))
                    lesson.teacher_id = request.GET.get('teacher_gen')
                    lesson.save()
                return JsonResponse({
                    'check': check
                })
            else:
                return HttpResponseRedirect('/access_denied/')
        else:
            if 'save_less' in request.GET:
                timestart = datetime.datetime(int(request.GET.get('date_less')[:4]),
                                              int(request.GET.get('date_less')[5:7]),
                                              int(request.GET.get('date_less')[8:10]),
                                              int(request.GET.get('time_start')[:2]),
                                              int(request.GET.get('time_start')[3:]))
                timefinish = datetime.datetime(int(request.GET.get('date_less')[:4]),
                                               int(request.GET.get('date_less')[5:7]),
                                               int(request.GET.get('date_less')[8:10]),
                                               int(request.GET.get('time_finish')[:2]),
                                               int(request.GET.get('time_finish')[3:]))
                lesson = CourseLessons.objects.get(id=request.GET.get('save_less'))
                LessonDay = datetime.date(int(request.GET.get('date_less')[:4]),
                                          int(request.GET.get('date_less')[5:7]),
                                          int(request.GET.get('date_less')[8:10]),)
                if CourseLessons.objects.filter(group_id=lesson.group_id).exclude(id=request.GET.get('save_less')).\
                        filter(lesson_time_start__lte=timestart).filter(lesson_time_finish__gt=timestart).exists():
                    check = True
                    CountPrevLessons = None
                    TargetLesson = None
                    NextLessons = None
                    if CourseLessons.objects.exclude(id=lesson.id).\
                    filter(group_id=lesson.group_id).\
                    filter(lesson_time_finish__day=LessonDay.day).\
                    filter(lesson_time_finish__month=LessonDay.month).\
                    filter(lesson_time_finish__year=LessonDay.year).\
                    filter(lesson_time_start__lt=timestart).exists():
                        CountPrevLessons = CourseLessons.objects.exclude(id=lesson.id).\
                            filter(group_id=lesson.group_id).\
                            filter(lesson_time_finish__day=LessonDay.day).\
                            filter(lesson_time_finish__month=LessonDay.month).\
                            filter(lesson_time_finish__year=LessonDay.year).\
                            filter(lesson_time_start__lt=timestart).count()
                        TargetLesson = CourseLessons.objects.filter(group_id=lesson.group_id). \
                            filter(lesson_time_finish__day=LessonDay.day). \
                            filter(lesson_time_finish__month=LessonDay.month). \
                            filter(lesson_time_finish__year=LessonDay.year). \
                            filter(lesson_time_start__lt=timestart).latest('lesson_time_start')
                    if CourseLessons.objects.filter(group_id=lesson.group_id).\
                        filter(lesson_time_finish__day=LessonDay.day).\
                        filter(lesson_time_finish__month=LessonDay.month).\
                        filter(lesson_time_finish__year=LessonDay.year).\
                        filter(lesson_time_start__gte=timestart).exists():
                            NextLessons = CourseLessons.objects.filter(group_id=lesson.group_id).\
                                filter(lesson_time_finish__day=LessonDay.day).\
                                filter(lesson_time_finish__month=LessonDay.month).\
                                filter(lesson_time_finish__year=LessonDay.year).\
                                filter(lesson_time_start__gte=timestart).order_by('lesson_time_start')
                    if NextLessons is not None:
                        if TargetLesson is not None:
                            new_finish = TargetLesson.lesson_time_finish.replace(tzinfo=None)
                            if check is True:
                                if CountPrevLessons % 2 == 0:
                                    if CountPrevLessons % 4 == 0:
                                        lesson.lesson_time_start = new_finish + timedelta(hours=8, minutes=30)
                                        lesson.lesson_time_finish = new_finish + timedelta(hours=8, minutes=75)
                                    else:
                                        lesson.lesson_time_start = new_finish + timedelta(hours=8, minutes=10)
                                        lesson.lesson_time_finish = new_finish + timedelta(hours=8, minutes=55)
                                else:
                                    lesson.lesson_time_start = new_finish + timedelta(hours=8)
                                    lesson.lesson_time_finish = new_finish + timedelta(hours=8, minutes=45)
                            for curr, next in zip(NextLessons, NextLessons[1:]):
                                curr.lesson_time_start = next.lesson_time_start
                                curr.lesson_time_finish = next.lesson_time_finish
                                curr.save()
                            LastNext = NextLessons.latest('lesson_time_start')
                            if (CountPrevLessons + NextLessons.count()) % 2 == 0:
                                if (CountPrevLessons + NextLessons.count()) % 4 == 0:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=75)
                                else:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=55)
                            else:
                                LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                            LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                        else:
                            lesson.lesson_time_start = timestart
                            lesson.lesson_time_finish = timefinish
                            for curr, next in zip(NextLessons, NextLessons[1:]):
                                curr.lesson_time_start = next.lesson_time_start
                                curr.lesson_time_finish = next.lesson_time_finish
                                curr.save()
                            FirstNext = NextLessons.latest('-lesson_time_start')
                            FirstNext.lesson_time_start = timefinish
                            FirstNext.save()
                            FirstNext.lesson_time_finish = FirstNext.lesson_time_start + timedelta(minutes=45)
                            FirstNext.save()
                            LastNext = NextLessons.latest('lesson_time_start')
                            if NextLessons.count() % 2 == 0:
                                if NextLessons.count() % 4 == 0:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=75)
                                else:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=55)
                            else:
                                LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                            LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                    else:
                        lesson.lesson_time_start = timestart
                        lesson.lesson_time_finish = timefinish
                else:
                    lesson.lesson_time_start = timestart
                    lesson.lesson_time_finish = timefinish
                str_teacher = request.GET.get('teacher')
                lesson.teacher_id = int(str_teacher[str_teacher.find(':') + 1:str_teacher.find(')')])
                lesson.control = request.GET.get('control')
                lesson.distance = request.GET.get('dist')
                lesson.save()
                messages.success(request, 'Информация о занятии успешно изменена')
                return HttpResponseRedirect('/centre/study/schedule/course_lessons_' + str(self.kwargs.get('group'))+'#close')
            return render(request, 'centre/study/schedule/lessons_course.html', context={
                'lessons': CourseLessons.objects.filter(group_id=self.kwargs.get('group')).order_by('lesson_time_start'),
                'group': StudentGroups.objects.get(id=self.kwargs.get('group')),
            })

    def post(self, request, *args, **kwargs):
        if 'delete_id' in request.POST:
            try:
                CourseLessons.objects.get(id=request.POST.get('delete_id')).delete()
                messages.success(request, 'Занятие успешно удалено')
            except BaseException:
                messages.error(request, 'Произошшла ошибка, повторите попытку позже')
        else:
            try:
                str_start = request.POST.get('lesson_day')+' '+request.POST.get('lesson_time_start')
                str_finish = request.POST.get('lesson_day')+' '+request.POST.get('lesson_time_finish')
                timestart = datetime.datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime.datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                               int(str_finish[11:13]), int(str_finish[14:16]))
                str_teacher = request.POST.get('teacher')
                new = CourseLessons()
                new.group_id = int(self.kwargs.get('group'))
                new.stschedule_id = int(request.POST.get('stschedule'))
                new.lecture_hours = int(request.POST.get('lecture_hours'))
                new.practice_hours = int(request.POST.get('practice_hours'))
                new.trainee_hours = int(request.POST.get('trainee_hours'))
                new.individual_hours = int(request.POST.get('individual_hours'))
                new.lesson_time_start = timestart
                new.lesson_time_finish = timefinish
                new.teacher_id = int(str_teacher[str_teacher.find(':')+1:str_teacher.find(')')])
                new.save()
                messages.success(request, 'Занятие успешно добавлено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/schedule/course_lessons_'+str(self.kwargs.get('group')))

    def get_context_data(self, **kwargs):
        context = super(CourseLessonsList, self).get_context_data(**kwargs)
        context.update({
            'course': Courses.objects.get(id=StudentGroups.objects.get(id=self.kwargs.get('group')).course_id)
        })
        return context


class EventLessonsList(CheckAdminMixin, ListView):
    login_url = '/'
    model = EventsLessons
    context_object_name = 'lessons'
    template_name = 'centre/study/schedule/lessons_event.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'timestart' in request.GET:
                check_timestart = True
                check_break = True
                check_lunch = True
                str_start = request.GET.get('timestart')
                timestart = datetime.datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]),
                                              int(str_start[11:13]) - 8, int(str_start[14:16]))
                if EventsLessons.objects.filter(group_id=self.kwargs.get('group')).exists():
                    timezone.make_aware(timestart, timezone.get_default_timezone())
                    lessons = EventsLessons.objects.exclude(id=request.GET.get('lesson')).filter(
                        group_id=self.kwargs.get('group')).order_by('lesson_time_start')
                    for lesson in lessons:
                        t_s = lesson.lesson_time_start.replace(tzinfo=None)
                        t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                        if t_s <= timestart < t_f:
                            check_timestart = False
                            break
                    for previous, current in zip(lessons, lessons[1:]):
                        t_s = current.lesson_time_start.replace(tzinfo=None)
                        t_f = previous.lesson_time_finish.replace(tzinfo=None)
                        if t_f <= timestart <= t_s:
                            main_tf = t_f
                            check_twohours = False
                            for les in lessons:
                                if les.lesson_time_finish == previous.lesson_time_start:
                                    check_twohours = True
                            if check_twohours is True:
                                break_between = timestart - previous.lesson_time_finish.replace(tzinfo=None)
                                if break_between.seconds // 60 < 10:
                                    check_break = False
                                    break
                    if check_timestart is True and check_break is True:
                        count_hours = 0
                        for lesson in lessons:
                            t_s = lesson.lesson_time_start.replace(tzinfo=None)
                            t_f = lesson.lesson_time_finish.replace(tzinfo=None)
                            if t_f.date() == timestart.date() and t_s < timestart:
                                count_hours += 1
                            if count_hours >= 4:
                                was_lunch = False
                                for previous, current in zip(lessons, lessons[1:]):
                                    t_s = current.lesson_time_start.replace(tzinfo=None)
                                    t_f = previous.lesson_time_finish.replace(tzinfo=None)
                                    if t_s.date() == t_f.date() and t_s < timestart:
                                        br = t_s - t_f
                                        if br.seconds // 60 >= 30:
                                            was_lunch = True
                                            break
                                if was_lunch is False:
                                    break_lunch = timestart - main_tf
                                    if break_lunch.seconds // 60 < 30:
                                        check_lunch = False
                    else:
                        pass
                return JsonResponse({
                    'check_timestart': check_timestart,
                    'check_break': check_break,
                    'check_lunch': check_lunch,
                })
            elif 'timefinish' in request.GET:
                check_timefinish = True
                if EventsLessons.objects.filter(group_id=self.kwargs.get('group')).exists():
                    str_finish = request.GET.get('timefinish')
                    timefinish = datetime.datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]),
                                                  int(str_finish[11:13]), int(str_finish[14:16]))
                    if EventsLessons.objects.filter(group_id=self.kwargs.get('group'))\
                            .filter(lesson_time_start__lte=timefinish).\
                            filter(lesson_time_finish__gte=timefinish).exists():
                        check_timefinish = False
                teachers = Profiles.objects.filter(teacher=True)
                list_t = []
                for teacher in teachers:
                    list_t.append(teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic + ' (ID:' + str(
                        teacher.id) + ')')
                return JsonResponse({
                    'check_timefinish': check_timefinish,
                    'list_t': list_t
                })
            elif 'id_teach' in request.GET:
                str_start = request.GET.get('time_start')
                str_finish = request.GET.get('time_finish')
                teacher = Profiles.objects.get(id=int(request.GET.get('id_teach')))
                check_free = True
                timestart = datetime.datetime(int(str_start[:4]), int(str_start[5:7]), int(str_start[8:10]), int(str_start[11:13]), int(str_start[14:16]))
                timefinish = datetime.datetime(int(str_finish[:4]), int(str_finish[5:7]), int(str_finish[8:10]), int(str_finish[11:13]), int(str_finish[14:16]))
                if CourseLessons.objects.filter(teacher_id=teacher.id).\
                        filter(lesson_time_start__lte=timestart).\
                        filter(lesson_time_finish__gte=timestart).exists():
                    check_free = False
                if CourseLessons.objects.filter(teacher_id=teacher.id).\
                        filter(lesson_time_start__lte=timefinish).\
                        filter(lesson_time_finish__gte=timefinish).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=teacher.id).\
                        filter(lesson_time_start__lte=timestart).\
                        filter(lesson_time_finish__gte=timestart).exists():
                    check_free = False
                if EventsLessons.objects.filter(teacher_id=teacher.id).\
                        filter(lesson_time_start__lte=timefinish).\
                        filter(lesson_time_finish__gte=timefinish).exists():
                    check_free = False
                return JsonResponse({
                    'fio': teacher.surname + ' ' + teacher.name + ' ' + teacher.patronymic,
                    'email': User.objects.get(id=teacher.user_id).email,
                    'phone': teacher.phone,
                    'freetime': check_free,
                })
            else:
                pass
        else:
            return render(request, 'centre/study/schedule/lessons_event.html', context={
                'lessons': EventsLessons.objects.filter(group_id=self.kwargs.get('group')).order_by('lesson_time_start'),
                'group': StudentGroups.objects.get(id=self.kwargs.get('group'))
            })

    def post(self, request, *args, **kwargs):
        if 'delete_id' in request.POST:
            try:
                EventsLessons.objects.get(id=request.POST.get('delete_id')).delete()
                messages.success(request, 'Занятие успешно удалено')
            except BaseException:
                messages.error(request, 'Произошшла ошибка, повторите попытку позже')
        else:
            try:
                gr = request.POST.get('group_id')
                lesson_day = request.POST.get('lesson_day')
                str_start = request.POST.get('lesson_time_start')
                str_finish = request.POST.get('lesson_time_finish')
                timestart = datetime.datetime(int(lesson_day[:4]), int(lesson_day[5:7]), int(lesson_day[8:10]),
                                              int(str_start[:2]), int(str_start[3:5]))
                timefinish = datetime.datetime(int(lesson_day[:4]), int(lesson_day[5:7]), int(lesson_day[8:10]),
                                               int(str_finish[:2]), int(str_finish[3:5]))
                str_teacher = request.POST.get('teacher')
                new = EventsLessons()
                new.group_id = int(self.kwargs.get('group'))
                new.theme = request.POST.get('theme')
                new.lecture_hours = int(request.POST.get('lecture_hours'))
                new.practice_hours = int(request.POST.get('practice_hours'))
                if EventsLessons.objects.filter(group_id=gr).\
                        filter(lesson_time_start__lte=timestart).filter(lesson_time_finish__gt=timestart).exists():
                    check = True
                    CountPrevLessons = None
                    TargetLesson = None
                    NextLessons = None
                    if EventsLessons.objects.\
                            filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__lt=timestart).exists():
                        CountPrevLessons = EventsLessons.objects.\
                            filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__lt=timestart).count()
                        TargetLesson = EventsLessons.objects.filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__lt=timestart).latest('lesson_time_start')
                    if EventsLessons.objects.filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__gte=timestart).exists():
                        NextLessons = EventsLessons.objects.filter(group_id=gr). \
                            filter(lesson_time_finish__day=timefinish.day). \
                            filter(lesson_time_finish__month=timefinish.month). \
                            filter(lesson_time_finish__year=timefinish.year). \
                            filter(lesson_time_start__gte=timestart).order_by('lesson_time_start')
                    if NextLessons is not None:
                        if TargetLesson is not None:
                            new_finish = TargetLesson.lesson_time_finish.replace(tzinfo=None)
                            if check is True:
                                if CountPrevLessons % 2 == 0:
                                    if CountPrevLessons % 4 == 0:
                                        new.lesson_time_start = new_finish + timedelta(hours=8, minutes=30)
                                        new.lesson_time_finish = new_finish + timedelta(hours=8, minutes=75)
                                    else:
                                        new.lesson_time_start = new_finish + timedelta(hours=8, minutes=10)
                                        new.lesson_time_finish = new_finish + timedelta(hours=8, minutes=55)
                                else:
                                    new.lesson_time_start = new_finish + timedelta(hours=8)
                                    new.lesson_time_finish = new_finish + timedelta(hours=8, minutes=45)
                            for curr, next in zip(NextLessons, NextLessons[1:]):
                                curr.lesson_time_start = next.lesson_time_start
                                curr.lesson_time_finish = next.lesson_time_finish
                                curr.save()
                            LastNext = NextLessons.latest('lesson_time_start')
                            if (CountPrevLessons + NextLessons.count()) % 2 == 0:
                                if (CountPrevLessons + NextLessons.count()) % 4 == 0:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=75)
                                else:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=55)
                            else:
                                LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                            LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                            LastNext.save()
                        else:
                            new.lesson_time_start = timestart
                            new.lesson_time_finish = timefinish
                            for curr, next in zip(NextLessons, NextLessons[1:]):
                                curr.lesson_time_start = next.lesson_time_start
                                curr.lesson_time_finish = next.lesson_time_finish
                                curr.save()
                            FirstNext = NextLessons.latest('-lesson_time_start')
                            FirstNext.lesson_time_start = timefinish
                            FirstNext.save()
                            FirstNext.lesson_time_finish = FirstNext.lesson_time_start + timedelta(minutes=45)
                            FirstNext.save()
                            if NextLessons.count() > 1:
                                LastNext = NextLessons.latest('lesson_time_start')
                                if NextLessons.count() % 2 == 0:
                                    if NextLessons.count() % 4 == 0:
                                        LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=75)
                                    else:
                                        LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=55)
                                else:
                                    LastNext.lesson_time_start = LastNext.lesson_time_start + timedelta(minutes=45)
                                LastNext.save()
                                LastNext.lesson_time_finish = LastNext.lesson_time_start + timedelta(minutes=45)
                                LastNext.save()
                    else:
                        new.lesson_time_start = timestart
                        new.lesson_time_finish = timefinish
                else:
                    new.lesson_time_start = timestart
                    new.lesson_time_finish = timefinish
                new.teacher_id = int(str_teacher[str_teacher.find(':')+1:str_teacher.find(')')])
                new.save()
                messages.success(request, 'Занятие успешно добавлено')
            except BaseException:
                messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/schedule/event_lessons_'+str(self.kwargs.get('group')))

    def get_context_data(self, **kwargs):
        context = super(EventLessonsList, self).get_context_data(**kwargs)
        context.update({
            'event': Events.objects.get(id=StudentGroups.objects.get(id=self.kwargs.get('group')).event_id)
        })
        return context


class PersonalSchedule(CheckAdminMixin, ListView):
    login_url = '/'
    model = CourseLessons
    context_object_name = 'courselessons'
    template_name = 'centre/guides/personal_schedule.html'

    def get_queryset(self):
        return CourseLessons.objects.filter(teacher_id=self.kwargs.get('teach')).order_by('lesson_time_start')

    def get_context_data(self, **kwargs):
        context = super(PersonalSchedule, self).get_context_data(**kwargs)
        context.update({
            'eventlessons': EventsLessons.objects.filter(teacher_id=self.kwargs.get('teach')).order_by('lesson_time_start'),
            'teacher': Profiles.objects.get(id=self.kwargs.get('teach'))
        })
        return context


class ReportsView(CheckAdminMixin, View):
    login_url = '/'

    def get(self, request, *args, **kwargs):
        if 'dpp_year' in request.GET:
            document = Document(STATIC_ROOT + "\\doc_templates\\docx\\reports\\dpp.docx")
            tbl = document.tables[0]
            list_id = []
            for course in Courses.objects.filter(date_start__year=request.GET.get('dpp_year')):
                list_id.append(course.program_id)
            for index, dpp in enumerate(Programs.objects.filter(id__in=list_id).prefetch_related('categories').order_by('-id')):
                row_cells = tbl.add_row().cells
                row_cells[0].text = str(index + 1)
                row_cells[1].text = dpp.name
                row_cells[2].text = str(dpp.duration)
                cats = ''
                for cat in dpp.categories.all():
                    cats += cat.name + ';\n'
                row_cells[3].text = cats
                row_cells[4].text = dpp.annotation
                row_cells[5].text = '№' + dpp.order_id + '/' + dpp.order_date.strftime('%d.%m.%Y')
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(11)
            newpath = MEDIA_ROOT + "\\Отчеты\\Перечень ДПП\\"
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            document.save(newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx")
            filename = newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            strin = "ДПП_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            data = open(filename, "rb").read()
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Length'] = os.path.getsize(filename)
            response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
            return response
        elif 'dpp_month' in request.GET:
            document = Document(STATIC_ROOT + "\\doc_templates\\docx\\reports\\dpp.docx")
            tbl = document.tables[0]
            list_id = []
            for course in Courses.objects.filter(date_start__year=datetime.datetime.now().year).\
                    filter(date_start__month=request.GET.get('dpp_month')):
                list_id.append(course.program_id)
            for index, dpp in enumerate(Programs.objects.filter(id__in=list_id).prefetch_related('categories').order_by('-id')):
                row_cells = tbl.add_row().cells
                row_cells[0].text = str(index + 1)
                row_cells[1].text = dpp.name
                row_cells[2].text = str(dpp.duration)
                cats = ''
                for cat in dpp.categories.all():
                    cats += cat.name + ';\n'
                row_cells[3].text = cats
                row_cells[4].text = dpp.annotation
                row_cells[5].text = '№' + dpp.order_id + '/' + dpp.order_date.strftime('%d.%m.%Y')
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(11)
            newpath = MEDIA_ROOT + "\\Отчеты\\Перечень ДПП\\"
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            document.save(newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx")
            filename = newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            strin = "ДПП_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            data = open(filename, "rb").read()
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Length'] = os.path.getsize(filename)
            response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
            return response
        elif 'sch_year' in request.GET:
            groups = StudentGroups.objects.\
                filter(Q(course__in=Courses.objects.filter(date_start__year=request.GET.get('sch_year'))) | Q(event__in=Events.objects.filter(date_start__year=request.GET.get('sch_year')))).\
                select_related('course', 'event')
            document = Document(STATIC_ROOT + "\\doc_templates\\docx\\reports\\schedule.docx")
            tbl = document.tables[0]
            for index, group in enumerate(groups):
                row_cells = tbl.add_row().cells
                row_cells[0].text = str(index + 1)
                if group.event is None:
                    crs = group.course
                    row_cells[1].text = crs.program.department
                    if crs.program.type_dpp == 'Повышение квалификации':
                        row_cells[2].text = 'Курс повышения квалификации'
                    else:
                        row_cells[2].text = 'Курс профессиональной переподготовки'
                    row_cells[3].text = crs.program.name
                    row_cells[4].text = str(crs.program.duration)
                    row_cells[5].text = crs.date_start.strftime('%d.%m.%Y')
                    row_cells[6].text = crs.date_finish.strftime('%d.%m.%Y')
                else:
                    row_cells[1].text = group.event.department
                    row_cells[2].text = group.event.type.name
                    row_cells[3].text = group.event.name
                    row_cells[4].text = str(group.event.duration)
                    row_cells[5].text = group.event.date_start.strftime('%d.%m.%Y')
                    row_cells[6].text = group.event.date_finish.strftime('%d.%m.%Y')
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(11)
            newpath = MEDIA_ROOT + "\\Отчеты\\График оказания платных услуг\\"
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            document.save(newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx")
            filename = newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            strin = "ГрафикУслуг_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            data = open(filename, "rb").read()
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Length'] = os.path.getsize(filename)
            response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
            return response
        elif 'sch_month' in request.GET:
            groups = StudentGroups.objects. \
                filter(Q(course__in=Courses.objects.filter(date_start__month=request.GET.get('sch_month'))) | Q(
                event__in=Events.objects.filter(date_start__month=request.GET.get('sch_month')))). \
                select_related('course', 'event')
            document = Document(STATIC_ROOT + "\\doc_templates\\docx\\reports\\schedule.docx")
            tbl = document.tables[0]
            for index, group in enumerate(groups):
                row_cells = tbl.add_row().cells
                row_cells[0].text = str(index + 1)
                if group.event is None:
                    crs = group.course
                    row_cells[1].text = crs.program.department
                    if crs.program.type_dpp == 'Повышение квалификации':
                        row_cells[2].text = 'Курс повышения квалификации'
                    else:
                        row_cells[2].text = 'Курс профессиональной переподготовки'
                    row_cells[3].text = crs.program.name
                    row_cells[4].text = str(crs.program.duration)
                    row_cells[5].text = crs.date_start.strftime('%d.%m.%Y')
                    row_cells[6].text = crs.date_finish.strftime('%d.%m.%Y')
                else:
                    row_cells[1].text = group.event.department
                    row_cells[2].text = group.event.type.name
                    row_cells[3].text = group.event.name
                    row_cells[4].text = str(group.event.duration)
                    row_cells[5].text = group.event.date_start.strftime('%d.%m.%Y')
                    row_cells[6].text = group.event.date_finish.strftime('%d.%m.%Y')
            for row in tbl.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(11)
            newpath = MEDIA_ROOT + "\\Отчеты\\График оказания платных услуг\\"
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            document.save(newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx")
            filename = newpath + "Отчет_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            strin = "ГрафикУслуг_" + datetime.date.today().strftime('%d.%m.%Y') + ".docx"
            data = open(filename, "rb").read()
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Length'] = os.path.getsize(filename)
            response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
            return response
        elif 'pk-1' in request.GET:
            if not Courses.objects.filter(date_start__year=request.GET.get('pk-1')).exists():
                messages.error(request, 'Не найдено запланированных курсов в указанном году')
                return HttpResponseRedirect('/centre/reports/')
            if Reports.objects.filter(type_report='ПК-1'). \
                    filter(admin_id=Profiles.objects.get(user_id=request.user.id)).exists():
                old_report = Reports.objects.filter(type_report='ПК-1').get(admin_id=Profiles.objects.get(user_id=request.user.id).id)
                if old_report.report is not None and  old_report.report.name != '':
                    os.remove(old_report.report.name)
                old_report.delete()
            new_report = Reports()
            new_report.admin_id = Profiles.objects.get(user_id=request.user.id).id
            new_report.date_start = datetime.datetime.now() + timedelta(hours=8)
            new_report.type_report = 'ПК-1'
            new_report.save()
            id_rep = new_report.id
            prof = Profiles.objects.get(user_id=request.user.id)
            fio = prof.surname + ' ' + prof.name[:1] + '.' + prof.patronymic[:1]
            newpath = STATIC_ROOT + "\\Отчеты\\ПК-1\\" + fio + "\\"
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            generate_pk1.delay(request.GET.get('pk-1'), Profiles.objects.get(user_id=request.user.id).id, id_rep)
            messages.success(request, 'Запрос на формирование отчета отправлен. На вашу электронную почту будет '
                                      'отправлено письмо об успешном завершении формирования отчета')
            return HttpResponseRedirect('/')
        elif 'user_pk1' in request.GET:
            if Reports.objects.filter(type_report='ПК-1').\
                    filter(admin_id=Profiles.objects.get(user_id=request.GET.get('user_pk1')).id).exists():
                last_rep = Reports.objects.filter(type_report='ПК-1').\
                    filter(admin_id=Profiles.objects.get(user_id=request.GET.get('user_pk1')).id).latest('id')
                if last_rep.date_finish is not None:
                    irk_date_start = last_rep.date_start
                    irk_date_finish = last_rep.date_finish + timedelta(hours=8)
                    return JsonResponse({
                        'date_start': irk_date_start.strftime('%d.%m.%Y %H:%M'),
                        'date_finish': irk_date_finish.strftime('%d.%m.%Y %H:%M'),
                        'id_report': last_rep.id
                    })
                else:
                    return JsonResponse({
                        'date_start': last_rep.date_start.strftime('%d.%m.%Y %H:%M'),
                        'date_finish': 'Формируется',
                        'id_report': '-'
                    })
            else:
                return JsonResponse({
                    'no_reports': 'yes'
                })
        elif 'groups_frdo' in request.GET:
            if Courses.objects.filter(Q(date_start__month=request.GET.get('groups_frdo')) | Q(date_finish__month=request.GET.get('groups_frdo'))).exists():
                crses = Courses.objects.filter(Q(date_start__month=request.GET.get('groups_frdo')) | Q(date_finish__month=request.GET.get('groups_frdo')))
                if StudentGroups.objects.filter(event_id=None).filter(course__in=crses).exists():
                    stgrs = StudentGroups.objects.filter(event_id=None).filter(course__in=crses)
                    dict_grs = {}
                    for gr in stgrs:
                        l = []
                        l.append(gr.code)
                        l.append(gr.course.program.name)
                        l.append(gr.course.date_start.strftime('%d.%m.%Y') + ' - ' + gr.course.date_finish.strftime('%d.%m.%Y'))
                        dict_grs[gr.id] = l
                    return JsonResponse({
                        'groups': dict_grs
                    })
                else:
                    return JsonResponse({
                        'no_groups': 'yes'
                    })
            else:
                return JsonResponse({
                    'no_groups': 'yes'
                })
        elif 'get_report' in request.GET:
            rec_report = Reports.objects.get(id=request.GET.get('get_report'))
            path = os.path.join(STATIC_ROOT, rec_report.report.name)
            correct_path = path.replace("/", "")
            data = open(correct_path, "rb").read()
            if rec_report.type_report == 'ПК-1':
                filename = "Отчет-ПК1.xlsx"
            else:
                filename = 'Отчет ФИС ФРДО.xlsx'
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Length'] = os.path.getsize(correct_path)
            response['Content-Disposition'] = 'attachment; filename="' + escape_uri_path(filename) + '"'
            return response
        elif 'fis_frdo' in request.GET:
            if Reports.objects.filter(type_report='ФИС ФРДО').\
                    filter(admin_id=Profiles.objects.get(user_id=request.GET.get('fis_frdo')).id).exists():
                last_rep = Reports.objects.filter(type_report='ФИС ФРДО').\
                    filter(admin_id=Profiles.objects.get(user_id=request.GET.get('fis_frdo')).id).latest('id')
                if last_rep.date_finish is not None:
                    irk_date_start = last_rep.date_start
                    irk_date_finish = last_rep.date_finish + timedelta(hours=8)
                    return JsonResponse({
                        'date_start': irk_date_start.strftime('%d.%m.%Y %H:%M'),
                        'date_finish': irk_date_finish.strftime('%d.%m.%Y %H:%M'),
                        'id_report': last_rep.id
                    })
                else:
                    return JsonResponse({
                        'date_start': last_rep.date_start.strftime('%d.%m.%Y %H:%M'),
                        'date_finish': 'Формируется',
                        'id_report': '-'
                    })
            else:
                return JsonResponse({
                    'no_reports': 'yes'
                })
        else:
            return render(request, 'centre/reports/main.html')

    def post(self, request, *args, **kwargs):
        if 'get_frdo' in request.POST:
            list_groups = request.POST.getlist('ChoosenGroups')
            if len(list_groups) == 0:
                messages.error(request, 'Группы не выбраны')
                return HttpResponseRedirect('/')
            if Reports.objects.filter(type_report='ФИС ФРДО'). \
                    filter(admin_id=Profiles.objects.get(user_id=request.user.id)).exists():
                old_report = Reports.objects.filter(type_report='ФИС ФРДО').get(
                    admin_id=Profiles.objects.get(user_id=request.user.id).id)
                if old_report.report is not None and old_report.report.name != '':
                    os.remove(old_report.report.name)
                old_report.delete()
            new_report = Reports()
            new_report.admin_id = Profiles.objects.get(user_id=request.user.id).id
            new_report.date_start = datetime.datetime.now() + timedelta(hours=8)
            new_report.type_report = 'ФИС ФРДО'
            new_report.save()
            id_rep = new_report.id
            prof = Profiles.objects.get(user_id=request.user.id)
            fio = prof.surname + ' ' + prof.name[:1] + '.' + prof.patronymic[:1]
            newpath = STATIC_ROOT + "\\Отчеты\\ФИС ФРДО\\" + fio + "\\"
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            frdo_report.delay(list_groups, Profiles.objects.get(user_id=request.user.id).id, id_rep)
            messages.success(request, 'Запрос на формирование отчета отправлен. На вашу электронную почту будет '
                                      'отправлено письмо об успешном завершении формирования отчета')
            return HttpResponseRedirect('/')
        else:
            return HttpResponseRedirect('/access_denied/')


def get_age_range(age):
    if age < 25:
        range = 'Моложе 25 лет'
    elif 25 <= age <= 29:
        range = '25-29'
    elif 30 <= age <= 34:
        range = '30-34'
    elif 35 <= age <= 39:
        range = '35-39'
    elif 40 <= age <= 44:
        range = '40-44'
    elif 45 <= age <= 49:
        range = '45-49'
    elif 50 <= age <= 54:
        range = '50-54'
    elif 55 <= age <= 59:
        range = '55-59'
    elif 60 <= age <= 64:
        range = '60-64'
    else:
        range = '65 и более'
    return range


class student_course():
    def __init__(self, id, surname, name,
                 patronymic, region, mo, oo, pos_cat,
                 pos, age, birthday, sex,
                 snils, edu_level, dipsurname,
                 edu_serial, edu_number, type_pay,
                 reg_number, serial_blank, number_blank):
        self.id = id
        self.surname = surname
        self.name = name
        self.patronymic = patronymic
        self.region = region
        self.mo = mo
        self.oo = oo
        self.pos_cat = pos_cat
        self.pos = pos
        self.age = age
        self.birthday = birthday
        self.sex = sex
        self.snils = snils
        self.edu_level = edu_level
        self.dipsurname = dipsurname
        self.edu_serial = edu_serial
        self.edu_number = edu_number
        self.type_pay = type_pay
        self.reg_number = reg_number
        self.serial_blank = serial_blank
        self.number_blank = number_blank


class student_event():
    def __init__(self, id, surname, name, patronymic,
                 terr, oo, pos):
        self.id = id
        self.surname = surname
        self.name = name
        self.patronymic = patronymic
        self.terr =terr
        self.oo = oo
        self.pos = pos


def get_students(group_id):
    sts = []
    studs = StudentGroups.objects.get(id=group_id).students.all().order_by('surname')
    count = 1
    if StudentGroups.objects.get(id=group_id).event is None:
        for stud in studs:
            form = CoursesForms.objects.filter(group_id=group_id).get(profile_id=stud.id)
            region = form.region.name
            if form.mo is None:
                mo = ''
            else:
                mo = form.mo.name
            if form.oo is None:
                oo = form.oo_new
            else:
                oo = form.oo.full_name
            if form.position_cat is None:
                pos_cat = ''
            else:
                pos_cat = form.position_cat.name
            if form.position is None:
                pos = ''
            else:
                pos = form.position.name
            age = datetime.date.today().year - stud.birthday.year
            if stud.sex is True:
                sex = 'Мужской'
            else:
                sex = 'Женский'
            if form.edu_level.name == 'Студент':
                dipsurname = ''
                edu_serial = ''
                edu_number = ''
            else:
                dipsurname = form.check_surname
                edu_serial = form.edu_serial
                edu_number = form.edu_number
            if form.type is True:
                pay = 'Физ. лицо'
            else:
                pay = 'Юр. лицо'
            if StudentsCerts.objects.filter(group_id=group_id).filter(student_id=stud.id).exists():
                cert = StudentsCerts.objects.filter(group_id=group_id).get(student_id=stud.id)
                st = student_course(
                    str(count),
                    stud.surname,
                    stud.name,
                    stud.patronymic,
                    region,
                    mo,
                    oo,
                    pos_cat,
                    pos,
                    get_age_range(age),
                    stud.birthday.strftime('%d.%m.%Y'),
                    sex,
                    stud.snils,
                    form.edu_level.name,
                    dipsurname,
                    edu_serial,
                    edu_number,
                    pay,
                    cert.reg_number,
                    cert.blank_serial,
                    cert.blank_number
                )
            else:
                st = student_course(
                    str(count),
                    stud.surname,
                    stud.name,
                    stud.patronymic,
                    region,
                    mo,
                    oo,
                    pos_cat,
                    pos,
                    get_age_range(age),
                    stud.birthday.strftime('%d.%m.%Y'),
                    sex,
                    stud.snils,
                    form.edu_level.name,
                    dipsurname,
                    edu_serial,
                    edu_number,
                    pay,
                    '',
                    '',
                    ''
                )
            count += 1
            sts.append(st)
    else:
        for stud in studs:
            form = EventsForms.objects.filter(group_id=group_id).get(profile_id=stud.id)
            if form.region.name == 'Иркутская область':
                terr = "Иркутская область ("+form.mo.name+")"
            else:
                terr = form.region.name
            if form.oo_new is None:
                oo = form.oo.full_name
            elif form.oo is None:
                oo = '-'
            else:
                oo = form.oo_new
            if form.position is None:
                pos = '-'
            else:
                pos = form.position.name
            item = student_event(
                str(count),
                stud.surname,
                stud.name,
                stud.patronymic,
                terr,
                oo,
                pos
            )
            sts.append(item)
            count += 1
    return sts


class teacher_course():
    def __init__(self, id, fio, lecture, practice, trainee,
                 individual, total):
        self.id = id
        self.fio = fio
        self.lecture = lecture
        self.practice = practice
        self.trainee = trainee
        self.individual = individual
        self.total = total


class teacher_event():
    def __init__(self, id, fio, lecture,
                 practice, total):
        self.id = id
        self.fio = fio
        self.lecture = lecture
        self.practice = practice
        self.total = total


def get_teachers(group_id):
    tchs = []
    gr = StudentGroups.objects.get(id=group_id)
    if gr.event is None:
        lessons = CourseLessons.objects.filter(group_id=group_id)
    else:
        lessons = EventsLessons.objects.filter(group_id=group_id)
    teachers = lessons.values('teacher_id').distinct()
    count = 1
    for teacher in teachers:
        prof = Profiles.objects.get(id=teacher['teacher_id'])
        teacher_lessons = lessons.filter(teacher_id=teacher['teacher_id'])
        lecture = practice = individual = trainee = total = 0
        for el in teacher_lessons:
            lecture += el.lecture_hours
            practice += el.practice_hours
            if gr.event is None:
                individual += el.individual_hours
                trainee += el.trainee_hours
                total += el.lecture_hours + el.practice_hours + el.individual_hours + el.trainee_hours
            else:
                total += el.lecture_hours + el.practice_hours
        if gr.event is None:
            t = teacher_course(
                str(count),
                prof.surname+' '+prof.name+' '+prof.patronymic,
                str(lecture),
                str(practice),
                str(trainee),
                str(individual),
                str(total)
            )
        else:
            t = teacher_event(
                str(count),
                prof.surname + ' ' + prof.name + ' ' + prof.patronymic,
                str(lecture),
                str(practice),
                str(total)
            )
        tchs.append(t)
        count += 1
    return tchs


class StudentsGroupList(CheckAdminMixin, ListView):
    login_url = '/'
    model = StudentGroups
    context_object_name = 'groups'
    template_name = 'centre/study/studentgroups/list.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'change_certreg' in request.GET:
                stgr = StudentGroups.objects.get(id=request.GET.get('group'))
                if StudentsCerts.objects.filter(group_id=stgr.id).filter(
                        student_id=request.GET.get('change_certreg')).exists():
                    rec = StudentsCerts.objects.filter(group_id=stgr.id).get(
                        student_id=request.GET.get('change_certreg'))
                else:
                    rec = StudentsCerts()
                    rec.group_id = stgr.id
                    rec.student_id = request.GET.get('change_certreg')
                rec.reg_number = request.GET.get('reg')
                rec.save()
                return JsonResponse({})
            if 'change_certser' in request.GET:
                stgr = StudentGroups.objects.get(id=request.GET.get('group'))
                if StudentsCerts.objects.filter(group_id=stgr.id).filter(
                        student_id=request.GET.get('change_certser')).exists():
                    rec = StudentsCerts.objects.filter(group_id=stgr.id).get(
                        student_id=request.GET.get('change_certser'))
                else:
                    rec = StudentsCerts()
                    rec.group_id = stgr.id
                    rec.student_id = request.GET.get('change_certser')
                rec.blank_serial = request.GET.get('ser')
                rec.save()
                return JsonResponse({})
            if 'change_certnumb' in request.GET:
                stgr = StudentGroups.objects.get(id=request.GET.get('group'))
                if StudentsCerts.objects.filter(group_id=stgr.id).filter(
                        student_id=request.GET.get('change_certnumb')).exists():
                    rec = StudentsCerts.objects.filter(group_id=stgr.id).get(
                        student_id=request.GET.get('change_certnumb'))
                else:
                    rec = StudentsCerts()
                    rec.group_id = stgr.id
                    rec.student_id = request.GET.get('change_certnumb')
                rec.blank_number = request.GET.get('numb')
                rec.save()
                return JsonResponse({})
            if 'generate_certs' in request.GET:
                reg_number = int(request.GET.get('reg'))
                ser_blank = request.GET.get('ser')
                numb_blank = int(request.GET.get('numb'))
                stgr = StudentGroups.objects.prefetch_related('students').get(id=request.GET.get('generate_certs'))
                for stud in stgr.students.all().order_by('surname'):
                    if StudentsCerts.objects.filter(group_id=stgr.id).filter(student_id=stud.id).exists():
                        StudentsCerts.objects.filter(group_id=stgr.id).get(student_id=stud.id).delete()
                    new = StudentsCerts()
                    new.group_id = stgr.id
                    new.student_id = stud.id
                    new_reg = "%04d" % reg_number
                    new.reg_number = new_reg
                    new.blank_serial = ser_blank
                    new.blank_number = str(numb_blank)
                    reg_number += 1
                    numb_blank += 1
                    new.save()
                messages.success(request, 'Данные успешно сгенерированы')
                return JsonResponse({})
            if 'enroll_exp' in request.GET:
                gr = StudentGroups.objects.get(id=request.GET.get('enroll_exp'))
                if gr.date_enroll is None:
                    date_enr = None
                else:
                    date_enr = gr.date_enroll.strftime('%Y-%m-%d')
                if gr.date_exp is None:
                    date_ex = None
                else:
                    date_ex = gr.date_exp.strftime('%Y-%m-%d')
                return JsonResponse({
                    'enroll_number': gr.enroll_number,
                    'exp_number': gr.exp_number,
                    'date_enroll': date_enr,
                    'date_exp': date_ex,
                })
            teacher = Profiles.objects.get(id=request.GET.get('id_teach'))
            return JsonResponse({
                'fio': teacher.surname+' '+teacher.name+' '+teacher.patronymic,
                'email': User.objects.get(id=teacher.user_id).email,
                'phone': teacher.phone
            })
        if 'cert_list' in request.GET:
            group = StudentGroups.objects.select_related('course').get(id=request.GET.get('cert_list'))
            dep = group.course.program.department
            date_start = group.course.date_start
            date_finish = group.course.date_finish
            date_enroll = datetime.date(int(request.GET.get('enroll_date')[:4]),
                                        int(request.GET.get('enroll_date')[-5:7]),
                                        int(request.GET.get('enroll_date')[8:]))
            date_expl = datetime.date(int(request.GET.get('expl_date')[:4]),
                                      int(request.GET.get('expl_date')[-5:7]),
                                      int(request.GET.get('expl_date')[8:]))
            group.date_enroll = date_enroll
            group.date_exp = date_expl
            group.enroll_number = request.GET.get('enroll_number')
            group.exp_number = request.GET.get('expl_number')
            group.save()
            if group.course.program.type_dpp == 'Повышение квалификации':
                type_dpp = 'повышения квалификации'
            else:
                type_dpp = 'профессиональной переподготовки'
            deps = GetDepsWithManagerFromAD()
            for key, value in deps.items():
                if key == dep:
                    manager = value
            if group.course.program.type_dpp == 'Повышение квалификации':
                type_dpp = 'повышении квалификации'
            else:
                type_dpp = 'профессиональной переподготовке'
            path = STATIC_ROOT + '\\doc_templates\\xlsx\\cert_list.xlsx'
            writer = BookWriter(path)
            info = {
                'dep': dep,
                'type_dpp': type_dpp,
                'name_dpp': group.course.program.name,
                'duration': group.course.program.duration,
                'code': group.code,
                'manager': manager,
                'day_start': date_start.strftime('%d'),
                'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                'year_start': date_start.strftime('%Y'),
                'day_finish': date_finish.strftime('%d'),
                'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                'year_finish': date_finish.strftime('%Y'),
                'date_get': group.course.date_finish.strftime('%d.%m.%Y'),
                'date_enroll': date_enroll.strftime('%d.%m.%Y'),
                'date_expl': date_expl.strftime('%d.%m.%Y'),
                'number_enroll': request.GET.get('enroll_number'),
                'number_expl': request.GET.get('expl_number'),
            }
            info['students'] = get_students(group.id)
            writer.render_sheet(info, 'Ведомость выдачи удостоверений', 0)
            newpath = MEDIA_ROOT + '\\Ведомости\\' + group.code
            if not os.path.exists(newpath):
                os.makedirs(newpath)
            writer.save(newpath + "\\ВыдачаУдостоверений_.xlsx")
            filename = newpath + "\\ВыдачаУдостоверений_.xlsx"
            strin = "ВыдачаУдостоверений_" + group.code + ".xlsx"
            data = open(filename, "rb").read()
            response = HttpResponse(data,
                                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Length'] = os.path.getsize(filename)
            response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
            return response
        qs = StudentGroups.objects.all().order_by('-id')
        if 'code' in request.GET:
            if len(request.GET.get('code')) > 0:
                qs = qs.filter(code__contains=request.GET.get('code'))
        if 'name' in request.GET:
            if len(request.GET.get('name')) > 0:
                events = Events.objects.filter(name__contains=request.GET.get('name'))
                courses = Courses.objects.filter(program__in=Programs.objects.filter(name__contains=request.GET.get('name')))
                qs = qs.filter(Q(event__in=events) | Q(course__in=courses))
        if 'type' in request.GET:
            typ = uri_to_iri(request.GET.get('type'))
            if typ in ['Повышение квалификации', 'Профессиональная переподготовка']:
                qs = qs.filter(event_id=None).filter(course__in=Courses.objects.filter(program__in=Programs.objects.filter(type_dpp=request.GET.get('type'))))
            elif request.GET.get('type') == 'all':
                pass
            else:
                qs = qs.filter(event__in=Events.objects.filter(type_id=request.GET.get('type')))
        if 'grstatuses' in request.GET:
            qs = qs.filter(status_id__in=request.GET.getlist('grstatuses'))
            return render(request, 'centre/study/studentgroups/list.html', context={
                'groups': qs[:25],
                'filter_status': 'yes'
            })
        return render(request, 'centre/study/studentgroups/list.html', context={
            'groups': qs[:25],
        })

    def post(self, request, *args, **kwargs):
        #try:
            if is_ajax(request=request):
                group = StudentGroups.objects.get(id=request.POST.get('id_group'))
                if 'url_study' in request.POST:
                    group.event_url = request.POST.get('url_study')
                    group.status = StGroupStatuses.objects.get(name='Идет обучение')
                    group.save()
                    messages.success(request, 'Ссылка успешно сохранена')
                    return JsonResponse({})
                elif 'url_survey' in request.POST:
                    group.survey_url = request.POST.get('url_survey')
                    group.save()
                    messages.success(request, 'Ссылка на опрос успешно изменена')
                    return JsonResponse({})
                else:
                    group.offer = request.FILES.get('file')
                    group.status = StGroupStatuses.objects.get(name='Ожидает ссылку на обучение')
                    group.save()
                    students = group.students.all()
                    emails = []
                    for student in students:
                        app = Apps.objects.filter(group_id=request.POST.get('id_group')).get(profile_id=student.id)
                        if app.status_id != Statuses.objects.get(name='Ждем оплату').id:
                            app.status_id = Statuses.objects.get(name='Ждем оплату').id
                            app.save()
                            emails.append(student.user.email)
                    if group.event is None:
                        if group.course.program.type_dpp == 'Повышение квалификации':
                            typ = 'курсе повышения квалификации'
                        else:
                            typ = 'курсе профессиональной переподготовки'
                        name = group.course.program.name
                        start = group.course.date_start
                    else:
                        type_classic = group.event.type.name
                        morph = pymorphy2.MorphAnalyzer()
                        type = morph.parse(type_classic)[0]
                        word = type.inflect({'loct'})
                        typ = word.word
                        name = group.event.name
                        start = group.event.date_start
                    pay_date = start - BDay(5)
                    pay_deadline = pay_date.strftime('%d')+' '+month_from_ru_to_eng(pay_date.strftime('%B'))+' '+pay_date.strftime('%Y')
                    EmailOfferPay.delay(emails, typ, name, pay_deadline)
                    messages.success(request, 'Скан договора оферты успешно загружен!')
                    return JsonResponse({})
            else:
                if 'sz_ou' in request.POST:
                    data = GetDepsWithManagerFromAD()
                    group = StudentGroups.objects.get(id=int(request.POST.get('sz_ou')))
                    department = group.course.program.department
                    for key, value in data.items():
                        if key == department:
                            manager = value
                            dep = key
                            break
                    p = Petrovich()
                    spl = manager.split(' ')
                    fio = spl[0]+' '+spl[1][:1]+'.'+spl[2][:1]+'.'
                    surname = p.lastname(spl[0], Case.GENITIVE)
                    initials = spl[1][:1]+'.'+spl[2][:1]+'.'
                    list_dep = dep.split(' ')
                    check_first = True
                    if 'центр' in list_dep:
                        if list_dep.index('центр') != 0:
                            check_first = False
                    if 'Центр' in list_dep:
                        if list_dep.index('Центр') != 0:
                            check_first = False
                    if check_first is False:
                        for el in list_dep:
                            if el != 'центр' and el != 'Центр':
                                strin = el
                                strin.replace('ый', 'ого')
                                list_dep[list_dep.index(el)] = str
                            else:
                                break
                    position = 'руководителя '
                    position_footer = 'Руководитель '
                    for el in list_dep:
                        if el != 'центр' and el != 'Центр':
                            position += el.lower()+' '
                            position_footer += el.lower()+' '
                        else:
                            position += el.lower()+'a '
                            position_footer += el.lower()+'a '
                    list_dep = dep.split(' ')
                    check_first = True
                    if 'центр' in list_dep:
                        if list_dep.index('центр') != 0:
                            check_first = False
                    if 'Центр' in list_dep:
                        if list_dep.index('Центр') != 0:
                            check_first = False
                    if check_first is False:
                        for el in list_dep:
                            if el != 'центр' and el != 'Центр':
                                strin = el
                                strin.replace('ый', 'ым')
                                list_dep[list_dep.index(el)] = str
                            else:
                                break
                    dep_str = ''
                    for el in list_dep:
                        if el != 'центр' and el != 'Центр':
                            dep_str += el.lower() + ' '
                        else:
                            dep_str += el.lower() + 'ом '
                    sz = group.course.date_start - timedelta(days=4)
                    date_sz = str(sz.day)+' '+month_from_ru_to_eng(sz.strftime('%B'))+' '+str(sz.year)+' г.'
                    dpp_name = group.course.program.name
                    date_st = group.course.date_start
                    date_f = group.course.date_finish
                    date_start = date_st.strftime('%d.%m.%Y')+' г.'
                    date_finish = date_f.strftime('%d.%m.%Y')+' г.'
                    duration = group.course.program.duration
                    if group.students_number is not None:
                        number_students = group.students_number
                    else:
                        number_students = group.students.count()
                    curator = ''
                    if group.curator_id is not None:
                        prof = Profiles.objects.get(id=group.curator_id)
                        if prof.sex is True:
                            curator += p.lastname(prof.surname, Case.ACCUSATIVE, Gender.MALE)+' '+p.firstname(prof.name, Case.ACCUSATIVE, Gender.MALE)+' '+p.middlename(prof.patronymic, Case.ACCUSATIVE, Gender.MALE)
                        else:
                            curator += p.lastname(prof.surname, Case.ACCUSATIVE, Gender.FEMALE) + ' ' + p.firstname(prof.name,
                                                                                                                  Case.ACCUSATIVE,
                                                                                                                  Gender.FEMALE) + ' ' + p.middlename(
                                prof.patronymic, Case.ACCUSATIVE, Gender.FEMALE)
                    doc = DocxTemplate(STATIC_ROOT+"/doc_templates/docx/ou/sz_ou.docx")
                    context = {
                        'manager': surname+' '+initials,
                        'position': position,
                        'date_sz': date_sz,
                        'dep': dep_str,
                        'dpp_name': dpp_name,
                        'date_start': date_start,
                        'date_finish': date_finish,
                        'duration': str(duration),
                        'number_students': str(number_students),
                        'curator': curator,
                        'position_footer': position_footer,
                        'fio': fio
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\СЗ\\ОУ\\" + group.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT+"\\СЗ\\ОУ\\"+group.code+"\\СЗ.docx")
                    filename = MEDIA_ROOT+"\\СЗ\\ОУ\\"+group.code+"\\СЗ.docx"
                    strin = "СЗ_"+group.code+".docx"
                    data = open(filename, "rb").read()
                    response = HttpResponse(data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'sz_iku' in request.POST:
                    data = GetDepsWithManagerFromAD()
                    group = StudentGroups.objects.get(id=int(request.POST.get('sz_iku')))
                    department = group.event.department
                    for key, value in data.items():
                        if key == department:
                            manager = value
                            dep = key
                            break
                    p = Petrovich()
                    spl = manager.split(' ')
                    fio = spl[0]+' '+spl[1][:1]+'.'+spl[2][:1]+'.'
                    surname = p.lastname(spl[0], Case.GENITIVE)
                    initials = spl[1][:1]+'.'+spl[2][:1]+'.'
                    list_dep = dep.split(' ')
                    check_first = True
                    if 'центр' in list_dep:
                        if list_dep.index('центр') != 0:
                            check_first = False
                    if 'Центр' in list_dep:
                        if list_dep.index('Центр') != 0:
                            check_first = False
                    if check_first is False:
                        for el in list_dep:
                            if el != 'центр' and el != 'Центр':
                                strin = el
                                strin.replace('ый', 'ого')
                                list_dep[list_dep.index(el)] = str
                            else:
                                break
                    position = 'руководителя '
                    position_footer = 'Руководитель '
                    for el in list_dep:
                        if el != 'центр' and el != 'Центр':
                            position += el.lower()+' '
                            position_footer += el.lower()+' '
                        else:
                            position += el.lower()+'a '
                            position_footer += el.lower()+'a '
                    list_dep = dep.split(' ')
                    check_first = True
                    if 'центр' in list_dep:
                        if list_dep.index('центр') != 0:
                            check_first = False
                    if 'Центр' in list_dep:
                        if list_dep.index('Центр') != 0:
                            check_first = False
                    if check_first is False:
                        for el in list_dep:
                            if el != 'центр' and el != 'Центр':
                                strin = el
                                strin.replace('ый', 'ым')
                                list_dep[list_dep.index(el)] = str
                            else:
                                break
                    dep_str = ''
                    for el in list_dep:
                        if el != 'центр' and el != 'Центр':
                            dep_str += el.lower() + ' '
                        else:
                            dep_str += el.lower() + 'ом '
                    sz = group.event.date_start - timedelta(days=4)
                    date_sz = str(sz.day)+' '+month_from_ru_to_eng(sz.strftime('%B'))+' '+str(sz.year)+' г.'
                    event_name = group.event.name
                    date_st = group.event.date_start
                    date_f = group.event.date_finish
                    date_start = date_st.strftime('%d.%m.%Y')+' г.'
                    date_finish = date_f.strftime('%d.%m.%Y')+' г.'
                    duration = group.event.duration
                    number_students = group.students.count()
                    curator = ''
                    if group.curator_id is not None:
                        prof = Profiles.objects.get(id=group.curator_id)
                        if prof.sex is True:
                            curator += p.lastname(prof.surname, Case.ACCUSATIVE, Gender.MALE)+' '+p.firstname(prof.name, Case.ACCUSATIVE, Gender.MALE)+' '+p.middlename(prof.patronymic, Case.ACCUSATIVE, Gender.MALE)
                        else:
                            curator += p.lastname(prof.surname, Case.ACCUSATIVE, Gender.FEMALE) + ' ' + p.firstname(prof.name,
                                                                                                                  Case.ACCUSATIVE,
                                                                                                                  Gender.FEMALE) + ' ' + p.middlename(
                                prof.patronymic, Case.ACCUSATIVE, Gender.FEMALE)
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse('бутявка')[0]
                    gent = type.inflect({'gent'})
                    doc = DocxTemplate(STATIC_ROOT+"/doc_templates/docx/iku/sz_iku.docx")
                    context = {
                        'manager': surname+' '+initials,
                        'position': position,
                        'date_sz': date_sz,
                        'dep': dep_str,
                        'event_name': event_name,
                        'date_start': date_start,
                        'date_finish': date_finish,
                        'duration': str(duration),
                        'number_students': str(number_students),
                        'curator': curator,
                        'position_footer': position_footer,
                        'fio': fio,
                        'type': gent.word
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\СЗ\\ИКУ\\"+group.code+"\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT+"\\СЗ\\ИКУ\\"+group.code+"\\СЗ.docx")
                    filename = MEDIA_ROOT+"\\СЗ\\ИКУ\\"+group.code+"\\СЗ.docx"
                    strin = "СЗ_"+group.code+".docx"
                    data = open(filename, "rb").read()
                    response = HttpResponse(data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'order_iku' in request.POST:
                    data = GetDepsWithManagerFromAD()
                    group = StudentGroups.objects.get(id=int(request.POST.get('order_iku')))
                    department = group.event.department
                    for key, value in data.items():
                        if key == department:
                            manager = value
                            dep = key
                            break
                    spl = manager.split(' ')
                    fio = spl[0]+' '+spl[1][:1]+'.'+spl[2][:1]+'.'
                    list_dep = dep.split(' ')
                    dep = ''
                    for word in list_dep:
                        if word not in ['Центр', 'центр']:
                            dep += word+' '
                    code = group.code
                    event_name = group.event.name
                    date_st = group.event.date_start
                    date_fn = group.event.date_finish
                    date_start = date_st.strftime('%d.%m.%Y')
                    date_finish = date_fn.strftime('%d.%m.%Y')
                    duration = group.event.duration
                    type_classic = group.event.type.name
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(type_classic)[0]
                    gent = type.inflect({'gent'})
                    loct = type.inflect({'loct'})
                    price = group.event.price
                    doc = DocxTemplate(STATIC_ROOT+"/doc_templates/docx/iku/order_iku.docx")
                    context = {
                        'type': gent.word,
                        'dep': dep,
                        'manager': fio,
                        'code': code,
                        'event_name': event_name,
                        'date_start': date_start,
                        'date_finish': date_finish,
                        'duration': str(duration),
                        'type_classic': type_classic.lower(),
                        'type_loct': loct.word,
                        'price': str(price),
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\Приказы\\ИКУ\\"+group.code+"\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT+"\\Приказы\\ИКУ\\"+group.code+"\\Приказ.docx")
                    filename = MEDIA_ROOT+"\\Приказы\\ИКУ\\"+group.code+"\\Приказ.docx"
                    strin = "Приказ_"+group.code+".docx"
                    data = open(filename, "rb").read()
                    response = HttpResponse(data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'order_ou' in request.POST:
                    data = GetDepsWithManagerFromAD()
                    group = StudentGroups.objects.get(id=int(request.POST.get('order_ou')))
                    department = group.course.program.department
                    for key, value in data.items():
                        if key == department:
                            manager = value
                            dep = key
                            break
                    spl = manager.split(' ')
                    fio = spl[0]+' '+spl[1][:1]+'.'+spl[2][:1]+'.'
                    list_dep = dep.split(' ')
                    dep = ''
                    for word in list_dep:
                        if word not in ['Центр', 'центр']:
                            dep += word+' '
                    code = group.code
                    program = group.course.program.name
                    duration = group.course.program.duration
                    date_st = group.course.date_start
                    day_start = date_st.day
                    month_start = month_from_ru_to_eng(date_st.strftime('%B'))
                    date_fn = group.course.date_finish
                    day_finish = date_fn.day
                    month_finish = month_from_ru_to_eng(date_fn.strftime('%B'))
                    year = date_fn.year
                    price = group.course.program.price
                    doc = DocxTemplate(STATIC_ROOT+"/doc_templates/docx/ou/order_ou.docx")
                    context = {
                        'dep': dep,
                        'manager': fio,
                        'code': code,
                        'program': program,
                        'duration': duration,
                        'day_start': day_start,
                        'day_finish': day_finish,
                        'month_start': month_start,
                        'month_finish': month_finish,
                        'year': year,
                        'price': str(price),
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT+"\\Приказы\\ОУ\\"+group.code+"\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT+"\\Приказы\\ОУ\\"+group.code+"\\Приказ.docx")
                    filename = MEDIA_ROOT+"\\Приказы\\ОУ\\"+group.code+"\\Приказ.docx"
                    strin = 'Приказ_'+group.code+'.docx'
                    data = open(filename, "rb").read()
                    response = HttpResponse(data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'offer_ou' in request.POST:
                    group = StudentGroups.objects.get(id=int(request.POST.get('offer_ou')))
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        type_dpp = 'повышения квалификации'
                        cert_type = 'повышении квалификации'
                    else:
                        type_dpp = 'профессиональной переподготовки'
                        cert_type = 'профессиональной переподготовке'
                    start = group.course.date_start
                    finish = group.course.date_finish
                    order_date = start - BDay(10)
                    pay_date = start - BDay(5)
                    doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/offer_ou.docx")
                    if group.curator is not None:
                        curator_email = group.curator.user.email
                    else:
                        curator_email = ''
                    context = {
                        'code': group.code,
                        'type_dpp': type_dpp,
                        'day_order': order_date.strftime('%d'),
                        'month_order': month_from_ru_to_eng(order_date.strftime('%B')),
                        'year_order': order_date.strftime('%Y'),
                        'dpp_name': group.course.program.name,
                        'duration': group.course.program.duration,
                        'day_start': start.strftime('%d'),
                        'month_start': month_from_ru_to_eng(start.strftime('%B')),
                        'year_start': start.strftime('%Y'),
                        'day_finish': finish.strftime('%d'),
                        'month_finish': month_from_ru_to_eng(finish.strftime('%B')),
                        'year_finish': finish.strftime('%Y'),
                        'cert_type': cert_type,
                        'curator_email': curator_email,
                        'price': str(group.course.program.price),
                        'price_letter': num2words(group.course.program.price, lang='ru'),
                        'day_pay': pay_date.strftime('%d'),
                        'month_pay': month_from_ru_to_eng(pay_date.strftime('%B')),
                        'year_pay': pay_date.strftime('%Y'),
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\Договора\\ОУ\\" + group.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT + "\\Договора\\ОУ\\" + group.code + "\\ДоговорОферты.docx")
                    filename = MEDIA_ROOT + "\\Договора\\ОУ\\" + group.code + "\\ДоговорОферты.docx"
                    strin = 'Договор_' + group.code + '.docx'
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'offer_iku' in request.POST:
                    group = StudentGroups.objects.get(id=int(request.POST.get('offer_iku')))
                    type_dpp = group.event.type.name
                    cert_type = group.event.type.name
                    start = group.event.date_start
                    finish = group.event.date_finish
                    order_date = start - BDay(10)
                    pay_date = start - BDay(5)
                    doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/offer_ou.docx")
                    if group.curator is not None:
                        curator_email = group.curator.user.email
                    else:
                        curator_email = ''
                    context = {
                        'code': group.code,
                        'type_dpp': type_dpp,
                        'day_order': order_date.strftime('%d'),
                        'month_order': month_from_ru_to_eng(order_date.strftime('%B')),
                        'year_order': order_date.strftime('%Y'),
                        'dpp_name': group.event.name,
                        'duration': group.event.duration,
                        'day_start': start.strftime('%d'),
                        'month_start': month_from_ru_to_eng(start.strftime('%B')),
                        'year_start': start.strftime('%Y'),
                        'day_finish': finish.strftime('%d'),
                        'month_finish': month_from_ru_to_eng(finish.strftime('%B')),
                        'year_finish': finish.strftime('%Y'),
                        'cert_type': cert_type,
                        'curator_email': curator_email,
                        'price': str(group.event.price),
                        'price_letter': num2words(group.event.price, lang='ru'),
                        'day_pay': pay_date.strftime('%d'),
                        'month_pay': month_from_ru_to_eng(pay_date.strftime('%B')),
                        'year_pay': pay_date.strftime('%Y'),
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\Договора\\ИКУ\\" + group.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT + "\\Договора\\ИКУ\\" + group.code + "\\ДоговорОферты.docx")
                    filename = MEDIA_ROOT + "\\Договора\\ИКУ\\" + group.code + "\\ДоговорОферты.docx"
                    strin = 'Договор_' + group.code + '.docx'
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'courses_forms' in request.POST:
                    group = StudentGroups.objects.get(id=int(request.POST.get('courses_forms')))
                    strin = 'Анкеты_' + group.code + '.xlsx'
                    response = HttpResponse(
                        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    )
                    response['Content-Disposition'] = 'attachment; filename="'+escape_uri_path(strin)+'"'
                    workbook = Workbook()
                    worksheet = workbook.active
                    worksheet.title = 'Анкеты'
                    cell = worksheet.cell(row=1, column=1)
                    cell.value = 'Телефон'
                    cell = worksheet.cell(row=1, column=2)
                    cell.value = 'Email'
                    cell = worksheet.cell(row=1, column=3)
                    cell.value = 'Фамилия'
                    cell = worksheet.cell(row=1, column=4)
                    cell.value = 'Имя'
                    cell = worksheet.cell(row=1, column=5)
                    cell.value = 'Отчество'
                    cell = worksheet.cell(row=1, column=6)
                    cell.value = 'Регион'
                    cell = worksheet.cell(row=1, column=7)
                    cell.value = 'Муниципальное образование'
                    cell = worksheet.cell(row=1, column=8)
                    cell.value = 'Организация'
                    cell = worksheet.cell(row=1, column=9)
                    cell.value = 'Категория должности'
                    cell = worksheet.cell(row=1, column=10)
                    cell.value = 'Должность'
                    cell = worksheet.cell(row=1, column=11)
                    cell.value = 'Уровень образования'
                    cell = worksheet.cell(row=1, column=12)
                    cell.value = 'Категория получаемого образования'
                    cell = worksheet.cell(row=1, column=13)
                    cell.value = 'Фамилия в дипломе'
                    cell = worksheet.cell(row=1, column=14)
                    cell.value = 'Cерия документа об образовании'
                    cell = worksheet.cell(row=1, column=15)
                    cell.value = 'Номер документа об образовании'
                    cell = worksheet.cell(row=1, column=16)
                    cell.value = 'Дата выдачи документа об образовании'
                    cell = worksheet.cell(row=1, column=17)
                    cell.value = 'Оплата'
                    cell = worksheet.cell(row=1, column=18)
                    cell.value = 'Получение удостоверения почтой'
                    cell = worksheet.cell(row=1, column=19)
                    cell.value = 'Физический адрес доставки удостоверения'
                    forms = CoursesForms.objects.filter(group_id=request.POST.get('courses_forms')).select_related('profile')
                    row_num = 2
                    for form in forms:
                        cell = worksheet.cell(row=row_num, column=1)
                        cell.value = form.profile.phone
                        cell = worksheet.cell(row=row_num, column=2)
                        cell.value = form.profile.user.email
                        cell = worksheet.cell(row=row_num, column=3)
                        cell.value = form.profile.surname
                        cell = worksheet.cell(row=row_num, column=4)
                        cell.value = form.profile.name
                        cell = worksheet.cell(row=row_num, column=5)
                        cell.value = form.profile.patronymic
                        cell = worksheet.cell(row=row_num, column=6)
                        cell.value = form.region.name
                        cell = worksheet.cell(row=row_num, column=7)
                        if form.mo is None:
                            cell.value = '-'
                        else:
                            cell.value = form.mo.name
                        cell = worksheet.cell(row=row_num, column=8)
                        if form.oo is None:
                            if form.oo_new == '':
                                cell.value = '-'
                            else:
                                cell.value = form.oo_new
                        else:
                            cell.value = form.oo.full_name
                        cell = worksheet.cell(row=row_num, column=9)
                        if form.position_cat is None:
                            cell.value = '-'
                        else:
                            cell.value = form.position_cat.name
                        cell = worksheet.cell(row=row_num, column=10)
                        if form.position is None:
                            cell.value = '-'
                        else:
                            cell.value = form.position.name
                        cell = worksheet.cell(row=row_num, column=11)
                        if form.edu_level is None:
                            cell.value = '-'
                        else:
                            cell.value = form.edu_level.name
                        cell = worksheet.cell(row=row_num, column=12)
                        if form.edu_cat is None:
                            cell.value = '-'
                        else:
                            cell.value = form.edu_cat.name
                        cell = worksheet.cell(row=row_num, column=13)
                        if form.check_surname == '':
                            cell.value = '-'
                        else:
                            cell.value = form.check_surname
                        cell = worksheet.cell(row=row_num, column=14)
                        if form.edu_serial == '':
                            cell.value = '-'
                        else:
                            cell.value = form.edu_serial
                        cell = worksheet.cell(row=row_num, column=15)
                        if form.edu_number == '':
                            cell.value = '-'
                        else:
                            cell.value = form.edu_number
                        cell = worksheet.cell(row=row_num, column=16)
                        if form.edu_date is None:
                            cell.value = '-'
                        else:
                            cell.value = form.edu_date.strftime('%d.%m.%Y')
                        cell = worksheet.cell(row=row_num, column=17)
                        if form.type is True:
                            cell.value = 'Физическое лицо'
                        else:
                            cell.value = 'Юридическое лицо'
                        cell = worksheet.cell(row=row_num, column=18)
                        if form.cert_mail is True:
                            cell.value = 'Да'
                        else:
                            cell.value = 'Нет'
                        cell = worksheet.cell(row=row_num, column=19)
                        if form.address == '':
                            cell.value = '-'
                        else:
                            cell.value = form.address
                        row_num += 1
                    workbook.save(response)
                    return response
                elif 'events_forms' in request.POST:
                    group = StudentGroups.objects.get(id=int(request.POST.get('events_forms')))
                    strin = 'Анкеты_' + group.code + '.xlsx'
                    response = HttpResponse(
                        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    )
                    response['Content-Disposition'] = 'attachment; filename="' + escape_uri_path(strin) + '"'
                    workbook = Workbook()
                    worksheet = workbook.active
                    worksheet.title = 'Анкеты'
                    cell = worksheet.cell(row=1, column=1)
                    cell.value = 'Телефон'
                    cell = worksheet.cell(row=1, column=2)
                    cell.value = 'Email'
                    cell = worksheet.cell(row=1, column=3)
                    cell.value = 'Фамилия'
                    cell = worksheet.cell(row=1, column=4)
                    cell.value = 'Имя'
                    cell = worksheet.cell(row=1, column=5)
                    cell.value = 'Отчество'
                    cell = worksheet.cell(row=1, column=6)
                    cell.value = 'Регион'
                    cell = worksheet.cell(row=1, column=7)
                    cell.value = 'Муниципальное образование'
                    cell = worksheet.cell(row=1, column=8)
                    cell.value = 'Организация'
                    cell = worksheet.cell(row=1, column=9)
                    cell.value = 'Категория должности'
                    cell = worksheet.cell(row=1, column=10)
                    cell.value = 'Должность'
                    cell = worksheet.cell(row=1, column=11)
                    cell.value = 'Оплата'
                    forms = EventsForms.objects.filter(group_id=request.POST.get('events_forms')).select_related('profile')
                    row_num = 2
                    for form in forms:
                        cell = worksheet.cell(row=row_num, column=1)
                        cell.value = form.profile.phone
                        cell = worksheet.cell(row=row_num, column=2)
                        cell.value = form.profile.user.email
                        cell = worksheet.cell(row=row_num, column=3)
                        cell.value = form.profile.surname
                        cell = worksheet.cell(row=row_num, column=4)
                        cell.value = form.profile.name
                        cell = worksheet.cell(row=row_num, column=5)
                        cell.value = form.profile.patronymic
                        cell = worksheet.cell(row=row_num, column=6)
                        cell.value = form.region.name
                        cell = worksheet.cell(row=row_num, column=7)
                        if form.mo is None:
                            cell.value = '-'
                        else:
                            cell.value = form.mo.name
                        cell = worksheet.cell(row=row_num, column=8)
                        if form.oo is None:
                            if form.oo_new == '':
                                cell.value = '-'
                            else:
                                cell.value = form.oo_new
                        else:
                            cell.value = form.oo.full_name
                        cell = worksheet.cell(row=row_num, column=9)
                        if form.position_cat is None:
                            cell.value = '-'
                        else:
                            cell.value = form.position_cat.name
                        cell = worksheet.cell(row=row_num, column=10)
                        if form.position is None:
                            cell.value = '-'
                        else:
                            cell.value = form.position.name
                        cell = worksheet.cell(row=row_num, column=11)
                        if form.type is True:
                            cell.value = 'Физическое лицо'
                        else:
                            cell.value = 'Юридическое лицо'
                        row_num += 1
                    workbook.save(response)
                    return response
                elif 'approve_list' in request.POST:
                    group = StudentGroups.objects.get(id=request.POST.get('approve_list'))
                    group.status = StGroupStatuses.objects.get(name='Ожидает подгрузку скана договора оферты')
                    group.save()
                    messages.success(request, 'Состав группы успешно утвержден')
                elif 'change_list' in request.POST:
                    group = StudentGroups.objects.get(id=request.POST.get('change_list'))
                    if group.event is None:
                        group.status = StGroupStatuses.objects.get(name='Идет регистрация')
                        group.save()
                        messages.success(request, 'Регистрация в учебную группу успешно открыта')
                    else:
                        count = group.students.count()
                        if count >= group.students_number:
                            messages.error(request, 'В учебной группе достигнуто плановое количество мест')
                        else:
                            group.status = StGroupStatuses.objects.get(name='Идет регистрация')
                            group.save()
                            messages.success(request, 'Регистрация в учебную группу успешно открыта')
                elif 'order_enroll' in request.POST:
                    group = StudentGroups.objects.get(id=request.POST.get('order_enroll'))
                    start = group.course.date_start
                    finish = group.course.date_finish
                    doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/order_enroll_ou.docx")
                    context = {
                        'prog_name': group.course.program.name,
                        'duration': group.course.program.duration,
                        'day_start': start.strftime('%d'),
                        'month_start': month_from_ru_to_eng(start.strftime('%B')),
                        'year_start': start.strftime('%Y'),
                        'day_finish': finish.strftime('%d'),
                        'month_finish': month_from_ru_to_eng(finish.strftime('%B')),
                        'year_finish': finish.strftime('%Y'),
                        'students_count': group.students.count(),
                        'code': group.code,
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_зачисление.docx")
                    document = Document(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_зачисление.docx")
                    tbl = document.tables[0]
                    for index, student in enumerate(group.students.all()):
                        row_cells = tbl.add_row().cells
                        row_cells[0].text = str(index+1)
                        row_cells[1].text = student.surname+' '+student.name+' '+student.patronymic
                    for row in tbl.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.name = 'Times New Roman'
                                    font.size = Pt(12)
                    document.save(MEDIA_ROOT + "\\Приказы\\ОУ\\Приказ_зачисление_" + group.code + ".docx")
                    filename = MEDIA_ROOT + "\\Приказы\\ОУ\\Приказ_зачисление_" + group.code + ".docx"
                    strin = 'Приказ_зачисление_' + group.code + '.docx'
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'order_expulsion' in request.POST:
                    group = StudentGroups.objects.get(id=request.POST.get('order_expulsion'))
                    finish = group.course.date_finish
                    if finish.strftime('%w') == '6':
                        finish = finish + timedelta(days=2)
                    if finish.strftime('%w') == '0':
                        finish = finish + timedelta(days=1)
                    doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/order_expulsion_ou.docx")
                    context = {
                        'prog_name': group.course.program.name,
                        'duration': group.course.program.duration,
                        'day_exp': finish.strftime('%d'),
                        'month_exp': month_from_ru_to_eng(finish.strftime('%B')),
                        'year_exp': finish.strftime('%Y'),
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx")
                    document = Document(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx")
                    tbl = document.tables[0]
                    for index, student in enumerate(group.students.all()):
                        row_cells = tbl.add_row().cells
                        row_cells[0].text = str(index+1)+'.'
                        row_cells[1].text = student.surname+' '+student.name+' '+student.patronymic
                    for row in tbl.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.name = 'Times New Roman'
                                    font.size = Pt(12)
                    document.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx")
                    filename = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Приказ_отчисление.docx"
                    strin = 'Приказ_отчисление_' + group.code + '.docx'
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'delete_gr' in request.POST:
                    StudentGroups.objects.get(id=request.POST.get('delete_gr')).delete()
                    messages.success(request, 'Группа успешно удалена')
                elif 'close_study' in request.POST:
                    group = StudentGroups.objects.get(id=request.POST.get('close_study'))
                    group.survey_show = True
                    group.status_id = StGroupStatuses.objects.get(name='Обучение завершено').id
                    group.save()
                    messages.success(request, 'Обучение в группе успешно завершено')
                elif 'journal_course' in request.POST:
                    if not CourseLessons.objects.filter(group_id=request.POST.get('journal_course')).exists():
                        messages.error(request, 'Не найдены занятия в расписании учебной группы')
                        return HttpResponseRedirect('/centre/study/studentgroups')
                    if not CourseLessons.objects.filter(group_id=request.POST.get('journal_course')).\
                            exclude(stschedule__in=StSchedule.objects.filter(control_form='')).exists():
                        messages.error(request, 'Не найдены занятия с формой контроля в расписании учебной группы')
                        return HttpResponseRedirect('/centre/study/studentgroups')
                    group = StudentGroups.objects.get(id=request.POST.get('journal_course'))
                    dep = group.course.program.department
                    list = re.split(' |-', dep)
                    short_dep = ''
                    for el in list:
                        short_dep += el[:1].upper()
                    date_start = group.course.date_start
                    date_finish = group.course.date_finish
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        type_dpp = 'повышения квалификации'
                    else:
                        type_dpp = 'профессиональной переподготовки'
                    cats = ''
                    for cat in group.course.program.categories.all():
                        cats += cat.name+'; '
                    deps = GetDepsWithManagerFromAD()
                    for key, value in deps.items():
                        if key == dep:
                            manager = value
                            fio = value.split(' ')
                            io_family_manager = fio[1][:1] + '.' + fio[2][:1] + '. ' + fio[0]
                            break
                    total_lecture = total_practice = total_individual = total_trainee = total = 0
                    lessons = CourseLessons.objects.filter(group_id=group.id)
                    for les in lessons:
                        total_lecture += les.lecture_hours
                        total_practice += les.practice_hours
                        total_trainee += les.trainee_hours
                        total_individual += les.individual_hours
                        total += les.lecture_hours + les.practice_hours + les.individual_hours + les.trainee_hours
                    path = STATIC_ROOT + '\\doc_templates\\xlsx\\journal_course.xlsx'
                    writer = BookWriter(path)
                    info = {
                        'type_dpp': type_dpp,
                        'name_dpp': group.course.program.name,
                        'duration': group.course.program.duration,
                        'place': group.course.place,
                        'code': group.code,
                        'dep': dep,
                        'cats': cats,
                        'manager': manager,
                        'day_start': date_start.strftime('%d'),
                        'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                        'year_start': date_start.strftime('%Y'),
                        'day_finish': date_finish.strftime('%d'),
                        'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                        'year_finish': date_finish.strftime('%Y'),
                        'io_family_manager': io_family_manager
                    }
                    info2 = {}
                    info2['students'] = get_students(group.id)
                    if group.study_form != 'Без использования ДОТ':
                        title = 'учебной аудиторной нагрузки,\nв т.ч. с использованием электронного обучения  и дистанционных образовательных технологий'
                    else:
                        title = 'учебной аудиторной нагрузки'
                    info3 = {
                        'total_lecture': total_lecture,
                        'total_practice': total_practice,
                        'total_trainee': total_trainee,
                        'total_individual': total_individual,
                        'total': total,
                        'short_dep': short_dep,
                        'manager': io_family_manager,
                        'title': title
                    }
                    info3['teachers'] = get_teachers(group.id)
                    writer.render_sheet(info, 'тит', 0)
                    writer.render_sheet(info2, 'список', 1)
                    writer.render_sheet(info3, 'часы ауд', 2)
                    count = 3
                    exams = lessons.exclude(stschedule__in=StSchedule.objects.filter(control_form='')).order_by('lesson_time_start')
                    newpath = MEDIA_ROOT + '\\Журналы\\' + group.code
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    for index, ex in enumerate(exams):
                        date_lesson = ex.lesson_time_start.date()
                        indexes = [m.start() for m in re.finditer(' ', ex.stschedule.name)]
                        info_pa = {
                            'theme': ex.stschedule.name[indexes[1]:],
                            'cats': cats,
                            'lector': ex.teacher.surname+' '+ex.teacher.name+' '+ex.teacher.patronymic,
                            'day_lesson': date_lesson.strftime('%d'),
                            'month_lesson': month_from_ru_to_eng(date_lesson.strftime('%B')),
                            'year_lesson': date_lesson.strftime('%Y'),
                            'control_form': ex.stschedule.control_form,
                            'manager': manager
                        }
                        info_pa['students'] = get_students(group.id)
                        if index < len(exams)-1:
                            writer.render_sheet(info_pa, 'зачетная ведомость ПА '+str(count-2), 3)
                        else:
                            writer.render_sheet(info_pa, 'зачетная ведомость ИА ', 4)
                        count += 1
                    writer.save(MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx")
                    wb = openpyxl.load_workbook(MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx")
                    ws = wb.worksheets[0]
                    img = openpyxl.drawing.image.Image(STATIC_ROOT + '\\doc_templates\\xlsx\\logo.png')
                    img.anchor = 'C1'
                    ws.add_image(img)
                    wb.save(MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx")
                    filename = MEDIA_ROOT + '\\Журналы\\' + group.code + "\\Журнал.xlsx"
                    strin = "Журнал_" + group.code + ".xlsx"
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'journal_event' in request.POST:
                    group = StudentGroups.objects.get(id=request.POST.get('journal_event'))
                    dep = group.event.department
                    list = re.split(' |-', dep)
                    short_dep = ''
                    for el in list:
                        short_dep += el[:1].upper()
                    date_start = group.event.date_start
                    date_finish = group.event.date_finish
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(group.event.type.name)[0]
                    word = type.inflect({'loct'})
                    cats = ''
                    for cat in group.event.categories.all():
                        cats += cat.name + '; '
                    deps = GetDepsWithManagerFromAD()
                    for key, value in deps.items():
                        if key == dep:
                            manager = value
                            fio = value.split(' ')
                            io_family_manager = fio[1][:1] + '.' + fio[2][:1] + '. ' + fio[0]
                            break
                    total_lecture = total_practice = total = 0
                    lessons = EventsLessons.objects.filter(group_id=group.id)
                    for les in lessons:
                        total_lecture += les.lecture_hours
                        total_practice += les.practice_hours
                        total += les.lecture_hours + les.practice_hours
                    path = STATIC_ROOT + '\\doc_templates\\xlsx\\journal_event.xlsx'
                    writer = BookWriter(path)
                    info = {
                        'type_event': word.word,
                        'event_name': group.event.name,
                        'code': group.code,
                        'place': group.event.place,
                        'dep': dep,
                        'cats': cats,
                        'manager': manager,
                        'day_start': date_start.strftime('%d'),
                        'month_start': month_from_ru_to_eng(date_start.strftime('%B')),
                        'year_start': date_start.strftime('%Y'),
                        'day_finish': date_finish.strftime('%d'),
                        'month_finish': month_from_ru_to_eng(date_finish.strftime('%B')),
                        'year_finish': date_finish.strftime('%Y'),
                    }
                    writer.render_sheet(info, 'тит', 0)
                    info2 = {}
                    info2['students'] = get_students(group.id)
                    writer.render_sheet(info2, 'Список', 1)
                    info3 = {}
                    total_l = total_p = total = 0
                    for lesson in EventsLessons.objects.filter(group_id=group.id):
                        total_l += lesson.lecture_hours
                        total_p += lesson.practice_hours
                        total += lesson.lecture_hours + lesson.practice_hours
                    info3['teachers'] = get_teachers(group.id)
                    info3['total_lecture'] = total_l
                    info3['total_practice'] = total_p
                    info3['total'] = total
                    writer.render_sheet(info3, 'часы ауд', 2)
                    info4 = {}
                    forms = EventsForms.objects.filter(group_id=group.id)
                    total_mo = 0
                    for mo in Mos.objects.all():
                        if forms.filter(mo_id=mo.id).exists():
                            info4[mo.tpl_name] = forms.filter(mo_id=mo.id).count()
                            total_mo += forms.filter(mo_id=mo.id).count()
                        else:
                            info4[mo.tpl_name] = '0'
                    info4['total_mo'] = str(total_mo)
                    total_oo = 0
                    for type in OoTypes.objects.all():
                        if forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).exists():
                            info4[type.tpl_name] = forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).count()
                            total_oo += forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).count()
                        else:
                            info4[type.tpl_name] = '0'
                    info4['total_oo'] = str(total_oo)
                    total_pos = 0
                    for type_pos in PositionCategories.objects.all():
                        if type_pos.tpl_name not in ['spec_1', 'spec_2']:
                            if forms.filter(position_cat_id=type_pos.id).exists():
                                info4[type_pos.tpl_name] = forms.filter(position_cat_id=type_pos.id).count()
                                total_pos += forms.filter(oo__in=Oos.objects.filter(type_oo_id=type.id)).count()
                            else:
                                info4[type_pos.tpl_name] = '0'
                    spec = 0
                    if forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_1').id).exists():
                        spec += forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_1').id).count()
                        total_pos += forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_1').id).count()
                    if forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_2').id).exists():
                        spec += forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_2').id).count()
                        total_pos += forms.filter(position_cat_id=PositionCategories.objects.get(tpl_name='spec_2').id).count()
                    info4['spec'] = spec
                    info4['total_pos'] = total_pos
                    writer.render_sheet(info4, 'cтат', 3)
                    newpath = MEDIA_ROOT + '\\Журналы\\' + group.code
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    writer.save(newpath + "\\Журнал.xlsx")
                    filename = newpath + "\\Журнал.xlsx"
                    strin = "Журнал_" + group.code + ".xlsx"
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'close_doc' in request.POST:
                    group = StudentGroups.objects.select_related('course').get(id=request.POST.get('close_doc'))
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        type_dpp = 'повышения квалификации'
                    else:
                        type_dpp = 'профессиональной переподготовки'
                    start = group.course.date_start
                    finish = group.course.date_finish
                    price = group.course.program.price
                    count = sum = 0
                    for student in group.students.all():
                        form = CoursesForms.objects.filter(group_id=group.id).get(profile_id=student.id)
                        if form.type is True:
                            count += 1
                            sum += price
                    doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/close_doc.docx")
                    context = {
                        'type_dpp': type_dpp,
                        'prog_name': group.course.program.name,
                        'duration': group.course.program.duration,
                        'date_start': group.course.date_start.strftime('%d.%m.%Y'),
                        'date_finish': group.course.date_finish.strftime('%d.%m.%Y'),
                        'code': group.code,
                        'students_count': count,
                        'sum': sum
                    }
                    doc.render(context)
                    newpath = MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    doc.save(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Закрывной.docx")
                    document = Document(MEDIA_ROOT + "\\Приказы\\ОУ\\" + group.code + "\\Закрывной.docx")
                    tbl = document.tables[1]
                    id = 1
                    for index, student in enumerate(group.students.all()):
                        form = CoursesForms.objects.filter(group_id=group.id).get(profile_id=student.id)
                        if form.type is True:
                            row_cells = tbl.add_row().cells
                            row_cells[0].text = str(id)
                            row_cells[1].text = student.surname + ' ' + student.name + ' ' + student.patronymic
                            id += 1
                    for row in tbl.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.name = 'Times New Roman'
                                    font.size = Pt(11)
                    document.save(MEDIA_ROOT + "\\Приказы\\ОУ\\Закрывной_" + group.code + ".docx")
                    filename = MEDIA_ROOT + "\\Приказы\\ОУ\\Закрывной_" + group.code + ".docx"
                    strin = 'Закрывной_' + group.code + '.docx'
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                elif 'print_certs' in request.POST:
                    stgr = StudentGroups.objects.prefetch_related('students').get(id=request.POST.get('print_certs'))
                    dict_studs = {}
                    for student in stgr.students.all().order_by('surname'):
                        l = []
                        l.append(student.surname + ' ' + student.name + ' ' + student.patronymic)
                        if StudentsCerts.objects.filter(group_id=stgr.id).filter(student_id=student.id).exists():
                            cert = StudentsCerts.objects.filter(group_id=stgr.id).get(student_id=student.id)
                            l.append(cert.reg_number)
                            l.append(cert.blank_serial)
                            l.append(cert.blank_number)
                        else:
                            l.append('')
                            l.append('')
                            l.append('')
                        dict_studs[student.id] = l
                    return render(request, 'centre/study/studentgroups/print_certs.html', {
                        'stgr': stgr,
                        'students': dict_studs
                    })
                elif 'print_file' in request.POST or 'to_print' in request.POST:
                    if 'print_file' in request.POST:
                        stgr = StudentGroups.objects.select_related('course').get(id=request.POST.get('print_file'))
                    else:
                        stgr = StudentGroups.objects.select_related('course').get(id=request.POST.get('to_print'))
                    crs = Courses.objects.select_related('program').get(id=stgr.course.id)
                    name_dpp = crs.program.name
                    duration = crs.program.duration
                    date_give = crs.date_finish.strftime('%d.%m.%Y')
                    day_start = crs.date_start.strftime('%d')
                    month_start = month_from_ru_to_eng(crs.date_start.strftime('%B'))
                    year_start = crs.date_start.strftime('%Y')
                    day_finish = crs.date_finish.strftime('%d')
                    month_finish = month_from_ru_to_eng(crs.date_finish.strftime('%B'))
                    year_finish = crs.date_finish.strftime('%Y')
                    data = GetDepsWithManagerFromAD()
                    department = stgr.course.program.department
                    for key, value in data.items():
                        if key == department:
                            manager = value
                            dep = key
                            break
                    spl = manager.split(' ')
                    dep_manager = spl[1][:1] + '.' + spl[2][:1] + '. ' + spl[0]
                    main = docx.Document(STATIC_ROOT + "/doc_templates/docx/ou/blank.docx")
                    newpath = MEDIA_ROOT + "\\Удостоверения\\" + stgr.code + "\\"
                    if not os.path.exists(newpath):
                        os.makedirs(newpath)
                    for stud in stgr.students.all().order_by('-surname'):
                        cert = StudentsCerts.objects.filter(group_id=stgr.id).get(student_id=stud.id)
                        doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/ou/cert_blank.docx")
                        stud = Profiles.objects.get(id=stud.id)
                        fio = stud.surname + ' ' + stud.name + ' ' + stud.patronymic
                        context = {
                            'fio': fio,
                            'name_dpp': name_dpp,
                            'duration': duration,
                            'reg_number': cert.reg_number,
                            'date_give': date_give,
                            'day_start': day_start,
                            'month_start': month_start,
                            'year_start': year_start,
                            'day_finish': day_finish,
                            'month_finish': month_finish,
                            'year_finish': year_finish,
                            'dep_manager': dep_manager
                        }
                        doc.render(context)
                        doc.save("new_cert.docx")
                        composer = Composer(main)
                        # filename_second_docx is the name of the second docx file
                        doc2 = docx.Document("new_cert.docx")
                        # append the doc2 into the master using composer.append function
                        composer.append(doc2)
                        # Save the combined docx with a name
                        composer.save(newpath + "Печать_" + stgr.code + ".docx")
                        os.remove("new_cert.docx")
                    if 'to_print' in request.POST:
                        email = EmailMessage(
                            "АИС «Учебный центр»: Файл печати удостоверений",
                            "Во вложении находится файл для печати удостоверений.\nДля более подробной информации"
                            " обратитесь к Загвоздиной Марине Николаевне (доб. 249)",
                            None,
                            ['print@coko38.ru',],
                        )
                        email.attach_file(newpath + "Печать_" + stgr.code + ".docx")
                        email.send()
                        messages.success(request, 'Файл успешно отправлен в Типографию')
                        return HttpResponseRedirect('/centre/study/studentgroups')
                    filename = newpath + "Печать_" + stgr.code + ".docx"
                    strin = "Печать_" + stgr.code + ".docx"
                    data = open(filename, "rb").read()
                    response = HttpResponse(data,
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Length'] = os.path.getsize(filename)
                    response['Content-Disposition'] = "attachment; filename=" + escape_uri_path(strin)
                    return response
                else:
                    str_teacher = request.POST.get('teacher')
                    group = StudentGroups.objects.get(id=request.POST.get('curator'))
                    group.curator_id = int(str_teacher[str_teacher.find(':')+1:str_teacher.find(')')])
                    group.save()
                    messages.success(request, 'Куратор успешно назначен')
        #except BaseException:
            #messages.error(request, 'Произошла ошибка, повторите попытку позже')
            return HttpResponseRedirect('/centre/study/studentgroups')


class StudentGroupCreate(CheckAdminMixin, CreateView):
    login_url = '/'
    model = StudentGroups
    context_object_name = 'group'
    template_name = 'centre/study/studentgroups/create.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'id_type' in request.GET:
                typ = uri_to_iri(request.GET.get('id_type'))
                flag = True
                if typ in ['Повышение квалификации', 'Профессиональная переподготовка']:
                    courses = Courses.objects.filter(program__in=Programs.objects.filter(type_dpp=typ)).order_by('-id')
                    if courses.count() != 0:
                        voc = {}
                        for course in courses:
                            voc[course.id] = course.program.name
                    else:
                        voc = None
                else:
                    flag = False
                    events = Events.objects.filter(type_id=int(request.GET.get('id_type'))).order_by('-id')
                    if events.count() != 0:
                        voc = {}
                        for event in events:
                            voc[event.id] = event.name
                    else:
                        voc = None
                return JsonResponse({
                    'voc': voc,
                    'type': flag
                })
        return render(request, 'centre/study/studentgroups/create.html', {
            'courses': Courses.objects.all().order_by('-id')[:50],
            'events': Events.objects.all().order_by('-id')[:50],
        })

    def post(self, request, *args, **kwargs):
        try:
            if 'type' in request.POST:
                type = request.POST.get('type')
                if type in ['Повышение квалификации', 'Профессиональная переподготовка']:
                    if type == 'Повышение квалификации':
                        typ = 'Курс повышения квалификации (онлайн)'
                    else:
                        typ = 'Курс профессиональной переподготовки (онлайн)'
                    crs = request.POST.get('course')
                    dep_name = Programs.objects.get(id=Courses.objects.get(id=crs).program_id).department
                    if Courses.objects.filter(program__in=Programs.objects.filter(department=dep_name)).exists():
                        count = Courses.objects.filter(program__in=Programs.objects.filter(department=dep_name)).count()
                    else:
                        count = 0
                    list = re.split(' |-', dep_name)
                    dep = ''
                    for el in list:
                        dep += el[:1].upper()
                    date_start = Courses.objects.get(id=crs).date_start
                    date_service = date_start.strftime('%B')+' '+str(date_start.year)+' года'
                    start = date_start.strftime('%d.%m.%Y')
                    finish = Courses.objects.get(id=crs).date_finish.strftime('%d.%m.%Y')
                    name = Courses.objects.get(id=crs).program.name
                    duration = Courses.objects.get(id=crs).program.duration
                    code = CourseGroupCode(dep_name, str(count+1), date_start, 'course')
                    new_group = StudentGroups()
                    new_group.code = code
                    new_group.course_id = crs
                    new_group.curator = None
                else:
                    typ = EventTypes.objects.get(id=type).name+' (онлайн)'
                    vnt = request.POST.get('event')
                    dep_name = Events.objects.get(id=vnt).department
                    if Events.objects.filter(department=dep_name).exists():
                        count = Events.objects.filter(department=dep_name).count()
                    else:
                        count = 0
                    list = re.split(' |-', dep_name)
                    dep = ''
                    for el in list:
                        dep += el[:1].upper()
                    date_start = Events.objects.get(id=vnt).date_start
                    date_service = date_start.strftime('%B') + ' ' + str(date_start.year) + ' года'
                    start = date_start.strftime('%d.%m.%Y')
                    finish = Events.objects.get(id=vnt).date_finish.strftime('%d.%m.%Y')
                    name = Events.objects.get(id=vnt).name
                    duration = Events.objects.get(id=vnt).duration
                    code = CourseGroupCode(dep_name, str(count+1), date_start, 'event')
                    new_group = StudentGroups()
                    new_group.code = code
                    new_group.event_id = vnt
                    new_group.curator = None
                new_group.students_number = int(request.POST.get('studentsnumber'))
                new_group.study_form = 'Без использования ДОТ'
                new_group.save()
                messages.success(request, 'Группа успешно создана')
                url = 'https://edu-dev.coko38.ru/student/detail/' + str(new_group.id)
                doc = DocxTemplate(STATIC_ROOT + "/doc_templates/docx/create_stgroup/email.docx")
                context = {
                    'date_service': date_service,
                    'dep': dep,
                    'type': typ,
                    'name': name,
                    'url': url,
                    'duration': duration,
                    'start': start,
                    'finish': finish,
                }
                doc.render(context)
                doc.save(MEDIA_ROOT + "\\Технические файлы\\Email\\" + code + ".docx")
                prof = Profiles.objects.get(user_id=request.user.id)
                fio = prof.surname+' '+prof.name+' '+prof.patronymic
                email = EmailMessage(
                    "АИС «Учебный центр»: Новая группа",
                    "В АИС была добавлена новая группа пользователем '"+prof+"'.\nИнформация на сайт во вложении",
                    None,
                    ['edu@coko38.ru', ],
                )
                email.attach_file(MEDIA_ROOT + "\\Технические файлы\\Email\\" + code + ".docx")
                email.send()
        except BaseException:
            messages.error(request, 'Произошла ошибка, повторите попытку позже...')
        return HttpResponseRedirect('/centre/study/studentgroups')


class StGrStudentsList(CheckCentreMixin, ListView):
    login_url = '/'
    model = StudentGroups
    context_object_name = 'students'
    template_name = 'centre/study/studentgroups/students_list.html'

    def get(self, request, *args, **kwargs):
        if is_ajax(request=request):
            if 'list_sts' in request.GET:
                apps = Apps.objects.filter(group_id=request.GET.get('group')).\
                    filter(check_diploma_info=False)
                list_st = []
                for app in apps:
                    list_st.append(app.profile_id)
                return JsonResponse({'list_st': list_st})
            elif 'list_pay' in request.GET:
                apps = Apps.objects.filter(group_id=request.GET.get('group')).filter(check_pay=False).\
                    exclude(pay_doc_id=None)
                list_pay = []
                for app in apps:
                    list_pay.append(app.profile_id)
                return JsonResponse({'list_pay': list_pay})
            elif 'list_oos' in request.GET:
                list_oos = []
                group = StudentGroups.objects.select_related('course', 'event').get(id=request.GET.get('group'))
                if group.course is not None:
                    forms = CoursesForms.objects.filter(group_id=group.id).select_related('region')
                else:
                    forms = EventsForms.objects.filter(group_id=group.id).select_related('region')
                for form in forms:
                    if form.region.name == 'Иркутская область' and form.oo_new is not None:
                        list_oos.append(form.id)
                return JsonResponse({
                    'list_oos': list_oos
                })
            elif 'pay_student' in request.GET:
                try:
                    pay_doc = Apps.objects.filter(group_id=request.GET.get('group')).\
                        get(profile_id=request.GET.get('pay_student')).pay_doc_id
                    status = Apps.objects.filter(group_id=request.GET.get('group')).\
                        get(profile_id=request.GET.get('pay_student')).status.name
                    return JsonResponse({'pay_doc': pay_doc, 'status': status})
                except BaseException:
                    return JsonResponse({'pay_doc': None})
            elif 'pay_accept' in request.GET:
                app = Apps.objects.get(id=request.GET.get('pay_accept'))
                app.status_id = Statuses.objects.get(name='Оплачено').id
                app.save()
                group = StudentGroups.objects.get(id=request.GET.get('group'))
                if group.event is None:
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        typ = 'курсе повышения квалификации'
                    else:
                        typ = 'курс профессиональной переподготовки'
                    name = group.course.program.name
                else:
                    type_classic = group.event.type.name
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(type_classic)[0]
                    word = type.inflect({'loct'})
                    typ = word.word
                    name = group.event.name
                mail = Profiles.objects.get(id=app.profile_id).user.email
                EmailAcceptPay.delay(mail, typ, name)
                app.check_pay = True
                app.save()
                messages.success(request, 'Оплата успешно подтверждена')
                return JsonResponse({})
            elif 'pay_denied' in request.GET:
                app = Apps.objects.get(id=request.GET.get('pay_denied'))
                id = app.pay_doc_id
                app.pay_doc_id = None
                app.save()
                Docs.objects.get(id=id).delete()
                app.status_id = Statuses.objects.get(name='Ждем оплату').id
                app.check_pay = False
                app.save()
                group = StudentGroups.objects.get(id=request.GET.get('group'))
                if group.event is None:
                    if group.course.program.type_dpp == 'Повышение квалификации':
                        typ = 'курсе повышения квалификации'
                    else:
                        typ = 'курс профессиональной переподготовки'
                    name = group.course.program.name
                else:
                    type_classic = group.event.type.name
                    morph = pymorphy2.MorphAnalyzer()
                    type = morph.parse(type_classic)[0]
                    word = type.inflect({'loct'})
                    typ = word.word
                    name = group.event.name
                mail = Profiles.objects.get(id=app.profile_id).user.email
                EmailDeniedPay.delay(mail, typ, name, request.GET.get('message'))
                messages.success(request, 'Оплата успешно отклонена')
                return JsonResponse({})
            elif 'pay_doc_id' in request.GET:
                app = Apps.objects.filter(group_id=request.GET.get('group')).get(profile_id=request.GET.get('pay_doc_id'))
                pd_id = app.pay_doc_id
                student = Profiles.objects.get(id=request.GET.get('pay_doc_id'))
                return JsonResponse({
                    'app': app.id,
                    'pd_id': pd_id,
                    'fio': student.surname+' '+student.name+' '+student.patronymic
                })
            elif 'oo_change_id' in request.GET:
                if StudentGroups.objects.get(id=request.GET.get('group')).course is None:
                    form = EventsForms.objects.select_related('mo').get(id=request.GET.get('oo_change_id'))
                else:
                    form = CoursesForms.objects.select_related('mo').get(id=request.GET.get('oo_change_id'))
                oos = Oos.objects.filter(mo_id=form.mo_id).select_related('type_oo').order_by('-id')
                list_oos = []
                for oo in oos:
                    list_oos.append([oo.id, oo.short_name, oo.full_name, oo.type_oo.name])
                student = Profiles.objects.get(id=form.profile_id)
                return JsonResponse({
                    'mo': form.mo.name,
                    'oo_new': form.oo_new,
                    'oos': list_oos,
                    'fio': student.surname + ' ' + student.name + ' ' + student.patronymic
                })
            elif 'selectoo' in request.GET:
                if StudentGroups.objects.get(id=request.GET.get('group')).course is not None:
                    form = CoursesForms.objects.get(id=request.GET.get('form'))
                else:
                    form = EventsForms.objects.get(id=request.GET.get('form'))
                form.oo_new = None
                form.oo_id = request.GET.get('selectoo')
                form.save()
                return JsonResponse({})
            elif 'searchoo' in request.GET:
                list_oos = []
                for oo in Oos.objects.filter(mo_id=Mos.objects.get(name=request.GET.get('searchoo')).id).\
                        filter(full_name__contains=request.GET.get('search')).select_related('type_oo'):
                    list_oos.append([oo.id, oo.short_name, oo.full_name, oo.type_oo.name])
                return JsonResponse({
                    'oos': list_oos,
                })
            else:
                form = CoursesForms.objects.filter(group_id=request.GET.get('group')).\
                    get(profile_id=request.GET.get('student'))
                if form.change_surname is not None:
                    change_sur = form.change_surname.id
                else:
                    change_sur = None
                student = Profiles.objects.get(id=request.GET.get('student'))
                if form.edu_cat is not None:
                    name_educat = form.edu_cat.name
                else:
                    name_educat = None
                return JsonResponse({
                    'fio': student.surname+' '+student.name+' '+student.patronymic,
                    'id_form': form.id,
                    'id_doc': form.edu_doc.id,
                    'surname': form.check_surname,
                    'id_changesur': change_sur,
                    'serial': form.edu_serial,
                    'number': form.edu_number,
                    'date': form.edu_date,
                    'edu_cat': name_educat
                })
        else:
            group = StudentGroups.objects.get(id=request.GET.get('id_group'))
            qs = group.students.all().order_by('surname')
            if 'surname' in request.GET:
                if len(request.GET.get('surname')) > 0:
                    qs = qs.filter(surname__contains=request.GET.get('surname'))
            if 'name' in request.GET:
                if len(request.GET.get('name')) > 0:
                    qs = qs.filter(name__contains=request.GET.get('name'))
            if 'patronymic' in request.GET:
                if len(request.GET.get('patronymic')) > 0:
                    qs = qs.filter(patronymic__contains=request.GET.get('patronymic'))
            if 'email' in request.GET:
                if len(request.GET.get('email')) > 0:
                    qs = qs.filter(user__in=User.objects.filter(email__contains=request.GET.get('email')))
            NoCheckDip = True
            if Apps.objects.filter(group_id=request.GET.get('id_group')).filter(check_diploma_info=False).exists():
                NoCheckDip = False
            NoCheckPay = True
            if Apps.objects.filter(group_id=request.GET.get('id_group')).filter(check_pay=False).\
                    exclude(pay_doc_id=None).exists():
                NoCheckPay = False
            return render(request, 'centre/study/studentgroups/students_list.html', context={
                'students': qs[:25],
                'group': group,
                'NoCheckDip': NoCheckDip,
                'NoCheckPay': NoCheckPay
            })

    def post(self, request, *args, **kwargs):
        if 'newoo' in request.POST:
            oo = Oos()
            oo.mo_id = Mos.objects.get(name=request.POST.get('newoo')).id
            oo.short_name = request.POST.get('short_name')
            oo.full_name = request.POST.get('full_name')
            oo.type_oo_id = OoTypes.objects.get(name=request.POST.get('type_oo')).id
            oo.form = request.POST.get('form')
            oo.save()
            if StudentGroups.objects.get(id=request.POST.get('group')).course is not None:
                form = CoursesForms.objects.get(id=request.POST.get('student_form'))
            else:
                form = EventsForms.objects.get(id=request.POST.get('student_form'))
            form.oo_new = None
            form.oo_id = oo.id
            form.save()
            return JsonResponse({})
        elif 'surname' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.check_surname = request.POST.get('surname')
            form.save()
            return JsonResponse({'check': 'ok'})
        elif 'serial' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.edu_serial = request.POST.get('serial')
            form.save()
            return JsonResponse({'check': 'ok'})
        elif 'number' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.edu_number = request.POST.get('number')
            form.save()
            return JsonResponse({'check': 'ok'})
        elif 'date' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            form.edu_date = request.POST.get('date')
            form.save()
            return JsonResponse({'check': 'ok'})
        elif 'end' in request.POST:
            form = CoursesForms.objects.get(id=request.POST.get('form'))
            app = Apps.objects.filter(group_id=form.group_id).get(profile_id=form.profile_id)
            app.check_diploma_info = True
            app.save()
            return JsonResponse({'check': 'ok'})
        elif 'del_student' in request.POST:
            app = Apps.objects.get(id=request.POST.get('del_student'))
            group = StudentGroups.objects.get(id=app.group_id)
            if group.course is None:
                EventsForms.objects.filter(group_id=group.id).get(profile_id=app.profile_id).delete()
            else:
                CoursesForms.objects.filter(group_id=group.id).get(profile_id=app.profile_id).delete()
            StInGr = group.students.get(id=app.profile_id)
            group.students.remove(StInGr)
            app.delete()
            messages.success(request, 'Обучающийся успешно удален из группы')
            return HttpResponseRedirect('/centre/study/students?id_group='+request.POST.get('id_group'))
        elif 'scan' in request.FILES:
            app = Apps.objects.filter(group_id=request.POST.get('group')).get(profile_id=request.POST.get('student'))
            new_doc = Docs()
            new_doc.profile_id = request.POST.get('student')
            new_doc.doc_type_id = DocsTypes.objects.get(name='Скан удостоверения').id
            new_doc.file = request.FILES.get('scan')
            new_doc.save()
            id_doc = new_doc.id
            app.certificate_id = id_doc
            app.status_id = Statuses.objects.get(name='Архив').id
            app.save()
            messages.success(request, 'Скан удостоверения успешно загружен')
            return JsonResponse({})
        else:
            return HttpResponseRedirect('/access_denied/')


def group_required(*group_names):
   "Проверка на вхождение пользователя в указанную группу"
   def in_groups(u):
       if bool(u.groups.filter(name__in=group_names)):
           return True
       else:
           return False

   return user_passes_test(in_groups, login_url='/access_denied/', redirect_field_name=None)


@login_required(login_url='/')
@group_required('Администраторы')
def doc_view(request):
    file = Programs.objects.get(id=request.GET.get('prog_id')).order_file
    if file.path[-3:] == 'pdf':
        return FileResponse(open(file.path, 'rb'), content_type='application/pdf')
    else:
        return FileResponse(open(file.path, 'rb'))


def month_from_ru_to_eng(month):
    out = ''
    if month == 'Январь': out = 'января'
    if month == 'Декабрь': out = 'декабря'
    if month == 'Февраль': out = 'февраля'
    if month == 'Март': out = 'марта'
    if month == 'Апрель': out = 'апреля'
    if month == 'Май': out = 'мая'
    if month == 'Июнь': out = 'июня'
    if month == 'Июль': out = 'июля'
    if month == 'Август': out = 'августа'
    if month == 'Сентябрь': out = 'сентября'
    if month == 'Октябрь': out = 'октября'
    if month == 'Ноябрь': out = 'ноября'
    if month == 'Декабрь': out = 'декабря'
    return out


def CourseGroupCode(name, number, date, type):
    list = re.split(' |-', name)
    short_name = ''
    for el in list:
        short_name += el[:1].upper()
    month = str(date.month)
    if len(month) == 1:
        month = '0'+month
    year = str(date.year)
    if type == 'course':
        code = short_name+'-ПК'+number+'-'+month+'-'+year[2:]
    else:
        code = short_name + '-С' + number + '-' + month + '-' + year[2:]
    if StudentGroups.objects.filter(code=code).exists():
        count = StudentGroups.objects.filter(code=code).count()
        code = code+'-'+str(count)
    return code


def is_ajax(request):
    return request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest'


@login_required(login_url='/')
@group_required('Администраторы')
def oo_change(request):
    if request.method == 'GET' or request.POST.get('oo_id') is None:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    else:
        oo = Oos.objects.get(id=request.POST.get('oo_id'))
        oo.mo_id = request.POST.get('munobr')
        oo.short_name = request.POST.get('shortname')
        oo.full_name = request.POST.get('fullname')
        oo.type_oo_id = request.POST.get('type_oo')
        oo.form = request.POST.get('form')
        oo.save()
        request.method = 'GET'
        messages.success(request, 'Информация успешно обновлена')
    return HttpResponseRedirect('/centre/guides/oos')


@login_required(login_url='/')
@group_required('Администраторы')
def oo_delete(request):
    if request.method == 'GET' or request.POST.get('oo_id') is None:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    else:
        Oos.objects.get(id=request.POST.get('oo_id')).delete()
        messages.success(request, 'Организация успешно удалена')
    return HttpResponseRedirect('/centre/guides/oos')


@login_required(login_url='/')
@group_required('Администраторы')
def oo_new(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    else:
        new = Oos()
        new.mo_id = request.POST.get('munobr')
        new.short_name = request.POST.get('shortname')
        new.full_name = request.POST.get('fullname')
        new.type_oo_id = request.POST.get('type_oo')
        new.form = request.POST.get('form')
        new.save()
        messages.success(request, 'Организация успешно добавлена!')
    return HttpResponseRedirect('/centre/guides/oos')


@login_required(login_url='/')
@group_required('Администраторы')
def new_pos(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    else:
        new = Positions()
        new.name = request.POST.get('name')
        try:
            new.save()
            messages.success(request, 'Должность успешно добавлена!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, повторите попытку позже')

    return HttpResponseRedirect('/centre/guides/positions')


@login_required(login_url='/')
@group_required('Администраторы')
def delete_pos(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            Positions.objects.get(id=request.POST.get('pos_id')).delete()
            messages.success(request, 'Должность успешно удалена!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/guides/positions')


@login_required(login_url='/')
@group_required('Администраторы')
def new_cat(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        new = PositionCategories()
        new.name = request.POST.get('name')
        request.method = 'GET'
        try:
            new.save()
            messages.success(request, 'Категория успешно добавлена!')
        except BaseException:
            messages.error(request, 'Категория с таким названием уже существует')
    return HttpResponseRedirect('/centre/guides/poscats')


@login_required(login_url='/')
@group_required('Администраторы')
def delete_cat(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            PositionCategories.objects.get(id=request.POST.get('cat_id')).delete()
            messages.success(request, 'Категория успешно удалена!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/guides/poscats')


@login_required(login_url='/')
@group_required('Администраторы')
def new_cat_aud(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        new = AudienceCategories()
        new.name = request.POST.get('name')
        try:
            new.save()
            messages.success(request, 'Категория успешно добавлена!')
        except BaseException:
            messages.error(request, 'Категория с таким названием уже существует')
    return HttpResponseRedirect('/centre/guides/audcats')


@login_required(login_url='/')
@group_required('Администраторы')
def delete_cat_aud(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            AudienceCategories.objects.get(id=request.POST.get('cat_id')).delete()
            messages.success(request, 'Категория успешно удалена!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/guides/audcats')


@login_required(login_url='/')
@group_required('Администраторы')
def choose_cats(request):
    return render(request, 'centre/guides/choose_cats.html')


@login_required(login_url='/')
@group_required('Администраторы')
def UserChangeStatus(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            user = User.objects.get(id=Profiles.objects.get(id=request.POST.get('user_id')).user_id)
            if user.groups.filter(name='Администраторы').exists():
                messages.error(request, 'Выбранный пользователь является администратором системы')
            else:
                prof = Profiles.objects.get(id=request.POST.get('user_id'))
                if prof.teacher is True:
                    prof.teacher = False
                    prof.save()
                    prof.refresh_from_db()
                    messages.success(request, 'Пользователь успешно убран из преподавателей!')
                else:
                    prof.teacher = True
                    prof.save()
                    prof.refresh_from_db()
                    messages.success(request, 'Пользователь успешно добавлен к преподавателям!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/guides/users')


@login_required(login_url='/')
@group_required('Администраторы')
def change_user(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            prof = Profiles.objects.get(id=request.POST.get('id'))
            user = User.objects.get(id=prof.user_id)
            user.email = request.POST.get('email')
            user.save()
            user.refresh_from_db()
            prof.state_id = request.POST.get('state')
            prof.snils = request.POST.get('snils')
            prof.phone = request.POST.get('phone')
            prof.surname = request.POST.get('surname')
            prof.name = request.POST.get('name')
            prof.patronymic = request.POST.get('patronymic')
            if request.POST.get('sex') == 'True':
                prof.sex = True
            else:
                prof.sex = False
            prof.birthday = request.POST.get('birthday')
            prof.save()
            prof.refresh_from_db()
            messages.success(request, 'Информация о пользователе успешно обновлена')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/guides/users')


@login_required(login_url='/')
@group_required('Администраторы')
def delete_user(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            user = User.objects.get(id=Profiles.objects.get(id=request.POST.get('id')).user_id)
            user.delete()
            messages.success(request, 'Пользователь успешно удален!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/guides/users')


@login_required(login_url='/')
@group_required('Администраторы')
def delete_program(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, попробуйте позже')
    else:
        try:
            program = Programs.objects.get(id=request.POST.get('id'))
            program.delete()
            messages.success(request, 'Программа успешно удалена!')
        except BaseException:
            messages.error(request, 'Произошла ошибка, попробуйте позже')
    return HttpResponseRedirect('/centre/study/programs')


@login_required(login_url='/')
@group_required('Администраторы')
def change_kug(request, id):
    dpp = Programs.objects.get(id=id)
    prof = Profiles.objects.get(user_id=request.user.id)
    if len(prof.surname) == 0 or len(prof.name) == 0 or len(prof.patronymic) == 0:
        messages.error(request, 'Пожалуйста, заполните ФИО в профиле для доступа к КУГ')
        return HttpResponseRedirect('/centre/study/programs')
    FIO = prof.surname+' '+prof.name[:1]+'.'+prof.patronymic[:1]+'.'
    UserKugEdit = dpp.kug_on_edit
    if len(UserKugEdit) != 0 and UserKugEdit != FIO:
        messages.error(request, 'КУГ находится на редактировании у пользователя\n "'+UserKugEdit+'"')
        return HttpResponseRedirect('/centre/study/programs')
    dpp.kug_on_edit = prof.surname+' '+prof.name[:1]+'.'+prof.patronymic[:1]+'.'
    dpp.save()
    return render(request, 'centre/study/programs/kug/edit.html', {
        'dpp': Programs.objects.get(id=id)
    })


def CheckSumHours(id, child_id, total, lecture, practice, trainee, individual):
    module = StSchedule.objects.get(id=id)
    if StSchedule.objects.filter(parent_id=id).exists():
        ch_total = ch_lecture = ch_practice = ch_individual = ch_trainee = 0
        childrens = StSchedule.objects.filter(parent_id=id)
        for child in childrens:
            if child_id is not None and int(child_id) == child.id:
                pass
            else:
                ch_total += child.total_hours
                ch_lecture += child.lecture_hours
                ch_practice += child.practice_hours
                ch_trainee += child.trainee_hours
                ch_individual += child.individual_hours
        if module.total_hours < ch_total + total:
            return 'total'
        if module.lecture_hours < ch_lecture + lecture:
            return 'lecture'
        if module.practice_hours < ch_practice + practice:
            return 'practice'
        if module.trainee_hours < ch_trainee + trainee:
            return 'trainee'
        if module.individual_hours < ch_individual + individual:
            return 'individual'
    else:
        if module.total_hours < total:
            return 'total'
        if module.lecture_hours < lecture:
            return 'lecture'
        if module.practice_hours < practice:
            return 'practice'
        if module.trainee_hours < trainee:
            return 'trainee'
        if module.individual_hours < individual:
            return 'individual'
    return None


@login_required(login_url='/')
@group_required('Администраторы')
def add_module(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/programs')
    total = int(request.POST.get('total_hours'))
    lecture = int(request.POST.get('lecture_hours'))
    practice = int(request.POST.get('practice_hours'))
    trainee = int(request.POST.get('trainee_hours'))
    individual = int(request.POST.get('individual_hours'))
    if lecture + practice + individual + trainee != total:
        messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
        return HttpResponseRedirect('/centre/study/kug_'+str(request.POST.get('dpp_id')))
    try:
        count = StSchedule.objects.filter(program_id=int(request.POST.get('dpp_id'))).filter(level=0).count()
        new = StSchedule()
        new.program_id = int(request.POST.get('dpp_id'))
        new.name = 'Раздел '+str(count+1)+'. '+request.POST.get('name')
        new.total_hours = request.POST.get('total_hours')
        new.lecture_hours = request.POST.get('lecture_hours')
        new.practice_hours = request.POST.get('practice_hours')
        new.trainee_hours = request.POST.get('trainee_hours')
        new.individual_hours = request.POST.get('individual_hours')
        new.control_form = request.POST.get('control_form')
        new.save()
        messages.success(request, 'Раздел успешно добавлен')
    except BaseException:
        messages.error(request, 'Произошла ошибка при сохранении, повторите попытку позже')
    return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))


@login_required(login_url='/')
@group_required('Администраторы')
def add_theme(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/programs')
    total = int(request.POST.get('total_hours'))
    lecture = int(request.POST.get('lecture_hours'))
    practice = int(request.POST.get('practice_hours'))
    trainee = int(request.POST.get('trainee_hours'))
    individual = int(request.POST.get('individual_hours'))
    if lecture + practice + individual + trainee != total:
        messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
        return HttpResponseRedirect('/centre/study/kug_'+str(request.POST.get('dpp_id')))
    check = CheckSumHours(int(request.POST.get('module_id')), None, total, lecture, practice, trainee, individual)
    if check is not None:
        if check == 'total':
            messages.error(request, 'Общая сумма часов всех тем превышает указанное в разделе значение')
        elif check == 'lecture':
            messages.error(request, 'Сумма лекционных часов всех тем превышает указанное в разделе значение')
        elif check == 'practice':
            messages.error(request, 'Сумма часов практики всех тем превышает указанное в разделе значение')
        elif check == 'trainee':
            messages.error(request, 'Сумма часов стажировок всех тем превышает указанное в разделе значение')
        else:
            messages.error(request, 'Сумма часов индивидуальных занятий всех тем превышает указанное в разделе значение')
        return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))
    try:
        count = StSchedule.objects.filter(program_id=int(request.POST.get('dpp_id')))\
            .filter(parent_id=int(request.POST.get('module_id'))).count()
        module_number = StSchedule.objects.get(id=int(request.POST.get('module_id'))).name
        new = StSchedule()
        new.program_id = int(request.POST.get('dpp_id'))
        new.name = 'Тема'+module_number[module_number.find(' '):module_number.find('.')]+'.'+str(count+1)+'. '+request.POST.get('name')
        new.total_hours = request.POST.get('total_hours')
        new.lecture_hours = request.POST.get('lecture_hours')
        new.practice_hours = request.POST.get('practice_hours')
        new.trainee_hours = request.POST.get('trainee_hours')
        new.individual_hours = request.POST.get('individual_hours')
        new.control_form = request.POST.get('control_form')
        new.parent_id = int(request.POST.get('module_id'))
        new.save()
        messages.success(request, 'Тема успешно добавлена')
    except BaseException:
        messages.error(request, 'Произошла ошибка при сохранении, повторите попытку позже')
    return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))


@login_required(login_url='/')
@group_required('Администраторы')
def change_module(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/programs')
    try:
        total = int(request.POST.get('total_hours'))
        lecture = int(request.POST.get('lecture_hours'))
        practice = int(request.POST.get('practice_hours'))
        trainee = int(request.POST.get('trainee_hours'))
        individual = int(request.POST.get('individual_hours'))
        if lecture + practice + individual + trainee != total:
            messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
            return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))
        module = StSchedule.objects.get(id=request.POST.get('module_id'))
        name = module.name
        module.name = name[:name.find('.')+2]+request.POST.get('name')
        module.total_hours = request.POST.get('total_hours')
        module.lecture_hours = request.POST.get('lecture_hours')
        module.practice_hours = request.POST.get('practice_hours')
        module.trainee_hours = request.POST.get('trainee_hours')
        module.individual_hours = request.POST.get('individual_hours')
        module.control_form = request.POST.get('control_form')
        module.save()
        messages.success(request, 'Информация о разделе успешно обновлена')
    except BaseException:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))


@login_required(login_url='/')
@group_required('Администраторы')
def change_theme(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/programs')
    total = int(request.POST.get('total_hours'))
    lecture = int(request.POST.get('lecture_hours'))
    practice = int(request.POST.get('practice_hours'))
    trainee = int(request.POST.get('trainee_hours'))
    individual = int(request.POST.get('individual_hours'))
    if lecture + practice + individual + trainee != total:
        messages.error(request, 'Сумма часов занятий не совпадает с общим количеством часов')
        return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))
    id_parent = StSchedule.objects.get(id=int(request.POST.get('theme_id'))).parent_id
    check = CheckSumHours(id_parent, request.POST.get('theme_id'), total, lecture, practice, trainee, individual)
    if check is not None:
        if check == 'total':
            messages.error(request, 'Общая сумма часов всех тем превышает указанное в разделе значение')
        elif check == 'lecture':
            messages.error(request, 'Сумма лекционных часов всех тем превышает указанное в разделе значение')
        elif check == 'practice':
            messages.error(request, 'Сумма часов практики всех тем превышает указанное в разделе значение')
        elif check == 'trainee':
            messages.error(request, 'Сумма часов стажировок всех тем превышает указанное в разделе значение')
        else:
            messages.error(request,
                           'Сумма часов индивидуальных занятий всех тем превышает указанное в разделе значение')
        return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))
    try:
        theme = StSchedule.objects.get(id=request.POST.get('theme_id'))
        name = theme.name
        theme.name = name[:name.find('.', name.find('.') + 1) + 2] + request.POST.get('name')
        theme.total_hours = request.POST.get('total_hours')
        theme.lecture_hours = request.POST.get('lecture_hours')
        theme.practice_hours = request.POST.get('practice_hours')
        theme.trainee_hours = request.POST.get('trainee_hours')
        theme.individual_hours = request.POST.get('individual_hours')
        theme.control_form = request.POST.get('control_form')
        theme.save()
        messages.success(request, 'Информация о теме успешно обновлена')
    except BaseException:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))


@login_required(login_url='/')
@group_required('Администраторы')
def KugDeleteElement(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/programs')
    try:
        el = StSchedule.objects.get(id=int(request.POST.get('id')))
        if StSchedule.objects.filter(parent_id=el.parent_id).count() > 1:
            children = StSchedule.objects.filter(parent_id=el.parent_id)
            for child in children:
                if child.lft > el.rght:
                    name = child.name
                    theme_numb = name[name.find('.'):name.find('.', name.find('.') + 1)]
                    newnumb = int(theme_numb[1:]) - 1
                    newname = name[:name.find('.')] + '.' + str(newnumb) + name[name.find('.', name.find('.') + 1):]
                    child.name = newname
                    #child.save()
        if el.parent_id is None:
            if StSchedule.objects.filter(tree_id__gt=el.tree_id).exists():
                other = StSchedule.objects.filter(tree_id__gt=el.tree_id).filter(level=0)
                for ot in other:
                    children = StSchedule.objects.filter(parent_id=ot.id)
                    for child in children:
                        name = child.name
                        theme_numb = name[name.find(' '):name.find('.')]
                        newnumb = int(theme_numb[1:]) - 1
                        newname = name[:name.find(' ')+1] + str(newnumb) + name[name.find('.'):]
                        child.name = newname
                        child.save()
                    name = ot.name
                    module_numb = name[name.find(' '):name.find('.')]
                    newnumb = int(module_numb) - 1
                    newname = name[:name.find(' ')+1]+str(newnumb)+name[name.find('.'):]
                    ot.name = newname
                    ot.save()
            messages.success(request, 'Раздел успешно удален')
        else:
            messages.success(request, 'Тема успешно удалена')
        StSchedule.objects.get(id=int(request.POST.get('id'))).delete()
    except BaseException:
        messages.error(request, 'Произошла ошибка во время удаления, повторите попытку позже')
    return HttpResponseRedirect('/centre/study/kug_' + str(request.POST.get('dpp_id')))


@login_required(login_url='/')
@group_required('Администраторы')
def KugEditFinish(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/programs')
    try:
        dpp = Programs.objects.get(id=int(request.POST.get('dpp_id')))
        dpp.kug_on_edit = ''
        dpp.save()
    except BaseException:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
    return HttpResponseRedirect('/centre/study/programs')


@login_required(login_url='/')
@group_required('Администраторы')
def planning_choose(request):
    if is_ajax(request=request):
        if 'get_pars' in request.GET:
            dict_pars = {}
            params = PlanningParameters.objects.all()
            for param in params:
                l = []
                l.append(param.value)
                l.append(param.alias)
                dict_pars[param.name] = l
            return JsonResponse({
                'params': dict_pars
            })
        if 'al' in request.GET:
            param = PlanningParameters.objects.get(alias=request.GET.get('al'))
            param.value = request.GET.get('val')
            param.save()
            return JsonResponse({})
    return render(request, 'centre/study/planning/choose.html')


@login_required(login_url='/')
@group_required('Администраторы')
def course_del(request):
    if request.GET:
        messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/')
    else:
        try:
            Courses.objects.get(id=request.POST.get('id')).delete()
            messages.success(request, 'Курс успешно удален')
        except BaseException:
            messages.error(request, 'Произошла ошибка, повторите попытку позже')
        return HttpResponseRedirect('/centre/study/planning/courses')
